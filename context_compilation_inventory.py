"""Bind an exact-context DOCX to its immutable JSONL source snapshot.

The Context Finder JSONL is the identity and text authority.  The companion
DOCX is a presentation of those records, so it cannot carry a durable region
identifier of its own.  :func:`bind_context_compilation` validates the pair in
JSONL order and returns a frozen inventory whose region identities are the
canonical JSONL ``region_id`` values.

This module deliberately performs no live-source validation, model calls, or
subset generation.  It binds two already-created artefacts and fails closed if
their counts, quoted text, metadata, links, query occurrences, highlighting, or
generation marker disagree.
"""

from __future__ import annotations

from dataclasses import dataclass
import hashlib
from io import BytesIO
import json
from pathlib import Path, PurePosixPath
import re
from typing import Any, Mapping, Sequence

from docx import Document  # type: ignore
from docx.enum.text import WD_COLOR_INDEX  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

from context_finder import (  # type: ignore
    COMPILATION_MARKER,
    SCHEMA_VERSION as CONTEXT_FINDER_SCHEMA_VERSION,
    ContextRegion,
    QuerySpec,
    compile_query_pattern,
    validate_query,
)


INVENTORY_SCHEMA_VERSION = "context-compilation-inventory-v1"
METADATA_STYLE = "Context Source Metadata"
HEADING_STYLE = "Heading 1"
BODY_STYLE = "Normal"
COMPILATION_AUTHOR = "AudioProcessor Context Finder"
_HYPERLINK_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)
_SOURCE_NOTE = (
    "Local-source edition. Quoted text is reproduced exactly from the source. "
    "Highlighting marks the requested word or phrase; source links open the "
    "local document on the machine where this compilation was created."
)
_SHA256_RE = re.compile(r"[0-9a-f]{64}")


class ContextCompilationInventoryError(RuntimeError):
    """Base class for an invalid or mismatched compilation pair."""


class ContextCompilationRecordError(ContextCompilationInventoryError):
    """The JSONL record set is malformed or internally inconsistent."""


class ContextCompilationStructureError(ContextCompilationInventoryError):
    """The DOCX is not a structurally valid Context Finder compilation."""


class ContextCompilationPairMismatchError(ContextCompilationInventoryError):
    """The DOCX presentation does not exactly represent the JSONL snapshot."""


@dataclass(frozen=True, slots=True)
class BoundSelectedParagraph:
    """One selected JSONL paragraph, including entries omitted from the DOCX."""

    number: int
    text: str
    page_number: int | None
    emitted: bool


@dataclass(frozen=True, slots=True)
class BoundContextSource:
    """Stored source identity aggregated across all of its context regions."""

    source_relative_path: str
    source_absolute_path: str
    source_sha256: str
    source_target: str
    region_count: int
    occurrence_count: int


@dataclass(frozen=True, slots=True)
class BoundContextRegion:
    """Immutable crosswalk from a canonical JSONL region to one DOCX section.

    ``selected_paragraphs`` preserves empty JSONL snapshots.  ``selected_text``
    is the exact text emitted to Word: non-empty selected paragraphs joined by
    one blank line.  ``selected_paragraphs_sha256`` commits the full structured
    selection, including any empty entries.
    """

    region_id: str
    ordinal: int
    source_region_ordinal: int
    source_relative_path: str
    source_absolute_path: str
    source_sha256: str
    source_target: str
    heading_text: str
    metadata_text: str
    locator: str
    selected_paragraphs: tuple[BoundSelectedParagraph, ...]
    selected_text: str
    selected_text_sha256: str
    selected_paragraphs_sha256: str
    emitted_paragraph_count: int
    empty_selected_paragraph_count: int
    occurrence_count: int
    occurrence_ids: tuple[str, ...]
    selection_start_paragraph: int
    selection_end_paragraph: int
    selection_method: str
    selection_model: str | None
    selection_confidence: float | None
    selection_note: str | None
    query_count: int
    highlight_count: int
    region_fingerprint: str

    @property
    def emitted_paragraphs(self) -> tuple[str, ...]:
        """Return the exact non-empty paragraph sequence represented in Word."""

        return tuple(
            paragraph.text for paragraph in self.selected_paragraphs if paragraph.emitted
        )


@dataclass(frozen=True, slots=True)
class BoundContextCompilation:
    """A content-addressed, immutable Context Finder DOCX/JSONL pair."""

    schema_version: str
    docx_path: Path
    jsonl_path: Path
    docx_sha256: str
    jsonl_sha256: str
    pair_fingerprint: str
    ordered_regions_sha256: str
    source_manifest_sha256: str
    query: str
    query_canonical: str
    query_word_count: int
    root: str
    scanned_files: int
    ignored_generated_files: int
    region_count: int
    occurrence_count: int
    source_count: int
    selected_paragraph_count: int
    emitted_paragraph_count: int
    empty_selected_paragraph_count: int
    highlight_count: int
    sources: tuple[BoundContextSource, ...]
    regions: tuple[BoundContextRegion, ...]


@dataclass(frozen=True, slots=True)
class _RecordsSnapshot:
    manifest: Mapping[str, Any]
    query: QuerySpec
    regions: tuple[ContextRegion, ...]


def bind_context_compilation(
    docx_path: Path | str,
    jsonl_path: Path | str | None = None,
) -> BoundContextCompilation:
    """Validate and bind a Context Finder DOCX to its JSONL companion.

    When ``jsonl_path`` is omitted, the companion is the DOCX path with a
    ``.jsonl`` suffix.  Both files are read once into immutable byte snapshots;
    their SHA-256 hashes therefore identify the exact bytes that were parsed.
    """

    docx = _resolve_input(docx_path, ".docx", "DOCX")
    records_path = (
        docx.with_suffix(".jsonl")
        if jsonl_path is None
        else _resolve_input(jsonl_path, ".jsonl", "JSONL")
    )
    if jsonl_path is None:
        records_path = _resolve_input(records_path, ".jsonl", "JSONL")

    docx_bytes = _read_stable_bytes(docx)
    jsonl_bytes = _read_stable_bytes(records_path)
    docx_sha256 = _sha256_bytes(docx_bytes)
    jsonl_sha256 = _sha256_bytes(jsonl_bytes)

    records = _parse_records(jsonl_bytes)
    document = _load_document(docx_bytes)
    bound_regions = _validate_document(document, records)
    sources = _build_sources(bound_regions)

    manifest = records.manifest
    selected_paragraph_count = sum(
        len(region.selected_paragraphs) for region in bound_regions
    )
    emitted_paragraph_count = sum(
        region.emitted_paragraph_count for region in bound_regions
    )
    empty_selected_paragraph_count = sum(
        region.empty_selected_paragraph_count for region in bound_regions
    )
    highlight_count = sum(region.highlight_count for region in bound_regions)
    occurrence_count = sum(region.occurrence_count for region in bound_regions)

    ordered_regions_sha256 = _sha256_json(
        [region.region_fingerprint for region in bound_regions]
    )
    source_manifest_sha256 = _sha256_json(
        [
            {
                "source_relative_path": source.source_relative_path,
                "source_absolute_path": source.source_absolute_path,
                "source_sha256": source.source_sha256,
                "source_target": source.source_target,
                "region_count": source.region_count,
                "occurrence_count": source.occurrence_count,
            }
            for source in sources
        ]
    )
    counts = {
        "region_count": len(bound_regions),
        "occurrence_count": occurrence_count,
        "source_count": len(sources),
        "selected_paragraph_count": selected_paragraph_count,
        "emitted_paragraph_count": emitted_paragraph_count,
        "empty_selected_paragraph_count": empty_selected_paragraph_count,
        "highlight_count": highlight_count,
    }
    pair_fingerprint = _sha256_json(
        {
            "schema_version": INVENTORY_SCHEMA_VERSION,
            "docx_sha256": docx_sha256,
            "jsonl_sha256": jsonl_sha256,
            "query": {
                "text": records.query.text,
                "canonical": records.query.canonical,
                "word_count": records.query.word_count,
            },
            "counts": counts,
            "ordered_regions_sha256": ordered_regions_sha256,
            "source_manifest_sha256": source_manifest_sha256,
        }
    )

    return BoundContextCompilation(
        schema_version=INVENTORY_SCHEMA_VERSION,
        docx_path=docx,
        jsonl_path=records_path,
        docx_sha256=docx_sha256,
        jsonl_sha256=jsonl_sha256,
        pair_fingerprint=pair_fingerprint,
        ordered_regions_sha256=ordered_regions_sha256,
        source_manifest_sha256=source_manifest_sha256,
        query=records.query.text,
        query_canonical=records.query.canonical,
        query_word_count=records.query.word_count,
        root=str(manifest["root"]),
        scanned_files=_manifest_int(manifest, "scanned_files", minimum=0),
        ignored_generated_files=_manifest_int(
            manifest, "ignored_generated_files", minimum=0
        ),
        region_count=counts["region_count"],
        occurrence_count=counts["occurrence_count"],
        source_count=counts["source_count"],
        selected_paragraph_count=selected_paragraph_count,
        emitted_paragraph_count=emitted_paragraph_count,
        empty_selected_paragraph_count=empty_selected_paragraph_count,
        highlight_count=highlight_count,
        sources=sources,
        regions=bound_regions,
    )


def _resolve_input(path: Path | str, suffix: str, label: str) -> Path:
    candidate = Path(path).expanduser().resolve()
    if candidate.suffix.casefold() != suffix:
        raise ValueError(f"{label} input must use the {suffix} extension")
    if not candidate.is_file():
        raise FileNotFoundError(f"{label} input does not exist: {candidate}")
    return candidate


def _read_stable_bytes(path: Path) -> bytes:
    try:
        before = path.stat()
        data = path.read_bytes()
        after = path.stat()
    except OSError as exc:
        raise ContextCompilationInventoryError(
            f"Could not read compilation input: {path}"
        ) from exc
    if (
        before.st_size != after.st_size
        or before.st_mtime_ns != after.st_mtime_ns
        or len(data) != after.st_size
    ):
        raise ContextCompilationInventoryError(
            f"Compilation input changed while it was being read: {path}"
        )
    return data


def _load_document(data: bytes) -> Any:
    try:
        return Document(BytesIO(data))
    except Exception as exc:
        raise ContextCompilationStructureError(
            "DOCX input is not a readable Word document"
        ) from exc


def _parse_records(data: bytes) -> _RecordsSnapshot:
    try:
        text = data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ContextCompilationRecordError("JSONL input is not valid UTF-8") from exc

    raw_records: list[Mapping[str, Any]] = []
    for line_number, line in enumerate(text.splitlines(), start=1):
        if not line.strip():
            continue
        try:
            record = json.loads(line)
        except json.JSONDecodeError as exc:
            raise ContextCompilationRecordError(
                f"JSONL line {line_number} is not valid JSON"
            ) from exc
        if not isinstance(record, Mapping):
            raise ContextCompilationRecordError(
                f"JSONL line {line_number} must contain an object"
            )
        raw_records.append(record)

    if not raw_records or raw_records[0].get("record_type") != "search_manifest":
        raise ContextCompilationRecordError(
            "Context records must begin with one search_manifest"
        )
    manifest = raw_records[0]
    if manifest.get("schema_version") != CONTEXT_FINDER_SCHEMA_VERSION:
        raise ContextCompilationRecordError(
            f"Unsupported context record schema: {manifest.get('schema_version')!r}"
        )
    if any(record.get("record_type") != "context_region" for record in raw_records[1:]):
        raise ContextCompilationRecordError(
            "Only context_region records may follow the search_manifest"
        )

    query_record = manifest.get("query")
    if not isinstance(query_record, Mapping):
        raise ContextCompilationRecordError("Search manifest query is malformed")
    try:
        query = validate_query(str(query_record["text"]))
    except (KeyError, ValueError) as exc:
        raise ContextCompilationRecordError("Search manifest query is invalid") from exc
    if (
        query_record.get("canonical") != query.canonical
        or query_record.get("word_count") != query.word_count
    ):
        raise ContextCompilationRecordError(
            "Search manifest query fields are internally inconsistent"
        )

    regions: list[ContextRegion] = []
    for ordinal, record in enumerate(raw_records[1:], start=1):
        try:
            region = ContextRegion.from_record(dict(record))
        except (KeyError, TypeError, ValueError) as exc:
            raise ContextCompilationRecordError(
                f"Context region {ordinal} is malformed"
            ) from exc
        regions.append(region)

    expected_regions = _manifest_int(manifest, "region_count", minimum=0)
    if expected_regions != len(regions):
        raise ContextCompilationRecordError(
            "Context region count does not match the search manifest"
        )

    try:
        _validate_record_regions(tuple(regions), query)
    except ContextCompilationRecordError:
        raise
    except (AttributeError, OverflowError, TypeError, ValueError) as exc:
        raise ContextCompilationRecordError(
            "Context region fields contain invalid types or values"
        ) from exc
    occurrence_count = sum(len(region.occurrences) for region in regions)
    source_count = len({region.source_relative_path for region in regions})
    if _manifest_int(manifest, "occurrence_count", minimum=0) != occurrence_count:
        raise ContextCompilationRecordError(
            "Occurrence count does not match the search manifest"
        )
    if _manifest_int(manifest, "source_count", minimum=0) != source_count:
        raise ContextCompilationRecordError(
            "Source count does not match the search manifest"
        )
    if _manifest_int(manifest, "scanned_files", minimum=0) < source_count:
        raise ContextCompilationRecordError(
            "Search manifest scanned_files is smaller than source_count"
        )
    _manifest_int(manifest, "ignored_generated_files", minimum=0)
    if not isinstance(manifest.get("root"), str) or not str(manifest["root"]).strip():
        raise ContextCompilationRecordError("Search manifest root is invalid")

    return _RecordsSnapshot(
        manifest=manifest,
        query=query,
        regions=tuple(regions),
    )


def _validate_record_regions(
    regions: Sequence[ContextRegion],
    query: QuerySpec,
) -> None:
    pattern = compile_query_pattern(query)
    region_ids: set[str] = set()
    occurrence_ids: set[str] = set()
    selected_spans: set[tuple[str, int, int]] = set()
    sources: dict[str, tuple[str, str, str, str]] = {}

    for ordinal, region in enumerate(regions, start=1):
        if not region.region_id or region.region_id in region_ids:
            raise ContextCompilationRecordError(
                f"Context region {ordinal} has a missing or duplicate region_id"
            )
        region_ids.add(region.region_id)
        if region.query != query.text:
            raise ContextCompilationRecordError(
                f"Context region {ordinal} query does not match the manifest"
            )
        if not region.source_relative_path or not region.source_absolute_path:
            raise ContextCompilationRecordError(
                f"Context region {ordinal} has invalid source paths"
            )
        if not Path(region.source_absolute_path).is_absolute():
            raise ContextCompilationRecordError(
                f"Context region {ordinal} source_absolute_path is not absolute"
            )
        if _SHA256_RE.fullmatch(region.source_sha256) is None:
            raise ContextCompilationRecordError(
                f"Context region {ordinal} source_sha256 is invalid"
            )
        expected_suffix = PurePosixPath(region.source_relative_path).suffix.casefold()
        if region.source_suffix.casefold() != expected_suffix:
            raise ContextCompilationRecordError(
                f"Context region {ordinal} source suffix does not match its path"
            )

        source_key = region.source_relative_path.casefold()
        source_identity = (
            region.source_relative_path,
            region.source_absolute_path,
            region.source_sha256,
            region.source_suffix.casefold(),
        )
        previous_source = sources.get(source_key)
        if previous_source is not None and previous_source != source_identity:
            raise ContextCompilationRecordError(
                f"Conflicting stored identity for source: {region.source_relative_path}"
            )
        sources[source_key] = source_identity

        if region.broad_start_paragraph > region.broad_end_paragraph:
            raise ContextCompilationRecordError(
                f"Context region {ordinal} has reversed broad boundaries"
            )
        paragraph_numbers = [paragraph.number for paragraph in region.paragraphs]
        if not paragraph_numbers or paragraph_numbers != list(
            range(region.broad_start_paragraph, region.broad_end_paragraph + 1)
        ):
            raise ContextCompilationRecordError(
                f"Context region {ordinal} paragraph map is not the complete broad span"
            )
        selection = region.selection
        if (
            not isinstance(selection.method, str)
            or not selection.method.strip()
            or selection.start_paragraph > selection.end_paragraph
            or selection.start_paragraph < region.broad_start_paragraph
            or selection.end_paragraph > region.broad_end_paragraph
        ):
            raise ContextCompilationRecordError(
                f"Context region {ordinal} selection is invalid"
            )
        if selection.confidence is not None and (
            isinstance(selection.confidence, bool)
            or not isinstance(selection.confidence, (int, float))
            or not 0 <= selection.confidence <= 1
        ):
            raise ContextCompilationRecordError(
                f"Context region {ordinal} selection confidence is invalid"
            )
        span_key = (
            region.source_relative_path.casefold(),
            selection.start_paragraph,
            selection.end_paragraph,
        )
        if span_key in selected_spans:
            raise ContextCompilationRecordError(
                f"Context region {ordinal} duplicates a selected source span"
            )
        selected_spans.add(span_key)

        paragraphs_by_number = {
            paragraph.number: paragraph for paragraph in region.paragraphs
        }
        expected_occurrences: list[tuple[int, int, int, str]] = []
        for paragraph in region.selected_paragraphs:
            expected_occurrences.extend(
                (
                    paragraph.number,
                    match.start(),
                    match.end(),
                    match.group(0),
                )
                for match in pattern.finditer(paragraph.text)
            )
        stored_occurrences: list[tuple[int, int, int, str]] = []
        for occurrence in region.occurrences:
            if not occurrence.occurrence_id or occurrence.occurrence_id in occurrence_ids:
                raise ContextCompilationRecordError(
                    f"Context region {ordinal} has a missing or duplicate occurrence_id"
                )
            occurrence_ids.add(occurrence.occurrence_id)
            paragraph = paragraphs_by_number.get(occurrence.paragraph_number)
            if paragraph is None:
                raise ContextCompilationRecordError(
                    f"Context region {ordinal} occurrence references a missing paragraph"
                )
            if not (
                selection.start_paragraph
                <= occurrence.paragraph_number
                <= selection.end_paragraph
            ):
                raise ContextCompilationRecordError(
                    f"Context region {ordinal} selection excludes an occurrence"
                )
            if (
                occurrence.start < 0
                or occurrence.end <= occurrence.start
                or paragraph.text[occurrence.start : occurrence.end]
                != occurrence.matched_text
                or pattern.fullmatch(occurrence.matched_text) is None
            ):
                raise ContextCompilationRecordError(
                    f"Context region {ordinal} occurrence offsets are invalid"
                )
            if occurrence.page_number != paragraph.page_number:
                raise ContextCompilationRecordError(
                    f"Context region {ordinal} occurrence page does not match its paragraph"
                )
            stored_occurrences.append(
                (
                    occurrence.paragraph_number,
                    occurrence.start,
                    occurrence.end,
                    occurrence.matched_text,
                )
            )
        if stored_occurrences != expected_occurrences:
            raise ContextCompilationRecordError(
                f"Context region {ordinal} occurrence records do not exactly cover its selection"
            )


def _validate_document(
    document: Any,
    records: _RecordsSnapshot,
) -> tuple[BoundContextRegion, ...]:
    properties = document.core_properties
    if (
        properties.subject != COMPILATION_MARKER
        or properties.keywords != COMPILATION_MARKER
    ):
        raise ContextCompilationStructureError(
            "DOCX does not contain both exact-context compilation markers"
        )
    expected_title = f'Context Finder: "{records.query.text}"'
    if properties.title != expected_title:
        raise ContextCompilationPairMismatchError(
            "DOCX core title query does not match the JSONL manifest"
        )
    if properties.author != COMPILATION_AUTHOR:
        raise ContextCompilationStructureError(
            "DOCX compilation author marker is invalid"
        )
    if document.tables:
        raise ContextCompilationStructureError(
            "Generated Context Finder compilation unexpectedly contains tables"
        )
    if len(document.sections) != 1:
        raise ContextCompilationStructureError(
            "Generated Context Finder compilation must contain exactly one Word section"
        )
    header_text = "\n".join(
        paragraph.text for paragraph in document.sections[0].header.paragraphs
    )
    if header_text != f'Context Finder | "{records.query.text}"':
        raise ContextCompilationPairMismatchError(
            "DOCX header query does not match the JSONL manifest"
        )

    paragraphs = list(document.paragraphs)
    heading_indexes = [
        index
        for index, paragraph in enumerate(paragraphs)
        if _style_name(paragraph) == HEADING_STYLE
    ]
    if len(heading_indexes) != len(records.regions):
        raise ContextCompilationPairMismatchError(
            "DOCX Heading 1 section count does not match JSONL region_count"
        )
    first_heading = heading_indexes[0] if heading_indexes else len(paragraphs)
    opening = paragraphs[:first_heading]
    _validate_opening(opening, records)

    query_pattern = compile_query_pattern(records.query)
    source_region_counts: dict[str, int] = {}
    bound: list[BoundContextRegion] = []
    used_relationship_ids: set[str] = set()

    for index, (heading_index, region) in enumerate(
        zip(heading_indexes, records.regions, strict=True),
        start=1,
    ):
        end = heading_indexes[index] if index < len(heading_indexes) else len(paragraphs)
        section = paragraphs[heading_index:end]
        expected_paragraphs = tuple(
            BoundSelectedParagraph(
                number=paragraph.number,
                text=paragraph.text,
                page_number=paragraph.page_number,
                emitted=bool(paragraph.text),
            )
            for paragraph in region.selected_paragraphs
        )
        emitted_texts = tuple(
            paragraph.text for paragraph in expected_paragraphs if paragraph.emitted
        )
        expected_section_length = 3 + len(emitted_texts)
        if len(section) != expected_section_length:
            raise ContextCompilationPairMismatchError(
                f"DOCX section {index} body paragraph count does not match JSONL selection"
            )
        heading, metadata, source_link = section[:3]
        body = tuple(section[3:])

        source_region_counts[region.source_relative_path] = (
            source_region_counts.get(region.source_relative_path, 0) + 1
        )
        source_region_ordinal = source_region_counts[region.source_relative_path]
        expected_heading = (
            f"{PurePosixPath(region.source_relative_path).stem} | "
            f"Context {source_region_ordinal}"
        )
        locator = _format_locator(region)
        expected_metadata = (
            f"Section {index} | {locator} | {len(region.occurrences)} occurrence"
            f"{'s' if len(region.occurrences) != 1 else ''}"
        )
        expected_target = Path(region.source_absolute_path).as_uri()

        if _style_name(heading) != HEADING_STYLE or heading.text != expected_heading:
            raise ContextCompilationPairMismatchError(
                f"DOCX section {index} heading does not match its JSONL region"
            )
        if _paragraph_hyperlinks(heading):
            raise ContextCompilationStructureError(
                f"DOCX section {index} heading unexpectedly contains a hyperlink"
            )
        if _style_name(metadata) != METADATA_STYLE or metadata.text != expected_metadata:
            raise ContextCompilationPairMismatchError(
                f"DOCX section {index} locator metadata does not match its JSONL region"
            )
        if _paragraph_hyperlinks(metadata):
            raise ContextCompilationStructureError(
                f"DOCX section {index} locator metadata unexpectedly contains a hyperlink"
            )
        if _style_name(source_link) != METADATA_STYLE:
            raise ContextCompilationPairMismatchError(
                f"DOCX section {index} source metadata style is invalid"
            )
        hyperlinks = _paragraph_hyperlinks(source_link)
        if len(hyperlinks) != 1:
            raise ContextCompilationPairMismatchError(
                f"DOCX section {index} must contain exactly one source hyperlink"
            )
        relationship_id, source_target, display_text = hyperlinks[0]
        if (
            source_link.text != f"Source: {region.source_relative_path}"
            or display_text != region.source_relative_path
            or source_target != expected_target
        ):
            raise ContextCompilationPairMismatchError(
                f"DOCX section {index} source hyperlink does not match its JSONL region"
            )
        used_relationship_ids.add(relationship_id)

        query_count = 0
        highlight_count = 0
        for body_index, (paragraph, expected_text) in enumerate(
            zip(body, emitted_texts, strict=True),
            start=1,
        ):
            if _style_name(paragraph) != BODY_STYLE or paragraph.text != expected_text:
                raise ContextCompilationPairMismatchError(
                    f"DOCX section {index} quotation paragraph {body_index} "
                    "does not exactly match the JSONL selection"
                )
            if _paragraph_hyperlinks(paragraph):
                raise ContextCompilationStructureError(
                    f"DOCX section {index} quotation body unexpectedly contains a hyperlink"
                )
            paragraph_query_count = len(tuple(query_pattern.finditer(expected_text)))
            paragraph_highlights = _validate_query_highlights(
                paragraph,
                query_pattern,
                section_index=index,
                body_index=body_index,
            )
            query_count += paragraph_query_count
            highlight_count += paragraph_highlights
        if query_count != len(region.occurrences):
            raise ContextCompilationPairMismatchError(
                f"DOCX section {index} query count does not match JSONL occurrences"
            )
        if highlight_count != query_count:
            raise ContextCompilationPairMismatchError(
                f"DOCX section {index} highlighted query count is invalid"
            )

        selected_text = "\n\n".join(emitted_texts)
        selected_paragraphs_payload = [
            {
                "number": paragraph.number,
                "text": paragraph.text,
                "page_number": paragraph.page_number,
                "emitted": paragraph.emitted,
            }
            for paragraph in expected_paragraphs
        ]
        selected_paragraphs_sha256 = _sha256_json(selected_paragraphs_payload)
        selected_text_sha256 = _sha256_text(selected_text)
        occurrence_ids = tuple(
            occurrence.occurrence_id for occurrence in region.occurrences
        )
        region_fingerprint = _sha256_json(
            {
                "schema_version": INVENTORY_SCHEMA_VERSION,
                "region_id": region.region_id,
                "ordinal": index,
                "source_region_ordinal": source_region_ordinal,
                "source": {
                    "relative_path": region.source_relative_path,
                    "absolute_path": region.source_absolute_path,
                    "sha256": region.source_sha256,
                    "target": source_target,
                },
                "heading": heading.text,
                "metadata": metadata.text,
                "locator": locator,
                "selection": {
                    "start_paragraph": region.selection.start_paragraph,
                    "end_paragraph": region.selection.end_paragraph,
                    "method": region.selection.method,
                    "model": region.selection.model,
                    "confidence": region.selection.confidence,
                    "note": region.selection.note,
                },
                "selected_paragraphs_sha256": selected_paragraphs_sha256,
                "selected_text_sha256": selected_text_sha256,
                "occurrences": [
                    {
                        "occurrence_id": occurrence.occurrence_id,
                        "paragraph_number": occurrence.paragraph_number,
                        "start": occurrence.start,
                        "end": occurrence.end,
                        "matched_text": occurrence.matched_text,
                        "page_number": occurrence.page_number,
                    }
                    for occurrence in region.occurrences
                ],
                "query_count": query_count,
                "highlight_count": highlight_count,
            }
        )
        bound.append(
            BoundContextRegion(
                region_id=region.region_id,
                ordinal=index,
                source_region_ordinal=source_region_ordinal,
                source_relative_path=region.source_relative_path,
                source_absolute_path=region.source_absolute_path,
                source_sha256=region.source_sha256,
                source_target=source_target,
                heading_text=heading.text,
                metadata_text=metadata.text,
                locator=locator,
                selected_paragraphs=expected_paragraphs,
                selected_text=selected_text,
                selected_text_sha256=selected_text_sha256,
                selected_paragraphs_sha256=selected_paragraphs_sha256,
                emitted_paragraph_count=len(emitted_texts),
                empty_selected_paragraph_count=sum(
                    not paragraph.emitted for paragraph in expected_paragraphs
                ),
                occurrence_count=len(region.occurrences),
                occurrence_ids=occurrence_ids,
                selection_start_paragraph=region.selection.start_paragraph,
                selection_end_paragraph=region.selection.end_paragraph,
                selection_method=region.selection.method,
                selection_model=region.selection.model,
                selection_confidence=region.selection.confidence,
                selection_note=region.selection.note,
                query_count=query_count,
                highlight_count=highlight_count,
                region_fingerprint=region_fingerprint,
            )
        )

    body_hyperlinks = document.element.body.xpath(".//w:hyperlink")
    if len(body_hyperlinks) != len(records.regions):
        raise ContextCompilationStructureError(
            "DOCX contains hyperlinks outside the validated source-link paragraphs"
        )
    external_hyperlink_relationships = {
        relationship_id
        for relationship_id, relationship in document.part.rels.items()
        if relationship.reltype == _HYPERLINK_RELATIONSHIP and relationship.is_external
    }
    if external_hyperlink_relationships != used_relationship_ids:
        raise ContextCompilationStructureError(
            "DOCX contains unused or unvalidated external hyperlink relationships"
        )
    if sum(region.highlight_count for region in bound) != _manifest_int(
        records.manifest, "occurrence_count", minimum=0
    ):
        raise ContextCompilationPairMismatchError(
            "DOCX total highlighted query count does not match the JSONL manifest"
        )
    return tuple(bound)


def _validate_opening(
    opening: Sequence[Paragraph],
    records: _RecordsSnapshot,
) -> None:
    if len(opening) != 3:
        raise ContextCompilationStructureError(
            "DOCX opening does not have the generated title, count, and source notice"
        )
    occurrence_count = _manifest_int(records.manifest, "occurrence_count", minimum=0)
    source_count = _manifest_int(records.manifest, "source_count", minimum=0)
    region_count = _manifest_int(records.manifest, "region_count", minimum=0)
    expected = (
        f'Context Finder: "{records.query.text}"',
        (
            f"{occurrence_count} exact occurrence"
            f"{'s' if occurrence_count != 1 else ''} in "
            f"{source_count} source document"
            f"{'s' if source_count != 1 else ''}, presented as "
            f"{region_count} distinct context section"
            f"{'s' if region_count != 1 else ''}."
        ),
        _SOURCE_NOTE,
    )
    if tuple(paragraph.text for paragraph in opening) != expected:
        raise ContextCompilationPairMismatchError(
            "DOCX opening query or expected counts do not match the JSONL manifest"
        )
    if any(_paragraph_hyperlinks(paragraph) for paragraph in opening):
        raise ContextCompilationStructureError(
            "DOCX opening unexpectedly contains a hyperlink"
        )


def _validate_query_highlights(
    paragraph: Paragraph,
    pattern: re.Pattern[str],
    *,
    section_index: int,
    body_index: int,
) -> int:
    runs = tuple(paragraph.runs)
    run_text = "".join(run.text for run in runs)
    if run_text != paragraph.text:
        raise ContextCompilationStructureError(
            f"DOCX section {section_index} quotation paragraph {body_index} "
            "contains unsupported nested text"
        )

    expected_mask = [False] * len(run_text)
    matches = tuple(pattern.finditer(run_text))
    for match in matches:
        expected_mask[match.start() : match.end()] = [True] * (
            match.end() - match.start()
        )

    actual_mask: list[bool] = []
    for run in runs:
        highlight = run.font.highlight_color
        if highlight is not None and highlight != WD_COLOR_INDEX.YELLOW:
            raise ContextCompilationPairMismatchError(
                f"DOCX section {section_index} quotation paragraph {body_index} "
                "contains non-yellow highlighting"
            )
        actual_mask.extend([highlight == WD_COLOR_INDEX.YELLOW] * len(run.text))
    if actual_mask != expected_mask:
        raise ContextCompilationPairMismatchError(
            f"DOCX section {section_index} quotation paragraph {body_index} "
            "does not highlight exactly the requested query"
        )
    return len(matches)


def _paragraph_hyperlinks(paragraph: Paragraph) -> tuple[tuple[str, str, str], ...]:
    details: list[tuple[str, str, str]] = []
    for hyperlink in paragraph._p.xpath(".//w:hyperlink"):
        relationship_id = hyperlink.get(qn("r:id"))
        if not relationship_id:
            raise ContextCompilationStructureError(
                "DOCX hyperlink has no relationship identifier"
            )
        relationship = paragraph.part.rels.get(str(relationship_id))
        if (
            relationship is None
            or relationship.reltype != _HYPERLINK_RELATIONSHIP
            or not relationship.is_external
        ):
            raise ContextCompilationStructureError(
                "DOCX source hyperlink is not an external hyperlink relationship"
            )
        display_text = "".join(hyperlink.xpath(".//w:t/text()"))
        details.append(
            (str(relationship_id), str(relationship.target_ref), display_text)
        )
    return tuple(details)


def _build_sources(
    regions: Sequence[BoundContextRegion],
) -> tuple[BoundContextSource, ...]:
    grouped: dict[str, list[BoundContextRegion]] = {}
    for region in regions:
        grouped.setdefault(region.source_relative_path, []).append(region)
    sources = [
        BoundContextSource(
            source_relative_path=relative_path,
            source_absolute_path=items[0].source_absolute_path,
            source_sha256=items[0].source_sha256,
            source_target=items[0].source_target,
            region_count=len(items),
            occurrence_count=sum(item.occurrence_count for item in items),
        )
        for relative_path, items in grouped.items()
    ]
    return tuple(
        sorted(sources, key=lambda source: source.source_relative_path.casefold())
    )


def _format_locator(region: ContextRegion) -> str:
    start = region.selection.start_paragraph
    end = region.selection.end_paragraph
    paragraph_text = f"Paragraph {start}" if start == end else f"Paragraphs {start}-{end}"
    pages = {
        paragraph.page_number
        for paragraph in region.selected_paragraphs
        if paragraph.page_number is not None
    }
    if len(pages) == 1:
        return f"Page {next(iter(pages))} | {paragraph_text}"
    if pages:
        return f"Pages {min(pages)}-{max(pages)} | {paragraph_text}"
    return paragraph_text


def _style_name(paragraph: Paragraph) -> str:
    return str(paragraph.style.name) if paragraph.style is not None else ""


def _manifest_int(
    manifest: Mapping[str, Any],
    key: str,
    *,
    minimum: int,
) -> int:
    value = manifest.get(key)
    if isinstance(value, bool) or not isinstance(value, int) or value < minimum:
        raise ContextCompilationRecordError(
            f"Search manifest {key} must be an integer of at least {minimum}"
        )
    return value


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _sha256_text(text: str) -> str:
    return _sha256_bytes(text.encode("utf-8"))


def _sha256_json(value: Any) -> str:
    encoded = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return _sha256_bytes(encoded)


__all__ = [
    "INVENTORY_SCHEMA_VERSION",
    "BoundContextCompilation",
    "BoundContextRegion",
    "BoundContextSource",
    "BoundSelectedParagraph",
    "ContextCompilationInventoryError",
    "ContextCompilationPairMismatchError",
    "ContextCompilationRecordError",
    "ContextCompilationStructureError",
    "bind_context_compilation",
]
