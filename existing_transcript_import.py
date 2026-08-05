"""Conservative, read-only import of legacy Faster-Whisper DOCX transcripts.

The older desktop transcriber wrote a labelled ``Transcript:`` section followed
by a generated ``Transcription Information`` footer.  This module extracts only
the labelled transcript body.  It intentionally rejects ambiguous documents
instead of guessing where lecture prose begins.
"""

from __future__ import annotations

from dataclasses import asdict, dataclass
import hashlib
from io import BytesIO
import os
from pathlib import Path
import re
import stat
from typing import Any
from zipfile import BadZipFile, ZipFile
from xml.etree import ElementTree

from docx import Document  # type: ignore


EXTRACTOR_VERSION = "legacy-faster-whisper-docx-v2"
MIN_TRANSCRIPT_WORDS = 20
MAX_DOCX_MEMBERS = 512
MAX_DOCX_MEMBER_UNCOMPRESSED_BYTES = 16 * 1024 * 1024
MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES = 64 * 1024 * 1024

_OLE_COMPOUND_FILE_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")
_TRANSCRIPT_MARKER_RE = re.compile(r"^transcript\s*:\s*$", re.IGNORECASE)
_INFORMATION_MARKER = "transcription information"
_DIVIDER_RE = re.compile(r"^[_=\-]{8,}$")
_WORKFLOW_FIELD_RE = re.compile(
    r"^(?:"
    r"model|device|processing\s+time|audio\s+preprocessing|"
    r"pipeline(?:\s+version)?|cleanup\s+model"
    r")\s*:\s*\S.*$",
    re.IGNORECASE,
)
_DELETION_NOTE_RE = re.compile(
    r"^\(?this\s+information\s+can\s+be\s+deleted\s+if\s+not\s+needed\)?[.!]?$",
    re.IGNORECASE,
)
_GENERATED_PROVENANCE_RE = re.compile(
    r"^processed\s+by\s+speech[- ]to[- ]text\s+from\s+a\s+digitised\s+tape\s+"
    r"recording\s+originally\s+(?:taken\s+from|recorded\s+in\s+person\s+by)\s+"
    r"[A-Z]{2,4}\s+on\s+.+[.]$",
    re.IGNORECASE,
)
_GENERATED_TIMESTAMP_RE = re.compile(
    r"^(?:cleaned\s+up|processed|generated)\s+at\s*:?\s*"
    r"\d{4}-\d{2}-\d{2}(?:[ T]\d{1,2}:\d{2}(?::\d{2})?)?.*$",
    re.IGNORECASE,
)
_WORD_RE = re.compile(r"\b[^\W_]+(?:['\u2019-][^\W_]+)*\b", re.UNICODE)
_HEADER_FOOTER_PART_RE = re.compile(r"^word/(?:header|footer)\d*\.xml$", re.IGNORECASE)
_SEMANTIC_TEXT_ELEMENTS = {"t", "delText", "instrText"}
_SEMANTIC_EMPTY_ELEMENTS = {
    "altChunk",
    "br",
    "cr",
    "drawing",
    "fldSimple",
    "object",
    "pict",
    "sym",
    "tab",
}
_TRACKED_CHANGE_ELEMENTS = {"del", "ins", "moveFrom", "moveTo"}
_TEXT_BOX_ELEMENTS = {"txbx", "txbxContent"}


class ExistingTranscriptImportError(ValueError):
    """Raised when a DOCX cannot be imported without unsafe assumptions."""


@dataclass(frozen=True)
class ExistingTranscriptImport:
    """Transcript text plus immutable provenance about the imported DOCX."""

    text: str
    source_docx: str
    source_sha256: str
    source_size: int
    source_mtime_ns: int
    word_count: int
    paragraph_count: int
    extractor_version: str = EXTRACTOR_VERSION

    @property
    def original_path(self) -> str:
        """Backward-readable alias for the original DOCX path."""

        return self.source_docx

    def to_dict(self) -> dict[str, Any]:
        """Return a JSON-serialisable audit record."""

        return asdict(self)


def _normalise_paragraph(text: str) -> str:
    """Normalise incidental Word whitespace but retain manual line breaks."""

    value = text.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    lines = [re.sub(r"[\t ]+", " ", line).strip() for line in value.split("\n")]
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return "\n".join(lines)


def _is_reparse_point(path_stat: Any) -> bool:
    reparse_flag = getattr(stat, "FILE_ATTRIBUTE_REPARSE_POINT", 0x400)
    return bool(getattr(path_stat, "st_file_attributes", 0) & reparse_flag)


def _same_file_snapshot(first: Any, second: Any) -> bool:
    fields = ("st_dev", "st_ino", "st_size", "st_mtime_ns")
    return all(getattr(first, field, None) == getattr(second, field, None) for field in fields)


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _parse_xml_part(package: ZipFile, name: str, source: Path) -> ElementTree.Element:
    try:
        return ElementTree.fromstring(package.read(name))
    except (ElementTree.ParseError, KeyError, OSError) as exc:
        raise ExistingTranscriptImportError(
            f"DOCX contains an invalid XML part ({name}): {source}"
        ) from exc


def _has_semantic_content(element: ElementTree.Element) -> bool:
    for descendant in element.iter():
        name = _local_name(descendant.tag)
        if name in _SEMANTIC_TEXT_ELEMENTS and (descendant.text or "").strip():
            return True
        if name in _SEMANTIC_EMPTY_ELEMENTS:
            return True
    return False


def _reject_omitted_semantic_content(
    package: ZipFile,
    names: set[str],
    source: Path,
) -> None:
    """Reject constructs not represented by ``Document.paragraphs``."""

    document_root = _parse_xml_part(package, "word/document.xml", source)

    for element in document_root.iter():
        name = _local_name(element.tag)
        if name in _TRACKED_CHANGE_ELEMENTS:
            raise ExistingTranscriptImportError(
                f"DOCX contains tracked insertions/deletions, which cannot be safely imported: {source}"
            )
        if name in _TEXT_BOX_ELEMENTS:
            raise ExistingTranscriptImportError(
                f"DOCX contains text-box content, which cannot be safely imported: {source}"
            )
        if name == "altChunk":
            raise ExistingTranscriptImportError(
                f"DOCX contains altChunk content, which cannot be safely imported: {source}"
            )
        if name == "tbl" and _has_semantic_content(element):
            raise ExistingTranscriptImportError(
                f"DOCX contains a nonempty table, which cannot be safely imported: {source}"
            )

    content_types_root = _parse_xml_part(package, "[Content_Types].xml", source)
    names_by_casefold = {name.casefold(): name for name in names}
    header_footer_parts = {
        name for name in names if _HEADER_FOOTER_PART_RE.fullmatch(name)
    }
    note_parts = {
        name
        for name in names
        if name.casefold() in {"word/footnotes.xml", "word/endnotes.xml"}
    }
    for override in content_types_root.iter():
        if _local_name(override.tag) != "Override":
            continue
        content_type = override.attrib.get("ContentType", "").casefold()
        declared_name = override.attrib.get("PartName", "").lstrip("/")
        if not declared_name:
            continue
        actual_name = names_by_casefold.get(declared_name.casefold(), declared_name)
        if content_type.endswith((".header+xml", ".footer+xml")):
            header_footer_parts.add(actual_name)
        if content_type.endswith((".footnotes+xml", ".endnotes+xml")):
            note_parts.add(actual_name)

    for name in sorted(header_footer_parts):
        root = _parse_xml_part(package, name, source)
        if _has_semantic_content(root):
            raise ExistingTranscriptImportError(
                f"DOCX contains a nonempty header/footer, which cannot be safely imported: {source}"
            )

    if note_parts:
        raise ExistingTranscriptImportError(
            "DOCX contains footnotes/endnotes, which cannot be safely imported "
            f"({', '.join(sorted(note_parts))}): {source}"
        )


def _validate_docx_package(payload: bytes, source: Path) -> None:
    if payload.startswith(_OLE_COMPOUND_FILE_MAGIC):
        raise ExistingTranscriptImportError(
            f"Encrypted or legacy binary Word documents are not supported: {source}"
        )

    try:
        with ZipFile(BytesIO(payload)) as package:
            members = package.infolist()
            if len(members) > MAX_DOCX_MEMBERS:
                raise ExistingTranscriptImportError(
                    f"DOCX has too many ZIP members ({len(members)}; maximum {MAX_DOCX_MEMBERS}): {source}"
                )

            member_names = [member.filename for member in members]
            names = set(member_names)
            if len(names) != len(member_names):
                raise ExistingTranscriptImportError(
                    f"DOCX contains duplicate ZIP member names: {source}"
                )
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ExistingTranscriptImportError(
                    f"File is not a valid DOCX package: {source}"
                )

            total_uncompressed = 0
            for member in members:
                if member.file_size > MAX_DOCX_MEMBER_UNCOMPRESSED_BYTES:
                    raise ExistingTranscriptImportError(
                        "DOCX ZIP member exceeds the uncompressed size limit "
                        f"({member.filename}: {member.file_size} bytes; maximum "
                        f"{MAX_DOCX_MEMBER_UNCOMPRESSED_BYTES}): {source}"
                    )
                total_uncompressed += member.file_size
                if total_uncompressed > MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES:
                    raise ExistingTranscriptImportError(
                        "DOCX exceeds the total uncompressed size limit "
                        f"({total_uncompressed} bytes; maximum "
                        f"{MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES}): {source}"
                    )

            if any(info.flag_bits & 0x1 for info in members):
                raise ExistingTranscriptImportError(
                    f"Encrypted DOCX packages are not supported: {source}"
                )
            bad_member = package.testzip()
            if bad_member is not None:
                raise ExistingTranscriptImportError(
                    f"DOCX package contains a corrupt member ({bad_member}): {source}"
                )
            _reject_omitted_semantic_content(package, names, source)
    except BadZipFile as exc:
        raise ExistingTranscriptImportError(f"File is not a valid DOCX package: {source}") from exc


def _is_footer_suffix(paragraphs: list[str], marker_index: int) -> bool:
    """Return true only for a strongly identifiable generated workflow footer."""

    suffix = [value for value in paragraphs[marker_index + 1 :] if value]
    workflow_fields = 0
    for value in suffix:
        if _WORKFLOW_FIELD_RE.fullmatch(value):
            workflow_fields += 1
            continue
        if _DELETION_NOTE_RE.fullmatch(value) or _GENERATED_TIMESTAMP_RE.fullmatch(value):
            continue
        return False

    divider_before = marker_index > 0 and bool(
        _DIVIDER_RE.fullmatch(paragraphs[marker_index - 1])
    )
    return workflow_fields >= 2 or (divider_before and not suffix)


def _find_workflow_footer(paragraphs: list[str], body_start: int) -> int | None:
    for index in range(len(paragraphs) - 1, body_start - 1, -1):
        if paragraphs[index].casefold() != _INFORMATION_MARKER:
            continue
        if _is_footer_suffix(paragraphs, index):
            return index
    return None


def _strip_narrow_generated_tail(paragraphs: list[str]) -> list[str]:
    """Strip only exact generated tail forms when no workflow footer is present."""

    result = list(paragraphs)
    while result and not result[-1]:
        result.pop()
    while result and (
        _GENERATED_PROVENANCE_RE.fullmatch(result[-1])
        or _GENERATED_TIMESTAMP_RE.fullmatch(result[-1])
    ):
        result.pop()
        while result and not result[-1]:
            result.pop()
    return result


def _extract_body(paragraphs: list[str]) -> list[str]:
    marker_indexes = [
        index for index, value in enumerate(paragraphs) if _TRANSCRIPT_MARKER_RE.fullmatch(value)
    ]
    if len(marker_indexes) != 1:
        detail = "missing" if not marker_indexes else "ambiguous"
        raise ExistingTranscriptImportError(
            f"Legacy DOCX has a {detail} Transcript: section marker"
        )

    body_start = marker_indexes[0] + 1
    footer_index = _find_workflow_footer(paragraphs, body_start)
    if footer_index is None:
        body = _strip_narrow_generated_tail(paragraphs[body_start:])
    else:
        body = paragraphs[body_start:footer_index]
        while body and not body[-1]:
            body.pop()
        if body and _DIVIDER_RE.fullmatch(body[-1]):
            body.pop()

    while body and not body[0]:
        body.pop(0)
    while body and not body[-1]:
        body.pop()
    return [value for value in body if value]


def import_existing_transcript(
    path: str | os.PathLike[str],
    *,
    minimum_words: int = MIN_TRANSCRIPT_WORDS,
) -> ExistingTranscriptImport:
    """Import a legacy Faster-Whisper DOCX without modifying the source file.

    A unique, explicit ``Transcript:`` marker is required.  Ambiguous or tiny
    documents are rejected so downstream cleanup never receives a title,
    workflow footer, or accidental non-transcript document.
    """

    if minimum_words < 1:
        raise ValueError("minimum_words must be at least 1")

    source = Path(path)
    if source.suffix.casefold() != ".docx":
        raise ExistingTranscriptImportError(f"Expected a .docx file: {source}")

    try:
        source_lstat = source.lstat()
    except FileNotFoundError as exc:
        raise ExistingTranscriptImportError(f"DOCX does not exist: {source}")
    except OSError as exc:
        raise ExistingTranscriptImportError(f"Unable to inspect DOCX path: {source}") from exc

    if stat.S_ISLNK(source_lstat.st_mode):
        raise ExistingTranscriptImportError(f"DOCX path must not be a symbolic link: {source}")
    if _is_reparse_point(source_lstat):
        raise ExistingTranscriptImportError(f"DOCX path must not be a reparse point: {source}")
    if not stat.S_ISREG(source_lstat.st_mode):
        raise ExistingTranscriptImportError(f"DOCX path is not a regular file: {source}")

    try:
        with source.open("rb") as stream:
            stat_before = os.fstat(stream.fileno())
            if not _same_file_snapshot(source_lstat, stat_before):
                raise ExistingTranscriptImportError(
                    f"DOCX changed before its snapshot could be read: {source}"
                )
            payload = stream.read()
            stat_after = os.fstat(stream.fileno())
    except OSError as exc:
        raise ExistingTranscriptImportError(f"Unable to read DOCX: {source}") from exc
    if not payload:
        raise ExistingTranscriptImportError(f"DOCX is empty: {source}")
    if (
        stat_before.st_size != stat_after.st_size
        or stat_before.st_mtime_ns != stat_after.st_mtime_ns
        or stat_after.st_size != len(payload)
    ):
        raise ExistingTranscriptImportError(f"DOCX changed while it was being read: {source}")

    try:
        source_lstat_after = source.lstat()
    except OSError as exc:
        raise ExistingTranscriptImportError(
            f"DOCX path changed while its snapshot was being read: {source}"
        ) from exc
    if (
        stat.S_ISLNK(source_lstat_after.st_mode)
        or _is_reparse_point(source_lstat_after)
        or not _same_file_snapshot(source_lstat_after, stat_after)
    ):
        raise ExistingTranscriptImportError(
            f"DOCX path changed while its snapshot was being read: {source}"
        )

    source_sha256 = hashlib.sha256(payload).hexdigest()
    _validate_docx_package(payload, source)
    try:
        document = Document(BytesIO(payload))
    except Exception as exc:
        raise ExistingTranscriptImportError(f"Unable to parse DOCX: {source}") from exc

    paragraphs = [_normalise_paragraph(paragraph.text) for paragraph in document.paragraphs]
    body_paragraphs = _extract_body(paragraphs)
    text = "\n\n".join(body_paragraphs).strip()
    word_count = len(_WORD_RE.findall(text))
    if not text:
        raise ExistingTranscriptImportError(f"Transcript body is empty: {source}")
    if word_count < minimum_words:
        raise ExistingTranscriptImportError(
            f"Transcript body is suspiciously tiny ({word_count} words; "
            f"minimum {minimum_words}): {source}"
        )

    return ExistingTranscriptImport(
        text=text,
        source_docx=os.path.abspath(os.fspath(source)),
        source_sha256=source_sha256,
        source_size=stat_after.st_size,
        source_mtime_ns=stat_after.st_mtime_ns,
        word_count=word_count,
        paragraph_count=len(body_paragraphs),
    )


__all__ = [
    "EXTRACTOR_VERSION",
    "MAX_DOCX_MEMBERS",
    "MAX_DOCX_MEMBER_UNCOMPRESSED_BYTES",
    "MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES",
    "MIN_TRANSCRIPT_WORDS",
    "ExistingTranscriptImport",
    "ExistingTranscriptImportError",
    "import_existing_transcript",
]
