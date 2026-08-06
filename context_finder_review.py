"""Create a structure-preserving GLM review copy of a Context Finder DOCX.

The exact-source Context Finder compilation remains immutable.  Only quotation
body paragraphs are submitted to the protected cleanup service and only those
paragraphs are replaced in the sibling review document.  Navigation, source
metadata, hyperlinks, headers and footers are fingerprinted before the cleanup
run and checked again on the staged output before it is atomically published.
"""

from __future__ import annotations

from concurrent.futures import Future, ThreadPoolExecutor, as_completed
from dataclasses import dataclass
import hashlib
import json
import os
from pathlib import Path
import re
import tempfile
from typing import Any, Callable, Mapping, Sequence

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from lxml import etree  # type: ignore

from cleanup_client import (  # type: ignore
    DEFAULT_CLEANUP_PROFILE,
    CleanupClient,
)
from context_finder import (  # type: ignore
    COMPILATION_MARKER,
    compile_query_pattern,
    validate_query,
)
from pipeline_control import (  # type: ignore
    CancelCheck,
    PipelineCancelledError,
    raise_if_cancelled,
)


REVIEW_SCHEMA_VERSION = "context-finder-glm-review-v1"
DEFAULT_MAX_WORKERS = 3
MAX_WORKERS = 8
GLM_REVIEW_SUFFIX = " - GLM Review"
METADATA_STYLE = "Context Source Metadata"
HEADING_STYLE = "Heading 1"
BODY_STYLE = "Normal"
SOURCE_NOTE_PREFIX = "Local-source edition. Quoted text is reproduced exactly"
REVIEW_NOTE_PREFIX = "GLM Review — Needs human review."
REVIEW_NOTE = (
    f"{REVIEW_NOTE_PREFIX} Quoted body text contains model-suggested corrections "
    "for checking against the linked source and the exact-source companion "
    "document. Highlighting marks the requested word or phrase."
)
_HYPERLINK_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)
_TITLE_RE = re.compile(r'^Context Finder: "(.+)"$')

ProgressCallback = Callable[["ReviewProgressUpdate"], None]


class ContextFinderReviewError(RuntimeError):
    """Base class for review-copy failures."""


class ReviewStructureError(ContextFinderReviewError):
    """The input or staged output does not match the protected structure."""


class ReviewSourceIntegrityError(ContextFinderReviewError):
    """The exact-source compilation changed during the review run."""


class ReviewOutputConflictError(ContextFinderReviewError):
    """The target exists but is not a proven prior tool publication."""


class ReviewCheckpointError(ContextFinderReviewError):
    """A resume manifest exists but is invalid or incompatible."""


@dataclass(frozen=True, slots=True)
class ReviewProgressUpdate:
    phase: str
    message: str
    completed: int = 0
    total: int = 0
    section_index: int | None = None
    heading: str | None = None
    status: str | None = None


@dataclass(frozen=True, slots=True)
class ContextFinderReviewOutcome:
    source_path: Path
    output_path: Path
    manifest_path: Path
    source_sha256: str
    source_fingerprint: str
    protected_fingerprint: str
    query: str
    region_count: int
    reviewed_regions: int
    unchanged_regions: int
    fallback_regions: int
    resumed_regions: int
    needs_review_regions: int
    occurrence_count: int
    model: str | None
    glossary_sha256: str
    glossary_count: int
    warnings: tuple[str, ...]

    @property
    def needs_human_review(self) -> bool:
        """Review copies are always provisional until a person approves them."""

        return True


@dataclass(frozen=True, slots=True)
class _CompilationSection:
    index: int
    section_id: str
    heading: Paragraph
    metadata: Paragraph
    source_link: Paragraph
    body: tuple[Paragraph, ...]
    source_target: str
    original_paragraphs: tuple[str, ...]
    original_occurrences: int

    @property
    def heading_text(self) -> str:
        return self.heading.text

    @property
    def input_text(self) -> str:
        return "\n\n".join(self.original_paragraphs)

    @property
    def input_sha256(self) -> str:
        return _sha256_text(self.input_text)


@dataclass(frozen=True, slots=True)
class _ParsedCompilation:
    query: str
    opening_note: Paragraph
    sections: tuple[_CompilationSection, ...]
    protected_fingerprint: str
    source_fingerprint: str


@dataclass(frozen=True, slots=True)
class _RegionReview:
    section_index: int
    section_id: str
    input_sha256: str
    paragraphs: tuple[str, ...]
    occurrence_count: int
    status: str
    changed: bool
    needs_review: bool
    resumed: bool
    model: str | None
    glossary_sha256: str | None
    glossary_count: int
    warnings: tuple[str, ...]

    def to_record(self) -> dict[str, Any]:
        return {
            "schema_version": REVIEW_SCHEMA_VERSION,
            "section_index": self.section_index,
            "section_id": self.section_id,
            "input_sha256": self.input_sha256,
            "paragraphs": list(self.paragraphs),
            "occurrence_count": self.occurrence_count,
            "status": self.status,
            "changed": self.changed,
            "needs_review": self.needs_review,
            "model": self.model,
            "glossary_sha256": self.glossary_sha256,
            "glossary_count": self.glossary_count,
            "warnings": list(self.warnings),
        }


__all__ = [
    "ContextFinderReviewError",
    "ContextFinderReviewOutcome",
    "ReviewCheckpointError",
    "ReviewOutputConflictError",
    "ReviewProgressUpdate",
    "ReviewSourceIntegrityError",
    "ReviewStructureError",
    "create_glm_review_copy",
    "default_review_output_path",
]


def default_review_output_path(source_docx: Path | str) -> Path:
    """Return the required sibling ``<stem> - GLM Review.docx`` path."""

    source = Path(source_docx).expanduser().resolve()
    return source.with_name(f"{source.stem}{GLM_REVIEW_SUFFIX}.docx")


def create_glm_review_copy(
    source_docx: Path | str,
    *,
    cleanup_client: CleanupClient | None = None,
    checkpoint_dir: Path | str | None = None,
    max_workers: int = DEFAULT_MAX_WORKERS,
    cancel_check: CancelCheck | None = None,
    progress_callback: ProgressCallback | None = None,
) -> ContextFinderReviewOutcome:
    """Create a resumable, structure-preserving sibling GLM review document.

    The source DOCX is only read.  A pre-existing output is replaceable only
    when the compatible atomic manifest records its exact byte hash as a prior
    publication (or as the candidate of an interrupted atomic publication).
    """

    source = Path(source_docx).expanduser().resolve()
    if source.suffix.casefold() != ".docx" or not source.is_file():
        raise FileNotFoundError(f"Context Finder source DOCX does not exist: {source}")
    if source.stem.casefold().endswith(GLM_REVIEW_SUFFIX.casefold()):
        raise ValueError("A GLM Review copy cannot be used as the exact-source input")
    output = default_review_output_path(source)
    if _normal_path(source) == _normal_path(output):
        raise ValueError("Review output must be separate from the source DOCX")

    workers = max(1, min(MAX_WORKERS, int(max_workers)))
    progress = progress_callback or (lambda _update: None)
    raise_if_cancelled(cancel_check, phase="context review preflight")
    progress(
        ReviewProgressUpdate(
            "preflight",
            f"Validating exact-source compilation: {source}",
        )
    )

    source_sha256 = _sha256_file(source)
    document = Document(str(source))
    parsed = _parse_compilation(document, require_review_note=False)
    source_fingerprint = parsed.source_fingerprint
    protected_fingerprint = parsed.protected_fingerprint
    if not parsed.sections:
        raise ReviewStructureError("Context Finder compilation contains no quote sections")

    checkpoint_base = _resolve_checkpoint_base(checkpoint_dir)
    run_key = _run_key(
        source=source,
        output=output,
        source_sha256=source_sha256,
        source_fingerprint=source_fingerprint,
    )
    run_dir = checkpoint_base / run_key
    manifest_path = run_dir / "review-manifest.json"
    existing_manifest = _load_manifest(manifest_path)
    # A completely unmanaged sibling is rejected before credentials are read or
    # any remote glossary/cleanup request can be made.
    if output.exists():
        _ensure_output_is_managed(output, existing_manifest or {}, manifest_path)

    active_client = cleanup_client or CleanupClient.from_environment()
    raise_if_cancelled(cancel_check, phase="context review glossary prefetch")
    glossary = active_client.ensure_glossary(cancel_check=cancel_check)
    if not bool(getattr(glossary, "pinned", False)):
        raise ContextFinderReviewError(
            "Context review requires one immutable, pinned glossary snapshot"
        )
    glossary_sha256 = getattr(glossary, "sha256", None)
    glossary_count = getattr(glossary, "count", None)
    if not isinstance(glossary_sha256, str) or not re.fullmatch(
        r"[0-9a-f]{64}", glossary_sha256
    ):
        raise ContextFinderReviewError("Pinned glossary returned no valid SHA-256")
    if not isinstance(glossary_count, int) or glossary_count < 0:
        raise ContextFinderReviewError("Pinned glossary returned an invalid term count")
    model = getattr(active_client, "model", None)
    if model is not None and not isinstance(model, str):
        raise ContextFinderReviewError("Cleanup client model id is invalid")

    expected_contract = _manifest_contract(
        source=source,
        output=output,
        source_sha256=source_sha256,
        parsed=parsed,
        model=model,
        glossary_sha256=glossary_sha256,
        glossary_count=glossary_count,
    )
    if existing_manifest is not None:
        _validate_manifest_contract(existing_manifest, expected_contract, manifest_path)
        manifest = existing_manifest
    else:
        manifest = expected_contract
    _ensure_output_is_managed(output, manifest, manifest_path)
    run_dir.mkdir(parents=True, exist_ok=True)
    manifest["status"] = "cleaning"
    manifest["last_error"] = None
    _write_json_atomically(manifest_path, manifest)

    progress(
        ReviewProgressUpdate(
            "cleaning",
            f"Reviewing {len(parsed.sections)} quote sections with up to {workers} concurrent regions.",
            total=len(parsed.sections),
        )
    )
    results: dict[int, _RegionReview] = {}
    completed = 0
    try:
        pending: list[_CompilationSection] = []
        for section in parsed.sections:
            resumed = _load_resumed_region(
                section,
                manifest,
                run_dir,
                glossary_sha256=glossary_sha256,
                glossary_count=glossary_count,
                model=model,
            )
            if resumed is None:
                pending.append(section)
                continue
            results[section.index] = resumed
            completed += 1
            progress(
                ReviewProgressUpdate(
                    "region",
                    f"Resumed {section.heading_text}",
                    completed=completed,
                    total=len(parsed.sections),
                    section_index=section.index,
                    heading=section.heading_text,
                    status="resumed",
                )
            )

        with ThreadPoolExecutor(
            max_workers=workers, thread_name_prefix="context-review-glm"
        ) as executor:
            futures: dict[Future[_RegionReview], _CompilationSection] = {
                executor.submit(
                    _review_region,
                    section,
                    query=parsed.query,
                    cleanup_client=active_client,
                    region_dir=_region_dir(run_dir, section),
                    glossary_sha256=glossary_sha256,
                    glossary_count=glossary_count,
                    cancel_check=cancel_check,
                ): section
                for section in pending
            }
            for future in as_completed(futures):
                section = futures[future]
                try:
                    result = future.result()
                except PipelineCancelledError:
                    for queued in futures:
                        queued.cancel()
                    raise
                results[section.index] = result
                completed += 1
                _record_region_result(manifest, result, run_dir)
                _write_json_atomically(manifest_path, manifest)
                progress(
                    ReviewProgressUpdate(
                        "region",
                        f"{result.status.replace('_', ' ').title()}: {section.heading_text}",
                        completed=completed,
                        total=len(parsed.sections),
                        section_index=section.index,
                        heading=section.heading_text,
                        status=result.status,
                    )
                )

        raise_if_cancelled(cancel_check, phase="context review publication")
        if len(results) != len(parsed.sections):
            raise ContextFinderReviewError("Not every quote section produced a review result")
        _assert_source_unchanged(
            source,
            expected_sha256=source_sha256,
            expected_source_fingerprint=source_fingerprint,
        )
        progress(
            ReviewProgressUpdate(
                "publish",
                f"Validating and publishing review copy: {output}",
                completed=len(results),
                total=len(parsed.sections),
            )
        )
        _apply_review_results(document, parsed, results)
        output_sha256 = _stage_validate_and_publish(
            document,
            parsed=parsed,
            results=results,
            source=source,
            output=output,
            source_sha256=source_sha256,
            source_fingerprint=source_fingerprint,
            manifest=manifest,
            manifest_path=manifest_path,
        )
        manifest["status"] = "complete"
        manifest["published_output_sha256"] = output_sha256
        manifest["candidate_output_sha256"] = None
        manifest["last_error"] = None
        _write_json_atomically(manifest_path, manifest)
    except BaseException as exc:
        manifest["status"] = (
            "cancelled" if isinstance(exc, PipelineCancelledError) else "failed"
        )
        manifest["last_error"] = _clean_error(f"{type(exc).__name__}: {exc}")
        _write_json_atomically(manifest_path, manifest)
        raise

    ordered = tuple(results[index] for index in sorted(results))
    warnings = tuple(
        warning
        for result in ordered
        for warning in result.warnings
    )
    outcome = ContextFinderReviewOutcome(
        source_path=source,
        output_path=output,
        manifest_path=manifest_path,
        source_sha256=source_sha256,
        source_fingerprint=source_fingerprint,
        protected_fingerprint=protected_fingerprint,
        query=parsed.query,
        region_count=len(ordered),
        reviewed_regions=sum(result.changed and result.status != "fallback" for result in ordered),
        unchanged_regions=sum(
            not result.changed and result.status != "fallback" for result in ordered
        ),
        fallback_regions=sum(result.status == "fallback" for result in ordered),
        resumed_regions=sum(result.resumed for result in ordered),
        needs_review_regions=sum(result.needs_review for result in ordered),
        occurrence_count=sum(result.occurrence_count for result in ordered),
        model=model,
        glossary_sha256=glossary_sha256,
        glossary_count=glossary_count,
        warnings=warnings,
    )
    progress(
        ReviewProgressUpdate(
            "complete",
            f"GLM Review copy ready: {output}",
            completed=len(ordered),
            total=len(ordered),
            status="complete",
        )
    )
    return outcome


def _parse_compilation(
    document: Any,
    *,
    require_review_note: bool,
) -> _ParsedCompilation:
    properties = document.core_properties
    if (
        properties.subject != COMPILATION_MARKER
        and properties.keywords != COMPILATION_MARKER
    ):
        raise ReviewStructureError("DOCX is not a generated exact-context compilation")
    title = properties.title or ""
    title_match = _TITLE_RE.fullmatch(title)
    if title_match is None:
        raise ReviewStructureError("Context Finder core title does not contain a valid query")
    try:
        query = validate_query(title_match.group(1)).text
    except ValueError as exc:
        raise ReviewStructureError("Context Finder query metadata is invalid") from exc
    query_pattern = compile_query_pattern(query)

    paragraphs = list(document.paragraphs)
    heading_indexes = [
        index
        for index, paragraph in enumerate(paragraphs)
        if _style_name(paragraph) == HEADING_STYLE
    ]
    if not heading_indexes:
        raise ReviewStructureError("Context Finder compilation contains no Heading 1 sections")
    first_heading = heading_indexes[0]
    opening = paragraphs[:first_heading]
    source_notes = [
        paragraph
        for paragraph in opening
        if paragraph.text.startswith(SOURCE_NOTE_PREFIX)
    ]
    review_notes = [
        paragraph
        for paragraph in opening
        if paragraph.text.startswith(REVIEW_NOTE_PREFIX)
    ]
    if require_review_note:
        if len(review_notes) != 1 or source_notes:
            raise ReviewStructureError("Staged review document has no unique review notice")
        opening_note = review_notes[0]
    else:
        if len(source_notes) != 1 or review_notes:
            raise ReviewStructureError(
                "Exact-source compilation has no unique exact-quotation notice"
            )
        opening_note = source_notes[0]

    sections: list[_CompilationSection] = []
    for ordinal, heading_index in enumerate(heading_indexes, start=1):
        end = (
            heading_indexes[ordinal]
            if ordinal < len(heading_indexes)
            else len(paragraphs)
        )
        section_paragraphs = paragraphs[heading_index:end]
        if len(section_paragraphs) < 4:
            raise ReviewStructureError(
                f"Section {ordinal} does not contain heading, metadata, source link and body"
            )
        heading, metadata, source_link = section_paragraphs[:3]
        body = tuple(section_paragraphs[3:])
        if _style_name(heading) != HEADING_STYLE or not heading.text.strip():
            raise ReviewStructureError(f"Section {ordinal} has an invalid heading")
        if _style_name(metadata) != METADATA_STYLE:
            raise ReviewStructureError(
                f"Section {ordinal} first metadata paragraph has an unexpected style"
            )
        if _style_name(source_link) != METADATA_STYLE:
            raise ReviewStructureError(
                f"Section {ordinal} source link paragraph has an unexpected style"
            )
        if _external_hyperlinks(metadata):
            raise ReviewStructureError(
                f"Section {ordinal} locator metadata unexpectedly contains a hyperlink"
            )
        targets = _external_hyperlinks(source_link)
        if len(targets) != 1:
            raise ReviewStructureError(
                f"Section {ordinal} must contain exactly one external source hyperlink"
            )
        if not source_link.text.startswith("Source: "):
            raise ReviewStructureError(
                f"Section {ordinal} source hyperlink paragraph is malformed"
            )
        if not body or any(
            _style_name(paragraph) != BODY_STYLE or not paragraph.text.strip()
            for paragraph in body
        ):
            raise ReviewStructureError(
                f"Section {ordinal} quotation body is empty or has an unexpected style"
            )
        original_paragraphs = tuple(paragraph.text for paragraph in body)
        occurrence_count = _count_occurrences(original_paragraphs, query_pattern)
        if occurrence_count < 1:
            raise ReviewStructureError(
                f"Section {ordinal} quotation body no longer contains the exact query"
            )
        section_identity = {
            "index": ordinal,
            "heading": heading.text,
            "metadata": metadata.text,
            "source_link": source_link.text,
            "source_target": targets[0],
        }
        section_id = "section_" + _sha256_json(section_identity)[:20]
        sections.append(
            _CompilationSection(
                index=ordinal,
                section_id=section_id,
                heading=heading,
                metadata=metadata,
                source_link=source_link,
                body=body,
                source_target=targets[0],
                original_paragraphs=original_paragraphs,
                original_occurrences=occurrence_count,
            )
        )

    protected = _protected_snapshot(
        document,
        query=query,
        opening=opening,
        opening_note=opening_note,
        sections=sections,
    )
    protected_fingerprint = _sha256_json(protected)
    source_fingerprint = _sha256_json(
        {
            "protected_fingerprint": protected_fingerprint,
            "sections": [
                {
                    "section_id": section.section_id,
                    "input_sha256": section.input_sha256,
                    "occurrence_count": section.original_occurrences,
                    "paragraph_count": len(section.original_paragraphs),
                }
                for section in sections
            ],
        }
    )
    return _ParsedCompilation(
        query=query,
        opening_note=opening_note,
        sections=tuple(sections),
        protected_fingerprint=protected_fingerprint,
        source_fingerprint=source_fingerprint,
    )


def _protected_snapshot(
    document: Any,
    *,
    query: str,
    opening: Sequence[Paragraph],
    opening_note: Paragraph,
    sections: Sequence[_CompilationSection],
) -> dict[str, Any]:
    protected_opening = [
        _canonical_xml(paragraph._p)
        for paragraph in opening
        if paragraph._p is not opening_note._p
    ]
    external_relationships = sorted(
        (relationship.reltype, str(relationship.target_ref))
        for relationship in document.part.rels.values()
        if relationship.is_external
    )
    section_parts = []
    for section in document.sections:
        section_parts.append(
            {
                "sect_pr": _canonical_xml(section._sectPr),
                "header": _canonical_xml(section.header._element),
                "footer": _canonical_xml(section.footer._element),
            }
        )
    properties = document.core_properties
    return {
        "query": query,
        "core": {
            "title": properties.title,
            "author": properties.author,
            "subject": properties.subject,
            "keywords": properties.keywords,
        },
        "opening": protected_opening,
        "sections": [
            {
                "section_id": section.section_id,
                "heading": _canonical_xml(section.heading._p),
                "metadata": _canonical_xml(section.metadata._p),
                "source_link": _canonical_xml(section.source_link._p),
                "source_target": section.source_target,
            }
            for section in sections
        ],
        "external_relationships": external_relationships,
        "section_parts": section_parts,
    }


def _external_hyperlinks(paragraph: Paragraph) -> tuple[str, ...]:
    relationship_ids = paragraph._p.xpath(".//w:hyperlink/@r:id")
    targets: list[str] = []
    for relationship_id in relationship_ids:
        relationship = paragraph.part.rels.get(str(relationship_id))
        if (
            relationship is None
            or relationship.reltype != _HYPERLINK_RELATIONSHIP
            or not relationship.is_external
        ):
            raise ReviewStructureError("Source hyperlink is not an external relationship")
        targets.append(str(relationship.target_ref))
    return tuple(targets)


def _style_name(paragraph: Paragraph) -> str:
    style = paragraph.style
    return str(style.name) if style is not None else ""


def _canonical_xml(element: Any) -> str:
    return etree.tostring(
        element,
        method="c14n",
        with_comments=True,
    ).decode("utf-8")


def _manifest_contract(
    *,
    source: Path,
    output: Path,
    source_sha256: str,
    parsed: _ParsedCompilation,
    model: str | None,
    glossary_sha256: str,
    glossary_count: int,
) -> dict[str, Any]:
    return {
        "schema_version": REVIEW_SCHEMA_VERSION,
        "source_path": str(source),
        "output_path": str(output),
        "source_sha256": source_sha256,
        "source_fingerprint": parsed.source_fingerprint,
        "protected_fingerprint": parsed.protected_fingerprint,
        "query": parsed.query,
        "cleanup_profile": DEFAULT_CLEANUP_PROFILE,
        "model": model,
        "glossary_sha256": glossary_sha256,
        "glossary_count": glossary_count,
        "status": "pending",
        "published_output_sha256": None,
        "candidate_output_sha256": None,
        "last_error": None,
        "sections": [
            {
                "section_index": section.index,
                "section_id": section.section_id,
                "input_sha256": section.input_sha256,
                "occurrence_count": section.original_occurrences,
                "paragraph_count": len(section.original_paragraphs),
                "status": "pending",
                "result_sha256": None,
                "result_path": None,
            }
            for section in parsed.sections
        ],
    }


def _validate_manifest_contract(
    manifest: Mapping[str, Any],
    expected: Mapping[str, Any],
    manifest_path: Path,
) -> None:
    fields = (
        "schema_version",
        "source_path",
        "output_path",
        "source_sha256",
        "source_fingerprint",
        "protected_fingerprint",
        "query",
        "cleanup_profile",
        "model",
        "glossary_sha256",
        "glossary_count",
    )
    for field in fields:
        if manifest.get(field) != expected.get(field):
            raise ReviewCheckpointError(
                f"Resume manifest contract mismatch for {field}: {manifest_path}"
            )
    actual_sections = manifest.get("sections")
    expected_sections = expected.get("sections")
    if not isinstance(actual_sections, list) or not isinstance(expected_sections, list):
        raise ReviewCheckpointError(f"Resume manifest has invalid sections: {manifest_path}")
    if len(actual_sections) != len(expected_sections):
        raise ReviewCheckpointError(
            f"Resume manifest section count changed: {manifest_path}"
        )
    identity_fields = (
        "section_index",
        "section_id",
        "input_sha256",
        "occurrence_count",
        "paragraph_count",
    )
    for actual, wanted in zip(actual_sections, expected_sections, strict=True):
        if not isinstance(actual, Mapping) or not isinstance(wanted, Mapping):
            raise ReviewCheckpointError(
                f"Resume manifest contains an invalid section record: {manifest_path}"
            )
        if any(actual.get(field) != wanted.get(field) for field in identity_fields):
            raise ReviewCheckpointError(
                f"Resume manifest section identity changed: {manifest_path}"
            )


def _load_manifest(path: Path) -> dict[str, Any] | None:
    if not path.exists():
        return None
    if not path.is_file() or path.is_symlink():
        raise ReviewCheckpointError(f"Resume manifest is not a regular file: {path}")
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ReviewCheckpointError(f"Could not read resume manifest: {path}") from exc
    if not isinstance(payload, dict):
        raise ReviewCheckpointError(f"Resume manifest is not a JSON object: {path}")
    return payload


def _ensure_output_is_managed(
    output: Path,
    manifest: Mapping[str, Any],
    manifest_path: Path,
) -> None:
    if not output.exists():
        return
    if not output.is_file() or output.is_symlink():
        raise ReviewOutputConflictError(
            f"GLM Review target is not a regular managed file: {output}"
        )
    current_hash = _sha256_file(output)
    allowed = {
        value.casefold()
        for value in (
            manifest.get("published_output_sha256"),
            manifest.get("candidate_output_sha256"),
        )
        if isinstance(value, str) and re.fullmatch(r"[0-9a-f]{64}", value)
    }
    if current_hash.casefold() not in allowed:
        raise ReviewOutputConflictError(
            "GLM Review target exists but is not a proven prior publication or "
            f"was manually changed: {output} (manifest: {manifest_path})"
        )


def _resolve_checkpoint_base(checkpoint_dir: Path | str | None) -> Path:
    if checkpoint_dir is not None:
        return Path(checkpoint_dir).expanduser().resolve()
    base = Path(os.environ.get("LOCALAPPDATA") or tempfile.gettempdir())
    return base / "AudioProcessor" / "context-finder-review"


def _run_key(
    *,
    source: Path,
    output: Path,
    source_sha256: str,
    source_fingerprint: str,
) -> str:
    identity = {
        "schema_version": REVIEW_SCHEMA_VERSION,
        "source": _normal_path(source),
        "output": _normal_path(output),
        "source_sha256": source_sha256,
        "source_fingerprint": source_fingerprint,
    }
    return "run_" + _sha256_json(identity)[:24]


def _region_dir(run_dir: Path, section: _CompilationSection) -> Path:
    return run_dir / "regions" / f"{section.index:05d}-{section.section_id}"


def _region_result_path(run_dir: Path, section: _CompilationSection) -> Path:
    return _region_dir(run_dir, section) / "review-result.json"


def _manifest_section(
    manifest: Mapping[str, Any], section_id: str
) -> dict[str, Any] | None:
    records = manifest.get("sections")
    if not isinstance(records, list):
        return None
    for record in records:
        if isinstance(record, dict) and record.get("section_id") == section_id:
            return record
    return None


def _load_resumed_region(
    section: _CompilationSection,
    manifest: Mapping[str, Any],
    run_dir: Path,
    *,
    glossary_sha256: str,
    glossary_count: int,
    model: str | None,
) -> _RegionReview | None:
    manifest_record = _manifest_section(manifest, section.section_id)
    if manifest_record is None or manifest_record.get("status") not in {
        "reviewed",
        "unchanged",
    }:
        return None
    result_path = _region_result_path(run_dir, section)
    recorded_relative = manifest_record.get("result_path")
    if recorded_relative != str(result_path.relative_to(run_dir)):
        return None
    expected_hash = manifest_record.get("result_sha256")
    if (
        not isinstance(expected_hash, str)
        or not re.fullmatch(r"[0-9a-f]{64}", expected_hash)
        or not result_path.is_file()
        or result_path.is_symlink()
        or _sha256_file(result_path) != expected_hash
    ):
        return None
    try:
        payload = json.loads(result_path.read_text(encoding="utf-8"))
    except (OSError, UnicodeDecodeError, json.JSONDecodeError):
        return None
    if not isinstance(payload, Mapping):
        return None
    result = _region_from_record(payload, resumed=True)
    if (
        result.section_index != section.index
        or result.section_id != section.section_id
        or result.input_sha256 != section.input_sha256
        or result.occurrence_count != section.original_occurrences
        or result.glossary_sha256 != glossary_sha256
        or result.glossary_count != glossary_count
        or result.model != model
        or result.status not in {"reviewed", "unchanged"}
    ):
        return None
    query = manifest.get("query")
    if not isinstance(query, str):
        return None
    pattern = compile_query_pattern(query)
    if _count_occurrences(result.paragraphs, pattern) != section.original_occurrences:
        return None
    return result


def _record_region_result(
    manifest: dict[str, Any], result: _RegionReview, run_dir: Path
) -> None:
    record = _manifest_section(manifest, result.section_id)
    if record is None:
        raise ReviewCheckpointError(
            f"Manifest contains no section record for {result.section_id}"
        )
    result_path = (
        run_dir
        / "regions"
        / f"{result.section_index:05d}-{result.section_id}"
        / "review-result.json"
    )
    record.update(
        {
            "status": result.status,
            "result_path": str(result_path.relative_to(run_dir)),
            "result_sha256": _sha256_file(result_path),
            "needs_review": result.needs_review,
            "changed": result.changed,
        }
    )


def _review_region(
    section: _CompilationSection,
    *,
    query: str,
    cleanup_client: CleanupClient,
    region_dir: Path,
    glossary_sha256: str,
    glossary_count: int,
    cancel_check: CancelCheck | None,
) -> _RegionReview:
    raise_if_cancelled(cancel_check, phase=f"context review section {section.index}")
    model = getattr(cleanup_client, "model", None)
    try:
        cleanup_result = cleanup_client.cleanup_text(
            section.input_text,
            checkpoint_dir=region_dir / "cleanup",
            reuse_checkpoints=True,
            cancel_check=cancel_check,
        )
        raise_if_cancelled(
            cancel_check, phase=f"context review section {section.index}"
        )
        cleaned_text = getattr(cleanup_result, "text", None)
        if not isinstance(cleaned_text, str) or not cleaned_text.strip():
            raise ContextFinderReviewError("cleanup returned no reviewed quotation text")
        result_model = getattr(cleanup_result, "model", None)
        if result_model != model:
            raise ContextFinderReviewError(
                f"cleanup returned model {result_model!r}, expected {model!r}"
            )
        result_glossary_sha = getattr(cleanup_result, "glossary_sha256", None)
        result_glossary_count = getattr(cleanup_result, "glossary_count", None)
        if (
            result_glossary_sha != glossary_sha256
            or result_glossary_count != glossary_count
        ):
            raise ContextFinderReviewError(
                "cleanup result did not use the prefetched pinned glossary"
            )
        paragraphs = _split_cleaned_paragraphs(cleaned_text)
        pattern = compile_query_pattern(query)
        occurrence_count = _count_occurrences(paragraphs, pattern)
        if occurrence_count != section.original_occurrences:
            raise ContextFinderReviewError(
                "reviewed quotation changed the exact-query occurrence count "
                f"from {section.original_occurrences} to {occurrence_count}"
            )
        warnings_value = getattr(cleanup_result, "warnings", ())
        warnings = tuple(_clean_error(str(value)) for value in warnings_value)
        changed = paragraphs != section.original_paragraphs
        result = _RegionReview(
            section_index=section.index,
            section_id=section.section_id,
            input_sha256=section.input_sha256,
            paragraphs=paragraphs,
            occurrence_count=occurrence_count,
            status="reviewed" if changed else "unchanged",
            changed=changed,
            needs_review=bool(getattr(cleanup_result, "needs_review", True)),
            resumed=False,
            model=result_model,
            glossary_sha256=result_glossary_sha,
            glossary_count=int(result_glossary_count),
            warnings=warnings,
        )
    except PipelineCancelledError:
        raise
    except Exception as exc:
        warning = (
            f"Section {section.index} ({section.heading_text}) retained exact-source "
            f"text because GLM review failed validation: {type(exc).__name__}: "
            f"{_clean_error(str(exc))}"
        )
        result = _RegionReview(
            section_index=section.index,
            section_id=section.section_id,
            input_sha256=section.input_sha256,
            paragraphs=section.original_paragraphs,
            occurrence_count=section.original_occurrences,
            status="fallback",
            changed=False,
            needs_review=True,
            resumed=False,
            model=model,
            glossary_sha256=glossary_sha256,
            glossary_count=glossary_count,
            warnings=(warning,),
        )
    region_dir.mkdir(parents=True, exist_ok=True)
    _write_json_atomically(region_dir / "review-result.json", result.to_record())
    return result


def _region_from_record(
    payload: Mapping[str, Any], *, resumed: bool
) -> _RegionReview:
    if payload.get("schema_version") != REVIEW_SCHEMA_VERSION:
        raise ReviewCheckpointError("Region result uses an unsupported schema")
    paragraphs_value = payload.get("paragraphs")
    warnings_value = payload.get("warnings")
    if (
        not isinstance(paragraphs_value, list)
        or not paragraphs_value
        or any(not isinstance(value, str) or not value.strip() for value in paragraphs_value)
        or not isinstance(warnings_value, list)
        or any(not isinstance(value, str) for value in warnings_value)
    ):
        raise ReviewCheckpointError("Region result contains invalid text fields")
    section_index = payload.get("section_index")
    glossary_count = payload.get("glossary_count")
    occurrence_count = payload.get("occurrence_count")
    if (
        not isinstance(section_index, int)
        or section_index < 1
        or not isinstance(glossary_count, int)
        or glossary_count < 0
        or not isinstance(occurrence_count, int)
        or occurrence_count < 1
    ):
        raise ReviewCheckpointError("Region result contains invalid numeric fields")
    status = payload.get("status")
    if status not in {"reviewed", "unchanged", "fallback"}:
        raise ReviewCheckpointError("Region result contains an invalid status")
    section_id = payload.get("section_id")
    input_sha256 = payload.get("input_sha256")
    glossary_sha256 = payload.get("glossary_sha256")
    model = payload.get("model")
    if (
        not isinstance(section_id, str)
        or not isinstance(input_sha256, str)
        or not re.fullmatch(r"[0-9a-f]{64}", input_sha256)
        or not isinstance(glossary_sha256, str)
        or not re.fullmatch(r"[0-9a-f]{64}", glossary_sha256)
        or (model is not None and not isinstance(model, str))
    ):
        raise ReviewCheckpointError("Region result contains invalid identity fields")
    return _RegionReview(
        section_index=section_index,
        section_id=section_id,
        input_sha256=input_sha256,
        paragraphs=tuple(paragraphs_value),
        occurrence_count=occurrence_count,
        status=status,
        changed=bool(payload.get("changed")),
        needs_review=bool(payload.get("needs_review")),
        resumed=resumed,
        model=model,
        glossary_sha256=glossary_sha256,
        glossary_count=glossary_count,
        warnings=tuple(warnings_value),
    )


def _split_cleaned_paragraphs(text: str) -> tuple[str, ...]:
    normalised = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalised:
        raise ContextFinderReviewError("cleanup returned empty quotation text")
    if re.search(r"\n[ \t]*\n", normalised):
        raw_blocks = re.split(r"\n[ \t]*\n+", normalised)
    else:
        lines = [line.strip() for line in normalised.split("\n") if line.strip()]
        raw_blocks = lines if len(lines) > 1 else [normalised]
    paragraphs = tuple(
        re.sub(r"[ \t]*\n[ \t]*", " ", block).strip()
        for block in raw_blocks
        if block.strip()
    )
    if not paragraphs:
        raise ContextFinderReviewError("cleanup returned no quotation paragraphs")
    if any("\x00" in paragraph for paragraph in paragraphs):
        raise ContextFinderReviewError("cleanup returned a NUL character")
    return paragraphs


def _count_occurrences(
    paragraphs: Sequence[str], pattern: re.Pattern[str]
) -> int:
    return sum(len(tuple(pattern.finditer(paragraph))) for paragraph in paragraphs)


def _apply_review_results(
    document: Any,
    parsed: _ParsedCompilation,
    results: Mapping[int, _RegionReview],
) -> None:
    note = parsed.opening_note
    note.clear()
    lead = note.add_run(REVIEW_NOTE_PREFIX)
    lead.bold = True
    note.add_run(REVIEW_NOTE[len(REVIEW_NOTE_PREFIX) :])
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pattern = compile_query_pattern(parsed.query)
    for section in parsed.sections:
        result = results[section.index]
        parent = section.source_link._parent
        for paragraph in section.body:
            element_parent = paragraph._p.getparent()
            if element_parent is None:
                raise ReviewStructureError(
                    f"Section {section.index} body paragraph is detached"
                )
            element_parent.remove(paragraph._p)
        anchor = section.source_link._p
        for text in result.paragraphs:
            element = OxmlElement("w:p")
            anchor.addnext(element)
            paragraph = Paragraph(element, parent)
            paragraph.style = BODY_STYLE
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_highlighted_text(paragraph, text, pattern)
            anchor = element


def _add_highlighted_text(
    paragraph: Paragraph, text: str, pattern: re.Pattern[str]
) -> None:
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        run = paragraph.add_run(match.group(0))
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _stage_validate_and_publish(
    document: Any,
    *,
    parsed: _ParsedCompilation,
    results: Mapping[int, _RegionReview],
    source: Path,
    output: Path,
    source_sha256: str,
    source_fingerprint: str,
    manifest: dict[str, Any],
    manifest_path: Path,
) -> str:
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="wb",
            prefix=f".{output.stem}.",
            suffix=".tmp.docx",
            dir=output.parent,
            delete=False,
        ) as handle:
            temporary = Path(handle.name)
        document.save(str(temporary))
        staged_document = Document(str(temporary))
        staged = _parse_compilation(staged_document, require_review_note=True)
        _validate_staged_review(staged_document, parsed, staged, results)
        candidate_sha256 = _sha256_file(temporary)

        _assert_source_unchanged(
            source,
            expected_sha256=source_sha256,
            expected_source_fingerprint=source_fingerprint,
        )
        _ensure_output_is_managed(output, manifest, manifest_path)
        manifest["status"] = "publishing"
        manifest["candidate_output_sha256"] = candidate_sha256
        manifest["last_error"] = None
        _write_json_atomically(manifest_path, manifest)
        os.replace(temporary, output)
        temporary = None
        published_sha256 = _sha256_file(output)
        if published_sha256 != candidate_sha256:
            raise ContextFinderReviewError(
                "Published GLM Review bytes do not match the validated staged document"
            )
        return published_sha256
    finally:
        if temporary is not None:
            temporary.unlink(missing_ok=True)


def _validate_staged_review(
    document: Any,
    source_parsed: _ParsedCompilation,
    staged: _ParsedCompilation,
    results: Mapping[int, _RegionReview],
) -> None:
    if staged.query != source_parsed.query:
        raise ReviewStructureError("Staged review changed the query metadata")
    if staged.protected_fingerprint != source_parsed.protected_fingerprint:
        raise ReviewStructureError(
            "Staged review changed headings, source metadata, hyperlinks, header or footer"
        )
    if len(staged.sections) != len(source_parsed.sections):
        raise ReviewStructureError("Staged review changed the section count")
    pattern = compile_query_pattern(source_parsed.query)
    for source_section, staged_section in zip(
        source_parsed.sections, staged.sections, strict=True
    ):
        result = results[source_section.index]
        if staged_section.section_id != source_section.section_id:
            raise ReviewStructureError(
                f"Staged review changed section identity {source_section.index}"
            )
        staged_texts = tuple(
            paragraph.text for paragraph in staged_section.body
        )
        if staged_texts != result.paragraphs:
            raise ReviewStructureError(
                f"Staged review body differs from validated result {source_section.index}"
            )
        if any(
            _style_name(paragraph) != BODY_STYLE
            or paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY
            for paragraph in staged_section.body
        ):
            raise ReviewStructureError(
                f"Staged review body formatting changed in section {source_section.index}"
            )
        occurrence_count = _count_occurrences(staged_texts, pattern)
        if occurrence_count != source_section.original_occurrences:
            raise ReviewStructureError(
                f"Staged review changed query count in section {source_section.index}"
            )
        highlighted = _highlighted_occurrence_count(staged_section.body, pattern)
        if highlighted != occurrence_count:
            raise ReviewStructureError(
                f"Staged review lost query highlighting in section {source_section.index}"
            )
    if not staged.opening_note.text.startswith(REVIEW_NOTE_PREFIX):
        raise ReviewStructureError("Staged review notice is missing")
    if document.core_properties.subject != COMPILATION_MARKER:
        raise ReviewStructureError("Staged review lost the generated-document marker")


def _highlighted_occurrence_count(
    paragraphs: Sequence[Paragraph], pattern: re.Pattern[str]
) -> int:
    count = 0
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                count += len(tuple(pattern.finditer(run.text)))
    return count


def _assert_source_unchanged(
    source: Path,
    *,
    expected_sha256: str,
    expected_source_fingerprint: str,
) -> None:
    if not source.is_file() or source.is_symlink():
        raise ReviewSourceIntegrityError(
            f"Exact-source compilation is missing or no longer a regular file: {source}"
        )
    actual_sha256 = _sha256_file(source)
    if actual_sha256 != expected_sha256:
        raise ReviewSourceIntegrityError(
            f"Exact-source compilation changed during GLM review: {source}"
        )
    try:
        reparsed = _parse_compilation(
            Document(str(source)), require_review_note=False
        )
    except Exception as exc:
        raise ReviewSourceIntegrityError(
            f"Exact-source compilation could not be revalidated: {source}"
        ) from exc
    if reparsed.source_fingerprint != expected_source_fingerprint:
        raise ReviewSourceIntegrityError(
            f"Exact-source structure changed during GLM review: {source}"
        )


def _write_json_atomically(path: Path, payload: Mapping[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="w",
            encoding="utf-8",
            newline="\n",
            prefix=f".{path.name}.",
            suffix=".tmp",
            dir=path.parent,
            delete=False,
        ) as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2, sort_keys=True)
            handle.write("\n")
            handle.flush()
            os.fsync(handle.fileno())
            temporary = Path(handle.name)
        os.replace(temporary, path)
        temporary = None
    finally:
        if temporary is not None:
            temporary.unlink(missing_ok=True)


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _sha256_json(payload: Any) -> str:
    encoded = json.dumps(
        payload,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _normal_path(path: Path) -> str:
    return os.path.normcase(os.path.abspath(str(path)))


def _clean_error(value: str, limit: int = 600) -> str:
    cleaned = " ".join(str(value).split())
    return cleaned[:limit] + ("..." if len(cleaned) > limit else "")
