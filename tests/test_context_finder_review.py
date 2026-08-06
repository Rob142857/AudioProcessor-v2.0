from __future__ import annotations

from dataclasses import dataclass
import json
from pathlib import Path
import tempfile
import threading
import time
from types import SimpleNamespace
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

from cleanup_client import DEFAULT_MODEL
from context_finder import (
    COMPILATION_MARKER,
    SearchOptions,
    compile_query_pattern,
    create_compilation_docx,
    find_contexts,
)
from context_finder_review import (
    REVIEW_NOTE_PREFIX,
    ContextFinderReviewError,
    ReviewOutputConflictError,
    ReviewSourceIntegrityError,
    ReviewStructureError,
    create_glm_review_copy,
    default_review_output_path,
)
from pipeline_control import PipelineCancelledError, raise_if_cancelled


GLOSSARY_SHA256 = "a" * 64


@dataclass(frozen=True)
class _Glossary:
    sha256: str = GLOSSARY_SHA256
    count: int = 3
    pinned: bool = True
    terms: tuple[str, ...] = ("awakening", "esotericism", "enneagram")


class _FakeCleanupClient:
    def __init__(
        self,
        transform=lambda text: text,
        *,
        needs_review: bool = False,
        delay: float = 0.0,
        on_call=None,
        raise_on_cleanup: bool = False,
    ) -> None:
        self.model = DEFAULT_MODEL
        self.transform = transform
        self.needs_review = needs_review
        self.delay = delay
        self.on_call = on_call
        self.raise_on_cleanup = raise_on_cleanup
        self.ensure_calls = 0
        self.cleanup_calls = 0
        self.inputs: list[str] = []
        self.checkpoint_dirs: list[Path] = []
        self.active = 0
        self.maximum_active = 0
        self._lock = threading.Lock()

    def ensure_glossary(self, *, cancel_check=None):
        raise_if_cancelled(cancel_check, phase="fake glossary")
        self.ensure_calls += 1
        return _Glossary()

    def cleanup_text(
        self,
        text: str,
        checkpoint_dir: Path | None = None,
        *,
        reuse_checkpoints: bool = True,
        cancel_check=None,
    ):
        del reuse_checkpoints
        raise_if_cancelled(cancel_check, phase="fake cleanup")
        with self._lock:
            self.cleanup_calls += 1
            self.inputs.append(text)
            self.checkpoint_dirs.append(Path(checkpoint_dir or "."))
            self.active += 1
            self.maximum_active = max(self.maximum_active, self.active)
            call_number = self.cleanup_calls
        try:
            if self.raise_on_cleanup:
                raise AssertionError("cleanup should have resumed without a remote call")
            if self.on_call is not None:
                self.on_call(call_number, text)
            if self.delay:
                time.sleep(self.delay)
            corrected = self.transform(text)
            raise_if_cancelled(cancel_check, phase="fake cleanup")
            return SimpleNamespace(
                text=corrected,
                model=self.model,
                glossary_sha256=GLOSSARY_SHA256,
                glossary_count=3,
                needs_review=self.needs_review,
                warnings=("service requested a human check",)
                if self.needs_review
                else (),
                chunks=(),
            )
        finally:
            with self._lock:
                self.active -= 1


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def _make_compilation(
    root: Path,
    *,
    source_count: int = 1,
    query: str = "awakening",
) -> Path:
    library = root / "library"
    library.mkdir()
    for index in range(1, source_count + 1):
        _write_docx(
            library / f"Lecture {index:02d}.docx",
            [
                f"{query.title()} is here in lecture {index} and needs polish.",
                f"This nearby sentence supplies context for lecture {index}.",
            ],
        )
    result = find_contexts(
        library,
        query,
        options=SearchOptions(context_words_each_side=100),
    )
    if len(result.regions) != source_count:
        raise AssertionError(
            f"fixture expected {source_count} regions, got {len(result.regions)}"
        )
    output = root / query.title().replace(" ", " ")
    output = output.with_suffix(".docx")
    create_compilation_docx(result, output)
    return output


def _style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style is not None else ""


def _quote_paragraphs(path: Path):
    document = Document(path)
    paragraphs = document.paragraphs
    quotes = []
    in_section = False
    metadata_left = 0
    for paragraph in paragraphs:
        if _style_name(paragraph) == "Heading 1":
            in_section = True
            metadata_left = 2
            continue
        if not in_section:
            continue
        if metadata_left:
            metadata_left -= 1
            continue
        quotes.append(paragraph)
    return document, quotes


def _navigation_snapshot(path: Path) -> dict[str, object]:
    document = Document(path)
    headings = [
        (paragraph.text, _style_name(paragraph), paragraph._p.xml)
        for paragraph in document.paragraphs
        if _style_name(paragraph) == "Heading 1"
    ]
    metadata = [
        (paragraph.text, _style_name(paragraph), paragraph._p.xml)
        for paragraph in document.paragraphs
        if _style_name(paragraph) == "Context Source Metadata"
    ]
    hyperlinks = sorted(
        (relationship.reltype, str(relationship.target_ref), relationship.is_external)
        for relationship in document.part.rels.values()
        if relationship.is_external
    )
    return {
        "headings": headings,
        "metadata": metadata,
        "hyperlinks": hyperlinks,
        "headers": [section.header._element.xml for section in document.sections],
        "footers": [section.footer._element.xml for section in document.sections],
        "subject": document.core_properties.subject,
    }


class ContextFinderReviewTests(unittest.TestCase):
    def test_review_copy_reparagraphs_body_and_preserves_navigation(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = _make_compilation(root)
            source_bytes = source.read_bytes()
            before = _navigation_snapshot(source)
            updates = []
            client = _FakeCleanupClient(
                lambda text: text.replace(
                    "needs polish.",
                    "has been carefully polished.\n\nA coherent closing paragraph follows.",
                )
            )

            outcome = create_glm_review_copy(
                source,
                cleanup_client=client,
                checkpoint_dir=root / "checkpoints",
                progress_callback=updates.append,
            )

            self.assertEqual(source_bytes, source.read_bytes())
            self.assertEqual(default_review_output_path(source), outcome.output_path)
            self.assertTrue(outcome.output_path.is_file())
            self.assertEqual(1, outcome.reviewed_regions)
            self.assertEqual(0, outcome.fallback_regions)
            self.assertEqual(1, outcome.occurrence_count)
            self.assertTrue(outcome.needs_human_review)
            self.assertEqual(1, client.ensure_calls)
            self.assertEqual(1, client.cleanup_calls)
            self.assertNotIn("Source:", client.inputs[0])
            self.assertNotIn("Section 1", client.inputs[0])
            self.assertEqual(before, _navigation_snapshot(outcome.output_path))
            reviewed, quotes = _quote_paragraphs(outcome.output_path)
            self.assertTrue(
                any(
                    paragraph.text.startswith(REVIEW_NOTE_PREFIX)
                    for paragraph in reviewed.paragraphs
                )
            )
            self.assertTrue(
                any("A coherent closing paragraph follows." in item.text for item in quotes)
            )
            self.assertTrue(
                all(item.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY for item in quotes)
            )
            highlighted = [
                run
                for paragraph in quotes
                for run in paragraph.runs
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW
            ]
            self.assertEqual(1, len(highlighted))
            self.assertEqual("complete", updates[-1].phase)
            self.assertEqual(COMPILATION_MARKER, reviewed.core_properties.subject)

    def test_query_count_change_falls_back_to_exact_body(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = _make_compilation(root)
            _document, exact_quotes = _quote_paragraphs(source)
            exact_text = [paragraph.text for paragraph in exact_quotes]
            client = _FakeCleanupClient(
                lambda text: text.replace("Awakening", "Sleeping")
            )

            outcome = create_glm_review_copy(
                source,
                cleanup_client=client,
                checkpoint_dir=root / "checkpoints",
            )

            _review, review_quotes = _quote_paragraphs(outcome.output_path)
            self.assertEqual(exact_text, [paragraph.text for paragraph in review_quotes])
            self.assertEqual(1, outcome.fallback_regions)
            self.assertEqual(1, outcome.needs_review_regions)
            self.assertIn("occurrence count", " ".join(outcome.warnings))
            pattern = compile_query_pattern("awakening")
            self.assertEqual(
                1,
                sum(len(tuple(pattern.finditer(item.text))) for item in review_quotes),
            )

    def test_default_three_workers_and_atomic_region_resume(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = _make_compilation(root, source_count=6)
            checkpoints = root / "checkpoints"
            client = _FakeCleanupClient(
                lambda text: text.replace("needs polish", "is polished"),
                delay=0.06,
            )

            first = create_glm_review_copy(
                source,
                cleanup_client=client,
                checkpoint_dir=checkpoints,
            )

            self.assertEqual(3, client.maximum_active)
            self.assertEqual(1, client.ensure_calls)
            self.assertEqual(6, client.cleanup_calls)
            self.assertEqual(6, len(set(client.checkpoint_dirs)))
            manifest = json.loads(first.manifest_path.read_text(encoding="utf-8"))
            self.assertEqual("complete", manifest["status"])
            self.assertEqual(
                list(range(1, 7)),
                [record["section_index"] for record in manifest["sections"]],
            )
            manifest_text = first.manifest_path.read_text(encoding="utf-8")
            self.assertNotIn("Awakening is here", manifest_text)
            self.assertFalse(list(first.manifest_path.parent.glob("*.tmp")))

            resumed_client = _FakeCleanupClient(raise_on_cleanup=True)
            second = create_glm_review_copy(
                source,
                cleanup_client=resumed_client,
                checkpoint_dir=checkpoints,
            )
            self.assertEqual(6, second.resumed_regions)
            self.assertEqual(0, resumed_client.cleanup_calls)
            self.assertEqual(1, resumed_client.ensure_calls)

    def test_unmanaged_and_manually_changed_outputs_fail_before_glossary(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = _make_compilation(root)
            output = default_review_output_path(source)
            output.write_bytes(b"not a managed review")
            client = _FakeCleanupClient()
            with self.assertRaises(ReviewOutputConflictError):
                create_glm_review_copy(
                    source,
                    cleanup_client=client,
                    checkpoint_dir=root / "checkpoints",
                )
            self.assertEqual(0, client.ensure_calls)
            self.assertEqual(b"not a managed review", output.read_bytes())

            output.unlink()
            first = create_glm_review_copy(
                source,
                cleanup_client=_FakeCleanupClient(),
                checkpoint_dir=root / "managed-checkpoints",
            )
            first.output_path.write_bytes(first.output_path.read_bytes() + b"manual edit")
            second_client = _FakeCleanupClient()
            with self.assertRaises(ReviewOutputConflictError):
                create_glm_review_copy(
                    source,
                    cleanup_client=second_client,
                    checkpoint_dir=root / "managed-checkpoints",
                )
            self.assertEqual(0, second_client.ensure_calls)

    def test_valid_needs_review_correction_is_retained(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = _make_compilation(root)
            client = _FakeCleanupClient(
                lambda text: text.replace("needs polish", "needs human checking"),
                needs_review=True,
            )

            outcome = create_glm_review_copy(
                source,
                cleanup_client=client,
                checkpoint_dir=root / "checkpoints",
            )

            _document, quotes = _quote_paragraphs(outcome.output_path)
            self.assertTrue(any("needs human checking" in item.text for item in quotes))
            self.assertEqual(1, outcome.reviewed_regions)
            self.assertEqual(1, outcome.needs_review_regions)
            self.assertIn("service requested a human check", outcome.warnings)

    def test_source_mutation_before_publication_fails_closed(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = _make_compilation(root)

            def mutate_source(_call_number, _text):
                source.write_bytes(source.read_bytes() + b"changed")

            client = _FakeCleanupClient(on_call=mutate_source)
            with self.assertRaises(ReviewSourceIntegrityError):
                create_glm_review_copy(
                    source,
                    cleanup_client=client,
                    checkpoint_dir=root / "checkpoints",
                )
            self.assertFalse(default_review_output_path(source).exists())

    def test_cancellation_publishes_no_partial_docx(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = _make_compilation(root, source_count=2)
            cancelled = threading.Event()

            def stop_after_call(_call_number, _text):
                cancelled.set()

            client = _FakeCleanupClient(on_call=stop_after_call)
            with self.assertRaises(PipelineCancelledError):
                create_glm_review_copy(
                    source,
                    cleanup_client=client,
                    checkpoint_dir=root / "checkpoints",
                    max_workers=1,
                    cancel_check=cancelled.is_set,
                )
            self.assertFalse(default_review_output_path(source).exists())
            manifests = list((root / "checkpoints").rglob("review-manifest.json"))
            self.assertEqual(1, len(manifests))
            manifest = json.loads(manifests[0].read_text(encoding="utf-8"))
            self.assertEqual("cancelled", manifest["status"])
            self.assertFalse(list(root.glob(".*.tmp.docx")))

    def test_plain_docx_is_rejected_without_remote_work(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = root / "Awakening.docx"
            _write_docx(source, ["Awakening is mentioned here."])
            client = _FakeCleanupClient()
            with self.assertRaises(ReviewStructureError):
                create_glm_review_copy(
                    source,
                    cleanup_client=client,
                    checkpoint_dir=root / "checkpoints",
                )
            self.assertEqual(0, client.ensure_calls)
            self.assertEqual(0, client.cleanup_calls)

    def test_unpinned_glossary_is_rejected(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            source = _make_compilation(root)
            client = _FakeCleanupClient()
            client.ensure_glossary = lambda **_kwargs: _Glossary(pinned=False)
            with self.assertRaises(ContextFinderReviewError):
                create_glm_review_copy(
                    source,
                    cleanup_client=client,
                    checkpoint_dir=root / "checkpoints",
                )
            self.assertEqual(0, client.cleanup_calls)


if __name__ == "__main__":
    unittest.main()
