from __future__ import annotations

from dataclasses import FrozenInstanceError
import json
from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

from context_compilation_inventory import (
    INVENTORY_SCHEMA_VERSION,
    ContextCompilationPairMismatchError,
    ContextCompilationRecordError,
    ContextCompilationStructureError,
    bind_context_compilation,
)
from context_finder import (
    SearchOptions,
    create_compilation_docx,
    find_contexts,
    write_result_records,
)


_HYPERLINK_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def _make_pair(root: Path) -> tuple[Path, Path]:
    library = root / "library"
    library.mkdir()
    _write_docx(
        library / "Lecture A.docx",
        [
            "The first theme begins.",
            "Awakening must occur.",
            "",
            "The first theme resolves.",
        ],
    )
    _write_docx(
        library / "Lecture B.docx",
        ["A second Awakening appears in another lecture."],
    )
    result = find_contexts(
        library,
        "awakening",
        options=SearchOptions(context_words_each_side=100),
    )
    jsonl = root / "Awakening.jsonl"
    docx = root / "Awakening.docx"
    write_result_records(result, jsonl)
    create_compilation_docx(result, docx)
    return docx, jsonl


def _first_paragraph_with_style(document, style_name: str):
    return next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style is not None and paragraph.style.name == style_name
    )


def _first_quote(document):
    in_section = False
    metadata_left = 0
    for paragraph in document.paragraphs:
        style = paragraph.style.name if paragraph.style is not None else ""
        if style == "Heading 1":
            in_section = True
            metadata_left = 2
            continue
        if not in_section:
            continue
        if metadata_left:
            metadata_left -= 1
            continue
        return paragraph
    raise AssertionError("fixture contains no quotation paragraph")


class ContextCompilationInventoryTests(unittest.TestCase):
    def test_binds_canonical_jsonl_order_and_preserves_empty_selection(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx, jsonl = _make_pair(root)

            bound = bind_context_compilation(docx)
            rebound = bind_context_compilation(docx, jsonl)
            jsonl_region_ids = [
                json.loads(line)["region_id"]
                for line in jsonl.read_text(encoding="utf-8").splitlines()[1:]
            ]

            self.assertEqual(INVENTORY_SCHEMA_VERSION, bound.schema_version)
            self.assertEqual(jsonl.resolve(), bound.jsonl_path)
            self.assertEqual(jsonl_region_ids, [item.region_id for item in bound.regions])
            self.assertEqual([1, 2], [item.ordinal for item in bound.regions])
            self.assertEqual(2, bound.region_count)
            self.assertEqual(2, bound.occurrence_count)
            self.assertEqual(2, bound.source_count)
            self.assertEqual(2, bound.highlight_count)
            self.assertEqual(1, bound.empty_selected_paragraph_count)
            self.assertEqual(
                bound.selected_paragraph_count - 1,
                bound.emitted_paragraph_count,
            )
            empty = [
                paragraph
                for region in bound.regions
                for paragraph in region.selected_paragraphs
                if not paragraph.emitted
            ]
            self.assertEqual(1, len(empty))
            self.assertEqual("", empty[0].text)
            self.assertNotIn("\n\n\n\n", bound.regions[0].selected_text)
            self.assertEqual(bound.pair_fingerprint, rebound.pair_fingerprint)
            for digest in (
                bound.docx_sha256,
                bound.jsonl_sha256,
                bound.pair_fingerprint,
                bound.ordered_regions_sha256,
                bound.source_manifest_sha256,
                bound.regions[0].selected_text_sha256,
                bound.regions[0].selected_paragraphs_sha256,
                bound.regions[0].region_fingerprint,
            ):
                self.assertRegex(digest, r"^[0-9a-f]{64}$")
            with self.assertRaises(FrozenInstanceError):
                bound.region_count = 99  # type: ignore[misc]

    def test_rejects_body_or_metadata_that_differs_from_same_ordinal_region(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx, jsonl = _make_pair(root)
            document = Document(docx)
            _first_quote(document).text = "Awakening was changed in the presentation."
            document.save(docx)
            with self.assertRaises(ContextCompilationPairMismatchError):
                bind_context_compilation(docx, jsonl)

        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx, jsonl = _make_pair(root)
            document = Document(docx)
            metadata = _first_paragraph_with_style(
                document, "Context Source Metadata"
            )
            metadata.text = metadata.text.replace("Section 1", "Section 99")
            document.save(docx)
            with self.assertRaises(ContextCompilationPairMismatchError):
                bind_context_compilation(docx, jsonl)

    def test_rejects_wrong_source_hyperlink_target(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx, jsonl = _make_pair(root)
            document = Document(docx)
            source_link = [
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.startswith("Source: ")
            ][0]
            hyperlink = source_link._p.xpath(".//w:hyperlink")[0]
            wrong_id = document.part.relate_to(
                "file:///C:/wrong-source.docx",
                _HYPERLINK_RELATIONSHIP,
                is_external=True,
            )
            hyperlink.set(qn("r:id"), wrong_id)
            document.save(docx)

            with self.assertRaises(ContextCompilationPairMismatchError):
                bind_context_compilation(docx, jsonl)

    def test_rejects_missing_or_extra_query_highlighting(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx, jsonl = _make_pair(root)
            document = Document(docx)
            highlighted_run = next(
                run
                for paragraph in document.paragraphs
                for run in paragraph.runs
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW
            )
            highlighted_run.font.highlight_color = None
            document.save(docx)

            with self.assertRaises(ContextCompilationPairMismatchError):
                bind_context_compilation(docx, jsonl)

    def test_rejects_missing_generation_marker(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx, jsonl = _make_pair(root)
            document = Document(docx)
            document.core_properties.keywords = "not the compilation marker"
            document.save(docx)

            with self.assertRaises(ContextCompilationStructureError):
                bind_context_compilation(docx, jsonl)

    def test_rejects_manifest_count_mismatch_and_duplicate_region_id(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx, jsonl = _make_pair(root)
            lines = jsonl.read_text(encoding="utf-8").splitlines()
            manifest = json.loads(lines[0])
            manifest["occurrence_count"] += 1
            lines[0] = json.dumps(manifest, sort_keys=True)
            jsonl.write_text("\n".join(lines) + "\n", encoding="utf-8")
            with self.assertRaises(ContextCompilationRecordError):
                bind_context_compilation(docx, jsonl)

        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx, jsonl = _make_pair(root)
            lines = jsonl.read_text(encoding="utf-8").splitlines()
            first = json.loads(lines[1])
            second = json.loads(lines[2])
            second["region_id"] = first["region_id"]
            lines[2] = json.dumps(second, sort_keys=True)
            jsonl.write_text("\n".join(lines) + "\n", encoding="utf-8")
            with self.assertRaises(ContextCompilationRecordError):
                bind_context_compilation(docx, jsonl)

    def test_jsonl_record_order_is_canonical_and_cannot_be_reordered_silently(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx, jsonl = _make_pair(root)
            lines = jsonl.read_text(encoding="utf-8").splitlines()
            lines[1], lines[2] = lines[2], lines[1]
            jsonl.write_text("\n".join(lines) + "\n", encoding="utf-8")

            with self.assertRaises(ContextCompilationPairMismatchError):
                bind_context_compilation(docx, jsonl)


if __name__ == "__main__":
    unittest.main()
