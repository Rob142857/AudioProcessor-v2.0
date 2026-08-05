from __future__ import annotations

import json
from pathlib import Path
import tempfile
import unittest
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from context_finder import (
    COMPILATION_MARKER,
    MAX_REGION_PARAGRAPH_MAP_BYTES,
    MAX_REGION_PARAGRAPH_SPAN,
    SearchOptions,
    SourceIntegrityError,
    apply_boundary_selection,
    create_compilation_docx,
    find_contexts,
    read_result_records,
    validate_source_integrity,
    validate_query,
    write_result_records,
)


def write_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


class QueryValidationTests(unittest.TestCase):
    def test_accepts_one_to_three_words_and_normalises_spacing(self):
        query = validate_query("  self   remembrance  ")
        self.assertEqual("self remembrance", query.text)
        self.assertEqual("self remembrance", query.canonical)
        self.assertEqual(2, query.word_count)
        self.assertEqual("Gurdjieff's", validate_query("Gurdjieff's").text)

    def test_rejects_empty_long_or_punctuation_wrapped_queries(self):
        for value in ("", "one two three four", '"awakening"'):
            with self.subTest(value=value), self.assertRaises(ValueError):
                validate_query(value)


class ExactContextSearchTests(unittest.TestCase):
    def test_glm_review_supersedes_only_its_same_directory_raw_sibling(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            nested = root / "nested"
            nested.mkdir()
            write_docx(root / "lecture.docx", ["Awakening raw copy."])
            write_docx(
                root / "lecture - GLM Review.docx",
                ["Awakening reviewed copy."],
            )
            write_docx(nested / "lecture.docx", ["Awakening unrelated copy."])
            write_docx(
                root / "orphan - GLM Review.docx",
                ["Awakening orphan review."],
            )

            result = find_contexts(
                root,
                "awakening",
                options=SearchOptions(context_words_each_side=0),
            )

            paths = {region.source_relative_path for region in result.regions}
            self.assertEqual(3, result.scanned_files)
            self.assertEqual(3, result.occurrence_count)
            self.assertNotIn("lecture.docx", paths)
            self.assertIn("lecture - GLM Review.docx", paths)
            self.assertIn("nested/lecture.docx", paths)
            self.assertIn("orphan - GLM Review.docx", paths)

    def test_dense_hits_split_without_losing_or_duplicating_occurrences(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            paragraphs = [f"Awakening marker {index}." for index in range(100)]
            write_docx(root / "dense.docx", paragraphs)

            result = find_contexts(
                root,
                "awakening",
                options=SearchOptions(
                    context_words_each_side=10_000,
                    max_region_paragraphs=6,
                    max_region_characters=100_000,
                ),
            )

            occurrence_ids = [
                occurrence.occurrence_id
                for region in result.regions
                for occurrence in region.occurrences
            ]
            self.assertGreater(len(result.regions), 1)
            self.assertEqual(100, result.occurrence_count)
            self.assertEqual(100, len(set(occurrence_ids)))
            for region in result.regions:
                self.assertLessEqual(
                    len([paragraph for paragraph in region.paragraphs if paragraph.text]),
                    6,
                )
                self.assertLessEqual(len(region.paragraphs), MAX_REGION_PARAGRAPH_SPAN)

    def test_region_paragraph_map_stays_inside_utf8_endpoint_budget(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            write_docx(
                root / "unicode.docx",
                [("😀" * 1_800) + f" awakening {index}." for index in range(12)],
            )

            result = find_contexts(
                root,
                "awakening",
                options=SearchOptions(
                    context_words_each_side=10_000,
                    max_region_characters=100_000,
                ),
            )

            self.assertEqual(12, result.occurrence_count)
            for region in result.regions:
                paragraph_map = [
                    {"number": paragraph.number, "text": paragraph.text}
                    for paragraph in region.paragraphs
                    if paragraph.text
                ]
                encoded = json.dumps(
                    paragraph_map,
                    ensure_ascii=False,
                    separators=(",", ":"),
                ).encode("utf-8")
                self.assertLessEqual(len(encoded), MAX_REGION_PARAGRAPH_MAP_BYTES)

    def test_recursive_whole_phrase_search_merges_overlapping_windows(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            nested = root / "nested"
            nested.mkdir()
            source = nested / "lecture.docx"
            write_docx(
                source,
                [
                    "Opening words for the lecture.",
                    "At the bus stop we waited.",
                    "A connecting paragraph remains exact.",
                    "The BUS STOP appeared in the example again.",
                    "Closing words for the lecture.",
                    "A bus stopper and bus-stop are not exact phrase matches.",
                ],
            )

            result = find_contexts(
                root,
                "bus stop",
                options=SearchOptions(context_words_each_side=6),
            )

            self.assertEqual(1, result.scanned_files)
            self.assertEqual(1, result.source_count)
            self.assertEqual(2, result.occurrence_count)
            self.assertEqual(1, len(result.regions))
            region = result.regions[0]
            self.assertEqual("nested/lecture.docx", region.source_relative_path)
            self.assertEqual(("bus stop", "BUS STOP"), tuple(
                occurrence.matched_text for occurrence in region.occurrences
            ))
            self.assertIn(
                "A connecting paragraph remains exact.",
                [paragraph.text for paragraph in region.paragraphs],
            )

    def test_whole_word_does_not_match_inside_a_longer_word(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            write_docx(
                root / "lecture.docx",
                ["Awakening is named. Reawakening and awakenings are different."],
            )
            result = find_contexts(
                root,
                "awakening",
                options=SearchOptions(context_words_each_side=0),
            )
            self.assertEqual(1, result.occurrence_count)
            self.assertEqual("Awakening", result.regions[0].occurrences[0].matched_text)

    def test_word_does_not_match_inside_hyphen_or_apostrophe_compounds(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            write_docx(
                root / "lecture.docx",
                ["Wake now. A wake-up follows; wake's echo and wake’s echo do not count."],
            )
            result = find_contexts(
                root,
                "wake",
                options=SearchOptions(context_words_each_side=0),
            )
            self.assertEqual(1, result.occurrence_count)
            self.assertEqual("Wake", result.regions[0].occurrences[0].matched_text)

    def test_txt_and_markdown_are_supported(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            (root / "notes.txt").write_text(
                "First paragraph.\n\nAwakening must occur.\n",
                encoding="utf-8",
            )
            (root / "notes.md").write_text(
                "# Notes\n\nA second awakening appears.\n",
                encoding="utf-8",
            )
            result = find_contexts(root, "awakening")
            self.assertEqual(2, result.scanned_files)
            self.assertEqual(2, result.occurrence_count)
            self.assertEqual(2, result.source_count)

    def test_unreadable_docx_is_reported_without_hiding_other_results(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            (root / "broken.docx").write_text("not a zip", encoding="utf-8")
            (root / "good.txt").write_text("awakening", encoding="utf-8")
            result = find_contexts(root, "awakening")
            self.assertEqual(1, result.scanned_files)
            self.assertEqual(1, result.occurrence_count)
            self.assertEqual(1, len(result.issues))
            self.assertEqual("broken.docx", result.issues[0].source_relative_path)


class ResumeAndBoundaryTests(unittest.TestCase):
    def test_ids_are_stable_and_jsonl_round_trip_is_lossless(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            write_docx(
                root / "lecture.docx",
                ["Before.", "Awakening must occur.", "After."],
            )
            first = find_contexts(root, "awakening")
            second = find_contexts(root, "awakening")
            self.assertEqual(
                [region.region_id for region in first.regions],
                [region.region_id for region in second.regions],
            )

            records = root / "results.jsonl"
            write_result_records(first, records)
            loaded = read_result_records(records)
            self.assertEqual(first, loaded)
            lines = records.read_text(encoding="utf-8").splitlines()
            self.assertEqual("search_manifest", json.loads(lines[0])["record_type"])

    def test_boundary_selection_can_only_choose_exact_stored_paragraphs(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            write_docx(
                root / "lecture.docx",
                [
                    "Distant introduction.",
                    "The theme begins here.",
                    "Awakening must occur.",
                    "The theme resolves here.",
                    "A new theme begins.",
                ],
            )
            region = find_contexts(root, "awakening").regions[0]
            refined = apply_boundary_selection(
                region,
                2,
                4,
                model="@cf/zai-org/glm-4.7-flash",
                confidence=0.98,
            )
            self.assertEqual(
                (
                    "The theme begins here.",
                    "Awakening must occur.",
                    "The theme resolves here.",
                ),
                tuple(paragraph.text for paragraph in refined.selected_paragraphs),
            )
            self.assertEqual("glm_boundary_refinement", refined.selection.method)
            with self.assertRaises(ValueError):
                apply_boundary_selection(region, 1, 2)
            with self.assertRaises(ValueError):
                find_contexts(root, "awakening").with_regions((region, region))

            payload = region.boundary_payload()
            self.assertEqual([3], payload["must_include_paragraphs"])
            self.assertEqual(
                "Awakening must occur.",
                payload["paragraphs"][2]["text"],
            )


class CompilationTests(unittest.TestCase):
    def test_publication_fails_closed_when_a_source_changed_after_scan(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "lecture.docx"
            write_docx(source, ["Awakening must occur."])
            result = find_contexts(root, "awakening")
            validate_source_integrity(result)

            write_docx(source, ["The source was edited after scanning."])
            output = root / "contexts.docx"
            with self.assertRaises(SourceIntegrityError):
                validate_source_integrity(result)
            with self.assertRaises(SourceIntegrityError):
                create_compilation_docx(result, output)
            self.assertFalse(output.exists())

    def test_compilation_is_verbatim_justified_highlighted_and_linked(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "lecture.docx"
            exact_quote = "Now Awakening must occur, without paraphrase."
            write_docx(source, [exact_quote])
            result = find_contexts(
                root,
                "awakening",
                options=SearchOptions(context_words_each_side=0),
            )
            output = root.parent / f"{root.name}-context.docx"
            create_compilation_docx(result, output)

            compiled = Document(output)
            quote_paragraphs = [
                paragraph for paragraph in compiled.paragraphs if paragraph.text == exact_quote
            ]
            self.assertEqual(1, len(quote_paragraphs))
            self.assertEqual(WD_ALIGN_PARAGRAPH.JUSTIFY, quote_paragraphs[0].alignment)
            self.assertEqual(COMPILATION_MARKER, compiled.core_properties.subject)

            with ZipFile(output) as package:
                document_xml = package.read("word/document.xml").decode("utf-8")
                rels_xml = package.read("word/_rels/document.xml.rels").decode("utf-8")
            self.assertIn('w:highlight w:val="yellow"', document_xml)
            self.assertIn("TargetMode=\"External\"", rels_xml)
            self.assertIn("file:///", rels_xml)

    def test_generated_compilation_is_ignored_on_a_later_recursive_search(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            write_docx(root / "source.docx", ["Awakening must occur."])
            first = find_contexts(root, "awakening")
            create_compilation_docx(first, root / "old-results.docx")
            second = find_contexts(root, "awakening")
            self.assertEqual(1, second.scanned_files)
            self.assertEqual(1, second.ignored_generated_files)
            self.assertEqual(1, second.occurrence_count)


if __name__ == "__main__":
    unittest.main()
