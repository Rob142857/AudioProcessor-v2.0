from __future__ import annotations

import hashlib
import os
from pathlib import Path
import stat
import tempfile
import types
import unittest
from unittest import mock
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

import existing_transcript_import as importer
from existing_transcript_import import (
    EXTRACTOR_VERSION,
    ExistingTranscriptImportError,
    import_existing_transcript,
)


BODY_ONE = (
    "Today we are considering the model: a map of consciousness, not a machine. "
    "The lecture wording must remain exactly within this first substantial paragraph."
)
BODY_TWO = (
    "Transcription Information is also an ordinary phrase when it appears in a lecture. "
    "This second paragraph ensures that such wording is never stripped from the body."
)


def write_legacy_docx(path: Path, *, body: tuple[str, ...] = (BODY_ONE, BODY_TWO)) -> None:
    document = Document()
    document.add_paragraph("The Enneagram")
    document.add_paragraph("Lecture 10 given on 22 January 1985")
    document.add_paragraph("Transcript:")
    for paragraph in body:
        document.add_paragraph(paragraph)
    document.add_paragraph("____________________________")
    document.add_paragraph("Transcription Information")
    document.add_paragraph("Model: Faster-Whisper large-v3")
    document.add_paragraph("Device: CUDA")
    document.add_paragraph("Processing Time: 15 minutes")
    document.add_paragraph("Audio Preprocessing: vintage tape")
    document.add_paragraph("(This information can be deleted if not needed)")
    document.save(path)


def rewrite_docx_part(path: Path, name: str, transform) -> None:
    rewritten = path.with_name(f"{path.stem}.rewritten.docx")
    with ZipFile(path, "r") as source_package:
        members = [
            (member, source_package.read(member.filename))
            for member in source_package.infolist()
        ]
    with ZipFile(rewritten, "w", compression=ZIP_DEFLATED) as output_package:
        for member, payload in members:
            if member.filename == name:
                payload = transform(payload)
            output_package.writestr(member, payload)
    rewritten.replace(path)


def add_docx_part(path: Path, name: str, payload: bytes) -> None:
    rewritten = path.with_name(f"{path.stem}.rewritten.docx")
    with ZipFile(path, "r") as source_package:
        members = [
            (member, source_package.read(member.filename))
            for member in source_package.infolist()
        ]
    with ZipFile(rewritten, "w", compression=ZIP_DEFLATED) as output_package:
        for member, existing_payload in members:
            output_package.writestr(member, existing_payload)
        output_package.writestr(name, payload)
    rewritten.replace(path)


def inject_document_xml(path: Path, fragment: str) -> None:
    encoded_fragment = fragment.encode("utf-8")

    def transform(payload: bytes) -> bytes:
        marker = b"</w:body>"
        if marker not in payload:
            raise AssertionError("synthetic DOCX has no w:body closing tag")
        return payload.replace(marker, encoded_fragment + marker, 1)

    rewrite_docx_part(path, "word/document.xml", transform)


class ExistingTranscriptImportTests(unittest.TestCase):
    def test_extracts_only_body_and_records_audit_metadata_read_only(self):
        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "lecture.docx"
            write_legacy_docx(source)
            before = source.read_bytes()

            result = import_existing_transcript(source)

            after = source.read_bytes()
            source_mtime_ns = source.stat().st_mtime_ns

        self.assertEqual(result.text, f"{BODY_ONE}\n\n{BODY_TWO}")
        self.assertEqual(result.paragraph_count, 2)
        self.assertEqual(result.word_count, 50)
        self.assertEqual(result.source_sha256, hashlib.sha256(before).hexdigest())
        self.assertEqual(result.source_size, len(before))
        self.assertEqual(result.source_mtime_ns, source_mtime_ns)
        self.assertEqual(result.extractor_version, EXTRACTOR_VERSION)
        self.assertEqual(result.source_docx, os.path.abspath(os.fspath(source)))
        self.assertEqual(result.original_path, os.path.abspath(os.fspath(source)))
        self.assertEqual(before, after)
        record = result.to_dict()
        self.assertEqual(record["text"], result.text)
        self.assertEqual(record["source_docx"], os.path.abspath(os.fspath(source)))
        self.assertEqual(record["source_sha256"], result.source_sha256)
        self.assertEqual(record["source_size"], result.source_size)
        self.assertEqual(record["source_mtime_ns"], result.source_mtime_ns)
        self.assertEqual(record["extractor_version"], EXTRACTOR_VERSION)

    def test_retains_lecture_lines_that_resemble_workflow_labels(self):
        body = (
            BODY_ONE,
            "Model: the enneagram is a teaching aid, and this is lecture wording.",
            "Transcription Information",
            "This phrase now introduces a discussion rather than a generated footer.",
            BODY_TWO,
        )
        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "lecture.docx"
            write_legacy_docx(source, body=body)
            result = import_existing_transcript(source)

        self.assertEqual(result.text, "\n\n".join(body))
        self.assertEqual(result.paragraph_count, len(body))

    def test_preserves_manual_line_breaks_and_word_paragraph_boundaries(self):
        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "lecture.docx"
            document = Document()
            document.add_paragraph("Transcript:")
            first = document.add_paragraph()
            first.add_run("First line with enough words to form a trustworthy transcript body.")
            first.add_run().add_break()
            first.add_run("Second line remains in the same Word paragraph for fidelity.")
            document.add_paragraph(
                "A separate Word paragraph becomes a deterministic blank-line boundary in text."
            )
            document.save(source)

            result = import_existing_transcript(source)

        self.assertEqual(result.paragraph_count, 2)
        self.assertIn("body.\nSecond line", result.text)
        self.assertIn("fidelity.\n\nA separate Word paragraph", result.text)

    def test_allows_marker_only_document_without_generated_footer(self):
        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "lecture.docx"
            document = Document()
            document.add_paragraph("Generated title")
            document.add_paragraph("Transcript:")
            document.add_paragraph(BODY_ONE)
            document.add_paragraph(BODY_TWO)
            document.save(source)

            result = import_existing_transcript(source)

        self.assertEqual(result.text, f"{BODY_ONE}\n\n{BODY_TWO}")

    def test_strips_only_exact_known_generated_provenance_without_footer(self):
        provenance = (
            "Processed by speech-to-text from a digitised tape recording originally "
            "recorded in person by MW on 22 January 1985."
        )
        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "lecture.docx"
            document = Document()
            document.add_paragraph("Transcript:")
            document.add_paragraph(BODY_ONE)
            document.add_paragraph(BODY_TWO)
            document.add_paragraph(provenance)
            document.save(source)

            result = import_existing_transcript(source)

        self.assertNotIn("speech-to-text", result.text)
        self.assertTrue(result.text.endswith(BODY_TWO))

    def test_rejects_missing_or_ambiguous_transcript_marker(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            missing = root / "missing.docx"
            document = Document()
            document.add_paragraph(BODY_ONE)
            document.add_paragraph(BODY_TWO)
            document.save(missing)

            ambiguous = root / "ambiguous.docx"
            document = Document()
            document.add_paragraph("Transcript:")
            document.add_paragraph(BODY_ONE)
            document.add_paragraph("Transcript:")
            document.add_paragraph(BODY_TWO)
            document.save(ambiguous)

            with self.assertRaisesRegex(ExistingTranscriptImportError, "missing"):
                import_existing_transcript(missing)
            with self.assertRaisesRegex(ExistingTranscriptImportError, "ambiguous"):
                import_existing_transcript(ambiguous)

    def test_rejects_empty_and_suspiciously_tiny_body(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            empty = root / "empty.docx"
            document = Document()
            document.add_paragraph("Transcript:")
            document.save(empty)

            tiny = root / "tiny.docx"
            document = Document()
            document.add_paragraph("Transcript:")
            document.add_paragraph("Only three tiny words")
            document.save(tiny)

            with self.assertRaisesRegex(ExistingTranscriptImportError, "empty"):
                import_existing_transcript(empty)
            with self.assertRaisesRegex(ExistingTranscriptImportError, "suspiciously tiny"):
                import_existing_transcript(tiny)

    def test_rejects_missing_non_docx_corrupt_and_encrypted_container(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            with self.assertRaisesRegex(ExistingTranscriptImportError, "Expected a .docx"):
                import_existing_transcript(root / "lecture.txt")
            with self.assertRaisesRegex(ExistingTranscriptImportError, "does not exist"):
                import_existing_transcript(root / "missing.docx")

            corrupt = root / "corrupt.docx"
            corrupt.write_bytes(b"this is not a zip package")
            with self.assertRaisesRegex(ExistingTranscriptImportError, "not a valid DOCX"):
                import_existing_transcript(corrupt)

            empty_file = root / "zero-bytes.docx"
            empty_file.write_bytes(b"")
            with self.assertRaisesRegex(ExistingTranscriptImportError, "DOCX is empty"):
                import_existing_transcript(empty_file)

            encrypted = root / "encrypted.docx"
            encrypted.write_bytes(bytes.fromhex("D0CF11E0A1B11AE1") + b"encrypted package")
            with self.assertRaisesRegex(ExistingTranscriptImportError, "Encrypted"):
                import_existing_transcript(encrypted)

    def test_minimum_word_threshold_must_be_positive(self):
        with self.assertRaisesRegex(ValueError, "at least 1"):
            import_existing_transcript("unused.docx", minimum_words=0)

    def test_rejects_symbolic_links_and_windows_reparse_points(self):
        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "lecture.docx"
            write_legacy_docx(source)
            regular_mode = source.lstat().st_mode

            with mock.patch.object(
                Path,
                "lstat",
                return_value=types.SimpleNamespace(
                    st_mode=stat.S_IFLNK,
                    st_file_attributes=0,
                ),
            ):
                with self.assertRaisesRegex(ExistingTranscriptImportError, "symbolic link"):
                    import_existing_transcript(source)

            with mock.patch.object(
                Path,
                "lstat",
                return_value=types.SimpleNamespace(
                    st_mode=regular_mode,
                    st_file_attributes=getattr(
                        stat, "FILE_ATTRIBUTE_REPARSE_POINT", 0x400
                    ),
                ),
            ):
                with self.assertRaisesRegex(ExistingTranscriptImportError, "reparse point"):
                    import_existing_transcript(source)

    def test_rejects_path_swap_while_snapshot_is_read(self):
        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "lecture.docx"
            write_legacy_docx(source)
            original = source.lstat()
            swapped = types.SimpleNamespace(
                st_mode=original.st_mode,
                st_file_attributes=0,
                st_dev=original.st_dev,
                st_ino=original.st_ino + 1,
                st_size=original.st_size,
                st_mtime_ns=original.st_mtime_ns,
            )

            with mock.patch.object(Path, "lstat", side_effect=(original, swapped)):
                with self.assertRaisesRegex(
                    ExistingTranscriptImportError, "path changed while its snapshot"
                ):
                    import_existing_transcript(source)

    def test_enforces_zip_member_and_uncompressed_size_limits(self):
        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "lecture.docx"
            write_legacy_docx(source)
            with ZipFile(source, "r") as package:
                members = package.infolist()
                member_count = len(members)
                largest_member = max(member.file_size for member in members)
                total_uncompressed = sum(member.file_size for member in members)

            with mock.patch.object(importer, "MAX_DOCX_MEMBERS", member_count - 1):
                with self.assertRaisesRegex(ExistingTranscriptImportError, "too many ZIP members"):
                    import_existing_transcript(source)

            with mock.patch.object(
                importer,
                "MAX_DOCX_MEMBER_UNCOMPRESSED_BYTES",
                largest_member - 1,
            ):
                with self.assertRaisesRegex(
                    ExistingTranscriptImportError, "member exceeds the uncompressed size limit"
                ):
                    import_existing_transcript(source)

            with mock.patch.object(
                importer,
                "MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES",
                total_uncompressed - 1,
            ):
                with self.assertRaisesRegex(
                    ExistingTranscriptImportError, "total uncompressed size limit"
                ):
                    import_existing_transcript(source)

    def test_rejects_nonempty_tables_but_accepts_empty_layout_tables(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            nonempty = root / "nonempty-table.docx"
            write_legacy_docx(nonempty)
            document = Document(nonempty)
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Omitted table words"
            document.save(nonempty)
            with self.assertRaisesRegex(ExistingTranscriptImportError, "nonempty table"):
                import_existing_transcript(nonempty)

            empty = root / "empty-table.docx"
            write_legacy_docx(empty)
            document = Document(empty)
            document.add_table(rows=1, cols=1)
            document.save(empty)
            result = import_existing_transcript(empty)

        self.assertEqual(result.text, f"{BODY_ONE}\n\n{BODY_TWO}")

    def test_rejects_nonempty_headers_and_footers(self):
        for location in ("header", "footer"):
            with self.subTest(location=location), tempfile.TemporaryDirectory() as temporary:
                source = Path(temporary) / f"{location}.docx"
                write_legacy_docx(source)
                document = Document(source)
                part = getattr(document.sections[0], location)
                part.paragraphs[0].text = f"Semantic {location} text"
                document.save(source)

                with self.assertRaisesRegex(ExistingTranscriptImportError, "header/footer"):
                    import_existing_transcript(source)

    def test_rejects_tracked_insertions_and_deletions(self):
        fragments = {
            "insertion": (
                '<w:p><w:ins w:id="1"><w:r><w:t>Tracked insertion</w:t>'
                "</w:r></w:ins></w:p>"
            ),
            "deletion": (
                '<w:p><w:del w:id="2"><w:r><w:delText>Tracked deletion</w:delText>'
                "</w:r></w:del></w:p>"
            ),
        }
        for label, fragment in fragments.items():
            with self.subTest(change=label), tempfile.TemporaryDirectory() as temporary:
                source = Path(temporary) / f"tracked-{label}.docx"
                write_legacy_docx(source)
                inject_document_xml(source, fragment)

                with self.assertRaisesRegex(
                    ExistingTranscriptImportError, "tracked insertions/deletions"
                ):
                    import_existing_transcript(source)

    def test_rejects_text_boxes_and_altchunk_content(self):
        fragments = {
            "text-box": (
                "<w:txbxContent><w:p><w:r><w:t>Text box words</w:t>"
                "</w:r></w:p></w:txbxContent>"
            ),
            "altChunk": "<w:altChunk/>",
        }
        expected_messages = {
            "text-box": "text-box content",
            "altChunk": "altChunk content",
        }
        for label, fragment in fragments.items():
            with self.subTest(content=label), tempfile.TemporaryDirectory() as temporary:
                source = Path(temporary) / f"{label}.docx"
                write_legacy_docx(source)
                inject_document_xml(source, fragment)

                with self.assertRaisesRegex(
                    ExistingTranscriptImportError, expected_messages[label]
                ):
                    import_existing_transcript(source)

    def test_rejects_footnotes_and_endnotes_parts(self):
        note_parts = {
            "footnotes": (
                "word/footnotes.xml",
                b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/'
                b'wordprocessingml/2006/main"><w:footnote w:id="1"><w:p><w:r>'
                b"<w:t>Footnote words</w:t></w:r></w:p></w:footnote></w:footnotes>",
            ),
            "endnotes": (
                "word/endnotes.xml",
                b'<w:endnotes xmlns:w="http://schemas.openxmlformats.org/'
                b'wordprocessingml/2006/main"><w:endnote w:id="1"><w:p><w:r>'
                b"<w:t>Endnote words</w:t></w:r></w:p></w:endnote></w:endnotes>",
            ),
        }
        for label, (part_name, payload) in note_parts.items():
            with self.subTest(notes=label), tempfile.TemporaryDirectory() as temporary:
                source = Path(temporary) / f"{label}.docx"
                write_legacy_docx(source)
                add_docx_part(source, part_name, payload)

                with self.assertRaisesRegex(ExistingTranscriptImportError, "footnotes/endnotes"):
                    import_existing_transcript(source)


if __name__ == "__main__":
    unittest.main()
