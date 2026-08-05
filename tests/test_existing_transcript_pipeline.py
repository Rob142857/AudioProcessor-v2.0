import hashlib
import json
import os
import shutil
import tempfile
import unittest
from datetime import datetime, timedelta, timezone
from pathlib import Path
from types import SimpleNamespace
from unittest import mock

from docx import Document

from archive_pipeline import (
    PipelineConfig,
    PipelineRunner,
    artifact_directory,
    artifact_paths,
    discover_audio,
    publish_source_docx_batch,
)
from existing_transcript_import import import_existing_transcript
from legacy_docx_replace import ReplacementError, plan_legacy_docx_replacements


BODY_PARAGRAPHS = (
    "Awareness begins with careful observation of sensation, feeling, and thought, "
    "without forcing experience into a theory or a premature conclusion.",
    "When attention remains present, each ordinary event can reveal relationships "
    "that habit usually conceals, and understanding develops through direct study.",
)
BODY_TEXT = "\n\n".join(BODY_PARAGRAPHS)


def sha256_file(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def write_legacy_docx(path: Path, paragraphs=BODY_PARAGRAPHS) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Sensing life exercise", level=0)
    document.add_paragraph(
        "Lecture 10 given on 10th of December 1985 by Dr Philip W Groves"
    )
    document.add_paragraph("Transcript:")
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.add_paragraph("________________________________________")
    document.add_paragraph("Transcription Information")
    document.add_paragraph("Model: Faster-Whisper large-v3")
    document.add_paragraph("Device: CUDA GPU")
    document.add_paragraph("Processing Time: 0m 26.86s")
    document.add_paragraph("Audio Preprocessing: Vintage tape preset")
    document.add_paragraph("(This information can be deleted if not needed)")
    document.save(path)


class FakeCleanupResult:
    def __init__(self, text: str):
        self.text = text
        self.model = "@cf/zai-org/glm-4.7-flash"
        self.glossary_sha256 = "a" * 64
        self.glossary_count = 1604
        self.chunks = (object(),)
        self.needs_review = False
        self.warnings = ()

    def to_dict(self):
        return {
            "model": self.model,
            "glossary_sha256": self.glossary_sha256,
            "glossary_count": self.glossary_count,
            "chunks": [
                {
                    "grounding": {"glossary_terms_considered": 1604},
                    "quality": {"status": "passed"},
                }
            ],
            "needs_review": False,
            "warnings": [],
        }


class FakeCleanupClient:
    def __init__(self, *, result_text=None):
        self.calls = []
        self.reuse_checkpoints = []
        self.glossary_sha256 = "a" * 64
        self.result_text = result_text

    def ensure_glossary(self, *, cancel_check=None):
        return SimpleNamespace(sha256=self.glossary_sha256)

    def cleanup_text(
        self,
        text,
        checkpoint_dir=None,
        *,
        reuse_checkpoints=True,
        cancel_check=None,
    ):
        self.calls.append(text)
        self.reuse_checkpoints.append(reuse_checkpoints)
        return FakeCleanupResult(
            text if self.result_text is None else self.result_text
        )


class ExistingTranscriptPipelineTests(unittest.TestCase):
    def make_archive(self, root: Path, *, name: str = "1210 Sensing life ex"):
        archive = root / "archive"
        year = archive / "1985 MW"
        year.mkdir(parents=True)
        audio = year / f"{name}.mp3"
        audio.write_bytes(b"synthetic audio placeholder; must never be opened")
        transcript = audio.with_suffix(".docx")
        write_legacy_docx(transcript)
        output = root / "polished"
        return archive, audio, transcript, output

    @staticmethod
    def config(
        archive: Path,
        output: Path,
        *,
        force: bool = False,
        publish: bool = False,
    ) -> PipelineConfig:
        return PipelineConfig(
            input_path=archive,
            output_root=output,
            existing_transcripts_only=True,
            existing_docx_mode="all",
            cleanup_enabled=True,
            force=force,
            publish_source_docx=publish,
        )

    @staticmethod
    def install_fakes(runner: PipelineRunner, cleanup: FakeCleanupClient):
        renders = []
        runner.cleanup_client = cleanup
        runner._transcribe = mock.Mock(
            side_effect=AssertionError("Whisper must not run in imported-DOCX mode")
        )

        def render(source, text, output_path, metadata):
            renders.append(
                {
                    "source": source,
                    "text": text,
                    "output_path": output_path,
                    "metadata": metadata,
                }
            )
            document = Document()
            document.add_heading("Polished lecture", level=0)
            for paragraph in text.split("\n\n"):
                document.add_paragraph(paragraph)
            document.save(output_path)
            return output_path

        runner._render_docx = render
        return renders

    def test_discovery_selects_and_deduplicates_source_adjacent_docx_and_rejects_skip(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            archive, audio, transcript, output = self.make_archive(root)
            # Multiple recording formats share one already-created Word transcript.
            audio.with_suffix(".flac").write_bytes(b"second synthetic recording variant")
            (audio.parent / "missing-transcript.wav").write_bytes(b"no transcript")
            review = audio.with_name(f"{audio.stem} - GLM Review.docx")
            shutil.copyfile(transcript, review)

            selected = discover_audio(
                archive,
                output,
                existing_transcripts_only=True,
                existing_docx_mode="all",
            )

            self.assertEqual(selected, [transcript.resolve()])
            with self.assertRaisesRegex(ValueError, "generated output|GLM Review"):
                discover_audio(
                    review,
                    output,
                    existing_transcripts_only=True,
                    existing_docx_mode="all",
                )
            with self.assertRaisesRegex(ValueError, "Skip existing|select nothing"):
                discover_audio(
                    archive,
                    output,
                    existing_transcripts_only=True,
                    existing_docx_mode="skip",
                )

    def test_import_route_skips_whisper_and_timestamps_then_resumes_and_force_recleans(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            archive, _audio, transcript, output = self.make_archive(root)
            source = transcript.resolve()
            paths = artifact_paths(artifact_directory(source, archive, output))

            cleanup = FakeCleanupClient()
            runner = PipelineRunner(self.config(archive, output))
            renders = self.install_fakes(runner, cleanup)
            try:
                with mock.patch(
                    "existing_transcript_import.import_existing_transcript",
                    wraps=import_existing_transcript,
                ) as importer, mock.patch(
                    "archive_pipeline.probe_audio_duration_seconds",
                    side_effect=AssertionError("audio duration must not be probed"),
                ) as duration_probe:
                    self.assertEqual(runner.process_one(source), "verified")
                    importer.assert_called_once_with(source)
                    duration_probe.assert_not_called()

                runner._transcribe.assert_not_called()
                self.assertEqual(cleanup.calls, [BODY_TEXT])
                self.assertEqual([item["text"] for item in renders], [BODY_TEXT])
                self.assertFalse(paths["segments"].exists())
                self.assertFalse(paths["vtt"].exists())
                self.assertFalse(paths["srt"].exists())

                manifest = json.loads(paths["manifest"].read_text(encoding="utf-8"))
                self.assertEqual(
                    manifest["approval_state"], "pending_human_review"
                )
                self.assertEqual(
                    manifest["publication"]["approval_state"],
                    "pending_human_review",
                )
                self.assertIs(manifest["stt"]["performed"], False)
                self.assertEqual(manifest["stt"]["backend"], "imported-docx")
                self.assertIsNone(manifest["artifacts"]["segments"])
                self.assertIsNone(manifest["artifacts"]["vtt"])
                self.assertIsNone(manifest["artifacts"]["srt"])
                self.assertEqual(
                    manifest["qa"]["stt_coverage"]["status"], "not_applicable"
                )
                self.assertEqual(manifest["qa"]["raw_input"]["status"], "passed")
                self.assertNotIn(
                    BODY_PARAGRAPHS[0],
                    paths["manifest"].read_text(encoding="utf-8"),
                )

                raw_before = paths["raw_text"].read_bytes()
                with mock.patch(
                    "existing_transcript_import.import_existing_transcript",
                    side_effect=AssertionError("resume must not re-import"),
                ) as importer, mock.patch(
                    "archive_pipeline.probe_audio_duration_seconds"
                ) as duration_probe:
                    self.assertEqual(runner.process_one(source), "skipped")
                    importer.assert_not_called()
                    duration_probe.assert_not_called()
                self.assertEqual(cleanup.calls, [BODY_TEXT])
                self.assertEqual(len(renders), 1)
                self.assertEqual(paths["raw_text"].read_bytes(), raw_before)
            finally:
                runner.close()

            forced_cleanup = FakeCleanupClient()
            forced = PipelineRunner(self.config(archive, output, force=True))
            forced_renders = self.install_fakes(forced, forced_cleanup)
            try:
                with mock.patch(
                    "existing_transcript_import.import_existing_transcript",
                    side_effect=AssertionError("force must reuse preserved raw text"),
                ) as importer, mock.patch(
                    "archive_pipeline.probe_audio_duration_seconds"
                ) as duration_probe:
                    self.assertEqual(forced.process_one(source), "verified")
                    importer.assert_not_called()
                    duration_probe.assert_not_called()
                forced._transcribe.assert_not_called()
                self.assertEqual(forced_cleanup.calls, [BODY_TEXT])
                self.assertEqual(forced_cleanup.reuse_checkpoints, [False])
                self.assertEqual([item["text"] for item in forced_renders], [BODY_TEXT])
                self.assertEqual(paths["raw_text"].read_bytes(), raw_before)
            finally:
                forced.close()

    def test_source_docx_mutation_blocks_resume_and_publication(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            archive, _audio, transcript, output = self.make_archive(root)
            source = transcript.resolve()
            config = self.config(archive, output, publish=True)
            cleanup = FakeCleanupClient()
            runner = PipelineRunner(config)
            self.install_fakes(runner, cleanup)
            try:
                self.assertEqual(runner.process_one(source), "verified")
            finally:
                runner.close()

            manifest_path = artifact_paths(
                artifact_directory(source, archive, output)
            )["manifest"]
            changed_body = (
                "This deliberately changed transcript contains enough distinct words to "
                "remain a valid synthetic document while changing the immutable source hash.",
                "Publication must stop because the document no longer matches the exact "
                "container that supplied the preserved raw text for cleanup and rendering.",
            )
            write_legacy_docx(source, changed_body)

            with self.assertRaisesRegex(
                ReplacementError, "changed after import|proven prior publication"
            ):
                plan_legacy_docx_replacements(
                    output,
                    archive,
                    manifest_paths=[manifest_path],
                )

            resumed_cleanup = FakeCleanupClient()
            resumed = PipelineRunner(config)
            self.install_fakes(resumed, resumed_cleanup)
            try:
                with mock.patch(
                    "existing_transcript_import.import_existing_transcript",
                    side_effect=AssertionError("changed source must not be re-imported"),
                ) as importer:
                    self.assertEqual(resumed.process_one(source), "failed")
                    importer.assert_not_called()
                self.assertEqual(resumed_cleanup.calls, [])
            finally:
                resumed.close()

            failed_manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(failed_manifest["status"], "failed")
            self.assertIn("refusing to re-import", failed_manifest["error"])

    def test_publication_keeps_source_and_backs_up_only_prior_review_copy(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            archive, _audio, transcript, output = self.make_archive(root)
            source = transcript.resolve()
            original_bytes = source.read_bytes()
            original_hash = sha256_file(source)
            review = source.with_name(f"{source.stem} - GLM Review.docx")
            config = self.config(archive, output, publish=True)
            runner = PipelineRunner(config)
            self.install_fakes(runner, FakeCleanupClient())
            try:
                self.assertEqual(runner.process_one(source), "verified")
            finally:
                runner.close()

            paths = artifact_paths(artifact_directory(source, archive, output))
            generated_bytes = paths["docx"].read_bytes()
            counts = {
                "discovered": 1,
                "queued": 0,
                "skipped": 0,
                "verified": 1,
                "needs_review": 0,
                "failed": 0,
                "cancelled": 0,
            }
            first_time = datetime(2026, 8, 5, 12, 0, 0, 123456, tzinfo=timezone.utc)
            first = publish_source_docx_batch(
                config,
                counts,
                manifest_paths=[paths["manifest"]],
                now=first_time,
            )
            self.assertIsNotNone(first)
            self.assertEqual(first["status"], "published")
            self.assertEqual(first["approval_state"], "pending_human_review")
            self.assertEqual(first["operations"]["create"], 1)
            self.assertTrue(Path(first["backup_root"]).is_dir())
            self.assertEqual(list(Path(first["backup_root"]).rglob("*.docx")), [])
            self.assertEqual(source.read_bytes(), original_bytes)
            self.assertEqual(sha256_file(source), original_hash)
            self.assertEqual(review.read_bytes(), generated_bytes)

            changed_text = BODY_TEXT.replace("careful observation", "patient observation")
            forced_config = self.config(archive, output, force=True, publish=True)
            forced = PipelineRunner(forced_config)
            self.install_fakes(
                forced, FakeCleanupClient(result_text=changed_text)
            )
            try:
                self.assertEqual(forced.process_one(source), "verified")
            finally:
                forced.close()
            regenerated_bytes = paths["docx"].read_bytes()
            self.assertNotEqual(regenerated_bytes, generated_bytes)

            second = publish_source_docx_batch(
                forced_config,
                counts,
                manifest_paths=[paths["manifest"]],
                now=first_time + timedelta(seconds=1),
            )
            self.assertIsNotNone(second)
            self.assertEqual(second["status"], "published")
            self.assertEqual(second["operations"]["replace"], 1)
            backup = (
                Path(second["backup_root"])
                / Path(second["plan"]["items"][0]["target_relative"])
            )
            self.assertEqual(backup.read_bytes(), generated_bytes)
            self.assertEqual(review.read_bytes(), regenerated_bytes)
            self.assertEqual(source.read_bytes(), original_bytes)

            third = publish_source_docx_batch(
                forced_config,
                counts,
                manifest_paths=[paths["manifest"]],
                now=first_time + timedelta(seconds=2),
            )
            self.assertEqual(third["operations"]["noop"], 1)
            self.assertFalse(Path(third["backup_root"]).exists())

    def test_manually_changed_review_copy_fails_closed(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            archive, _audio, transcript, output = self.make_archive(root)
            source = transcript.resolve()
            source_bytes = source.read_bytes()
            config = self.config(archive, output, publish=True)
            runner = PipelineRunner(config)
            self.install_fakes(runner, FakeCleanupClient())
            try:
                self.assertEqual(runner.process_one(source), "verified")
            finally:
                runner.close()
            paths = artifact_paths(artifact_directory(source, archive, output))
            counts = {
                "discovered": 1,
                "queued": 0,
                "skipped": 0,
                "verified": 1,
                "needs_review": 0,
                "failed": 0,
                "cancelled": 0,
            }
            publish_source_docx_batch(
                config, counts, manifest_paths=[paths["manifest"]]
            )
            review = source.with_name(f"{source.stem} - GLM Review.docx")
            write_legacy_docx(
                review,
                (
                    "A human deliberately changed this review copy and those edits must not be overwritten.",
                    "The pipeline must fail closed until the conflict is resolved explicitly.",
                ),
            )

            with self.assertRaisesRegex(
                ReplacementError, "manually changed|not a proven prior publication"
            ):
                plan_legacy_docx_replacements(
                    output, archive, manifest_paths=[paths["manifest"]]
                )
            self.assertEqual(source.read_bytes(), source_bytes)

    def test_hard_interruption_after_commit_is_recoverable_from_planned_journal(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            archive, _audio, transcript, output = self.make_archive(root)
            source = transcript.resolve()
            source_bytes = source.read_bytes()
            review = source.with_name(f"{source.stem} - GLM Review.docx")
            config = self.config(archive, output, publish=True)
            runner = PipelineRunner(config)
            self.install_fakes(runner, FakeCleanupClient())
            try:
                self.assertEqual(runner.process_one(source), "verified")
            finally:
                runner.close()

            paths = artifact_paths(artifact_directory(source, archive, output))
            generated_bytes = paths["docx"].read_bytes()
            counts = {
                "discovered": 1,
                "queued": 0,
                "skipped": 0,
                "verified": 1,
                "needs_review": 0,
                "failed": 0,
                "cancelled": 0,
            }
            crash_time = datetime(2026, 8, 5, 13, 0, 0, 123456, tzinfo=timezone.utc)

            def commit_then_lose_process(
                plan,
                *,
                expected_scope_root,
                backup_root,
                confirm,
                expected_count,
            ):
                item = plan.items[0]
                if item.original_sha256 is not None:
                    backup = Path(backup_root) / item.target_relative
                    backup.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copyfile(item.target, backup)
                staged = item.target.parent / f".{item.target.name}.power-loss-test"
                shutil.copyfile(item.generated, staged)
                os.replace(staged, item.target)
                # BaseException models termination that bypasses normal rollback
                # and the publisher's Exception handler/final receipt.
                raise SystemExit("simulated power loss")

            with mock.patch(
                "legacy_docx_replace.apply_legacy_docx_replacements",
                side_effect=commit_then_lose_process,
            ), self.assertRaisesRegex(SystemExit, "simulated power loss"):
                publish_source_docx_batch(
                    config,
                    counts,
                    manifest_paths=[paths["manifest"]],
                    now=crash_time,
                )

            journal = output / "source-docx-publication-20260805-130000-123456.json"
            self.assertEqual(json.loads(journal.read_text(encoding="utf-8"))["status"], "planned")
            self.assertEqual(source.read_bytes(), source_bytes)
            self.assertEqual(review.read_bytes(), generated_bytes)

            changed_text = BODY_TEXT.replace(
                "careful observation", "patient observation", 1
            )
            resumed_config = self.config(archive, output, force=True, publish=True)
            resumed_cleanup = FakeCleanupClient(result_text=changed_text)
            resumed = PipelineRunner(resumed_config)
            self.install_fakes(resumed, resumed_cleanup)
            try:
                with mock.patch(
                    "existing_transcript_import.import_existing_transcript",
                    side_effect=AssertionError("recovery must use preserved raw text"),
                ) as importer:
                    self.assertEqual(resumed.process_one(source), "verified")
                    importer.assert_not_called()
                resumed._transcribe.assert_not_called()
            finally:
                resumed.close()

            regenerated_bytes = paths["docx"].read_bytes()
            self.assertNotEqual(regenerated_bytes, generated_bytes)
            self.assertEqual(source.read_bytes(), source_bytes)
            self.assertEqual(review.read_bytes(), generated_bytes)
            recovered = publish_source_docx_batch(
                resumed_config,
                counts,
                manifest_paths=[paths["manifest"]],
                now=crash_time + timedelta(seconds=1),
            )
            self.assertEqual(recovered["status"], "published")
            self.assertEqual(recovered["operations"]["replace"], 1)
            self.assertEqual(review.read_bytes(), regenerated_bytes)
            self.assertEqual(source.read_bytes(), source_bytes)

    def test_incomplete_rollback_journal_remains_recoverable(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            archive, _audio, transcript, output = self.make_archive(root)
            source = transcript.resolve()
            source_bytes = source.read_bytes()
            review = source.with_name(f"{source.stem} - GLM Review.docx")
            config = self.config(archive, output, publish=True)
            runner = PipelineRunner(config)
            self.install_fakes(runner, FakeCleanupClient())
            try:
                self.assertEqual(runner.process_one(source), "verified")
            finally:
                runner.close()

            paths = artifact_paths(artifact_directory(source, archive, output))
            generated_bytes = paths["docx"].read_bytes()
            counts = {
                "discovered": 1,
                "queued": 0,
                "skipped": 0,
                "verified": 1,
                "needs_review": 0,
                "failed": 0,
                "cancelled": 0,
            }
            failure_time = datetime(
                2026, 8, 5, 14, 0, 0, 123456, tzinfo=timezone.utc
            )

            def commit_then_fail_rollback(
                plan,
                *,
                expected_scope_root,
                backup_root,
                confirm,
                expected_count,
            ):
                item = plan.items[0]
                if item.original_sha256 is not None:
                    backup = Path(backup_root) / item.target_relative
                    backup.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copyfile(item.target, backup)
                staged = item.target.parent / f".{item.target.name}.rollback-test"
                shutil.copyfile(item.generated, staged)
                os.replace(staged, item.target)
                raise ReplacementError(
                    "simulated commit failure; rollback also failed"
                )

            with mock.patch(
                "legacy_docx_replace.apply_legacy_docx_replacements",
                side_effect=commit_then_fail_rollback,
            ), self.assertRaisesRegex(ReplacementError, "rollback also failed"):
                publish_source_docx_batch(
                    config,
                    counts,
                    manifest_paths=[paths["manifest"]],
                    now=failure_time,
                )

            journal = output / "source-docx-publication-20260805-140000-123456.json"
            report = json.loads(journal.read_text(encoding="utf-8"))
            self.assertEqual(report["status"], "rollback_incomplete")
            self.assertEqual(source.read_bytes(), source_bytes)
            self.assertEqual(review.read_bytes(), generated_bytes)

            resumed = PipelineRunner(config)
            self.install_fakes(resumed, FakeCleanupClient())
            try:
                with mock.patch(
                    "existing_transcript_import.import_existing_transcript",
                    side_effect=AssertionError("recovery must use preserved raw text"),
                ) as importer:
                    self.assertEqual(resumed.process_one(source), "skipped")
                    importer.assert_not_called()
            finally:
                resumed.close()

            recovered = publish_source_docx_batch(
                config,
                counts,
                manifest_paths=[paths["manifest"]],
                now=failure_time + timedelta(seconds=1),
            )
            self.assertEqual(recovered["status"], "published")
            self.assertEqual(recovered["operations"]["noop"], 1)


if __name__ == "__main__":
    unittest.main()
