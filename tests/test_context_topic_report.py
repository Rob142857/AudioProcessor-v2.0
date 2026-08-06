from __future__ import annotations

import hashlib
import json
from pathlib import Path
import tempfile
from types import MappingProxyType
import unittest

from docx import Document
from docx.oxml.ns import qn

from context_compilation_inventory import (
    BoundContextCompilation,
    BoundContextRegion,
    BoundContextSource,
    BoundSelectedParagraph,
)
from context_topic_report import (
    REPORT_MARKER,
    ContextTopicAnalysisError,
    ContextTopicReportOutputError,
    create_subtopic_plan_report,
)


def _sha(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _sha_json(value: object) -> str:
    encoded = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _region(
    ordinal: int,
    region_id: str,
    source: Path,
    text: str,
) -> BoundContextRegion:
    selected = BoundSelectedParagraph(
        number=ordinal,
        text=text,
        page_number=ordinal,
        emitted=True,
    )
    relative = f"1985 MW/{source.name}"
    text_sha = _sha(text)
    return BoundContextRegion(
        region_id=region_id,
        ordinal=ordinal,
        source_region_ordinal=1,
        source_relative_path=relative,
        source_absolute_path=str(source),
        source_sha256=_sha(relative),
        source_target=source.as_uri(),
        heading_text=f"{source.stem} | Context 1",
        metadata_text=f"Section {ordinal} | Page {ordinal} | Paragraph {ordinal}",
        locator=f"Page {ordinal} | Paragraph {ordinal}",
        selected_paragraphs=(selected,),
        selected_text=text,
        selected_text_sha256=text_sha,
        selected_paragraphs_sha256=_sha(f"paragraphs:{text}"),
        emitted_paragraph_count=1,
        empty_selected_paragraph_count=0,
        occurrence_count=1,
        occurrence_ids=(f"occ_{ordinal}",),
        selection_start_paragraph=ordinal,
        selection_end_paragraph=ordinal,
        selection_method="local_test",
        selection_model=None,
        selection_confidence=1.0,
        selection_note=None,
        query_count=1,
        highlight_count=1,
        region_fingerprint=_sha(f"region:{region_id}"),
    )


def _inventory(root: Path) -> tuple[BoundContextCompilation, Path]:
    snapshot = root / "temporary read snapshot.docx"
    records = root / "temporary read snapshot.jsonl"
    real_master = root / "Awakening Complete Context.docx"
    for path, data in (
        (snapshot, b"snapshot-master"),
        (records, b"canonical-records"),
        (real_master, b"real-master"),
    ):
        path.write_bytes(data)
    source_paths = [root / f"Lecture {letter}.docx" for letter in "ABC"]
    for path in source_paths:
        path.write_bytes(path.name.encode("utf-8"))
    duplicate = "Awakening must occur in ordinary life."
    regions = (
        _region(1, "region_1", source_paths[0], duplicate),
        _region(2, "region_2", source_paths[1], duplicate),
        _region(3, "region_3", source_paths[2], "Practice prepares the whole being for awakening."),
    )
    sources = tuple(
        BoundContextSource(
            source_relative_path=region.source_relative_path,
            source_absolute_path=region.source_absolute_path,
            source_sha256=region.source_sha256,
            source_target=region.source_target,
            region_count=1,
            occurrence_count=1,
        )
        for region in regions
    )
    inventory = BoundContextCompilation(
        schema_version="context-compilation-inventory-v1",
        docx_path=snapshot,
        jsonl_path=records,
        docx_sha256=_sha("snapshot-master"),
        jsonl_sha256=_sha("canonical-records"),
        pair_fingerprint=_sha("pair"),
        ordered_regions_sha256=_sha("ordered"),
        source_manifest_sha256=_sha("sources"),
        query="awakening",
        query_canonical="awakening",
        query_word_count=1,
        root=str(root),
        scanned_files=3,
        ignored_generated_files=0,
        region_count=3,
        occurrence_count=3,
        source_count=3,
        selected_paragraph_count=3,
        emitted_paragraph_count=3,
        empty_selected_paragraph_count=0,
        highlight_count=3,
        sources=sources,
        regions=regions,
    )
    return inventory, real_master


def _analysis_payload(inventory: BoundContextCompilation, master: Path) -> dict:
    payload = {
        "analysis_profile": "context-topic-analysis-v1",
        "model": "@cf/zai-org/glm-4.7-flash",
        "operation": "classification",
        "query": "awakening",
        "status": "proposed",
        "recommendation": "Retain the master and review two smaller proposed reading volumes.",
        "corpus": {
            "inventory_pair_fingerprint": inventory.pair_fingerprint,
            "master_docx_path": str(master),
            "master_layout_pages": 37,
            "reading_words_per_page": 5,
        },
        "integrity": {
            "inventory_pair_fingerprint": inventory.pair_fingerprint,
            "request_sha256": "1" * 64,
            "model_response_sha256": "2" * 64,
        },
        "taxonomy_sha256": "",
        "taxonomy": {
            "families": [
                {
                    "family_id": "family_practice",
                    "label": "Practice and Preparation",
                    "definition": "Practical disciplines and conditions of preparation.",
                }
            ],
            "topics": [
                {
                    "topic_id": "topic_imperative",
                    "family_id": "family_practice",
                    "label": "The Imperative to Awaken",
                    "definition": "Direct statements that awakening is necessary.",
                    "include_cues": ["must occur", "necessity"],
                    "exclude_cues": ["technical exercises"],
                },
                {
                    "topic_id": "topic_preparation",
                    "family_id": "family_practice",
                    "label": "Preparation and Practice",
                    "definition": "Methods that prepare the whole being.",
                    "include_cues": ["practice", "prepares"],
                    "exclude_cues": ["abstract definition"],
                },
            ],
        },
        "classifications": [
            {
                "region_id": "region_1",
                "region_input_sha256": "4" * 64,
                "status": "classified",
                "primary_topic_id": "topic_imperative",
                "secondary_topic_ids": [],
                "evidence_paragraph_numbers": [1],
                "certainty": "high",
                "ambiguity": "none",
                "review_status": "review_recommended",
            },
            {
                "region_id": "region_2",
                "region_input_sha256": "5" * 64,
                "status": "classified",
                "primary_topic_id": "topic_imperative",
                "secondary_topic_ids": [],
                "evidence_paragraph_numbers": [2],
                "certainty": "low",
                "ambiguity": "insufficient_context",
                "review_status": "review_required",
            },
            {
                "region_id": "region_3",
                "region_input_sha256": "6" * 64,
                "status": "classified",
                "primary_topic_id": "topic_preparation",
                "secondary_topic_ids": ["topic_imperative"],
                "evidence_paragraph_numbers": [3],
                "certainty": "medium",
                "ambiguity": "topic_overlap",
                "review_status": "review_recommended",
            },
        ],
    }
    payload["taxonomy_sha256"] = _sha_json(payload["taxonomy"])
    return payload


def _write_analysis(path: Path, payload: dict) -> None:
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


class ContextTopicReportTests(unittest.TestCase):
    def test_creates_polished_proposal_from_local_join_without_quotes_or_input_changes(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            inventory, master = _inventory(root)
            analysis = root / "awakening-topic-analysis.json"
            _write_analysis(analysis, _analysis_payload(inventory, master))
            protected = {
                path: hashlib.sha256(path.read_bytes()).hexdigest()
                for path in (inventory.docx_path, inventory.jsonl_path, master, analysis)
            }

            output = create_subtopic_plan_report(analysis, inventory)

            self.assertTrue(output.parent.samefile(master.parent))
            self.assertEqual(f"{master.stem} - Subtopic Plan.docx", output.name)
            self.assertTrue(output.is_file())
            for path, digest in protected.items():
                self.assertEqual(digest, hashlib.sha256(path.read_bytes()).hexdigest())
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = "\n".join(
                cell.text for table in document.tables for row in table.rows for cell in row.cells
            )
            combined = text + "\n" + table_text
            self.assertIn("PROPOSED - NOT APPROVED", combined)
            self.assertIn("37-page exact-source master compilation", combined)
            self.assertIn("Corpus integrity ledger", text)
            self.assertIn("The Imperative to Awaken", combined)
            self.assertIn("Lecture A | Context 1", combined)
            self.assertIn("1985 MW/Lecture A.docx", combined)
            self.assertIn("2 primary regions; 1 secondary membership", combined)
            self.assertIn("1 unique passage", combined)
            self.assertIn("consolidated reading pages", combined)
            self.assertNotIn("Awakening must occur in ordinary life.", combined)
            self.assertEqual(REPORT_MARKER, document.core_properties.subject)
            self.assertEqual(REPORT_MARKER, document.core_properties.keywords)
            self.assertGreaterEqual(len(document.tables), 7)

    def test_applies_exact_compact_reference_guide_geometry(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            inventory, master = _inventory(root)
            analysis = root / "analysis.json"
            _write_analysis(analysis, _analysis_payload(inventory, master))
            output = create_subtopic_plan_report(analysis, inventory, root / "plan.docx")
            document = Document(output)

            section = document.sections[0]
            self.assertAlmostEqual(8.5, section.page_width.inches, places=3)
            self.assertAlmostEqual(11, section.page_height.inches, places=3)
            for margin in (
                section.top_margin,
                section.right_margin,
                section.bottom_margin,
                section.left_margin,
            ):
                self.assertAlmostEqual(1, margin.inches, places=3)
            self.assertAlmostEqual(0.492, section.header_distance.inches, places=3)
            self.assertAlmostEqual(0.492, section.footer_distance.inches, places=3)
            normal = document.styles["Normal"]
            self.assertEqual("Calibri", normal.font.name)
            self.assertAlmostEqual(11, normal.font.size.pt, places=2)
            self.assertAlmostEqual(6, normal.paragraph_format.space_after.pt, places=2)
            self.assertEqual(1.25, normal.paragraph_format.line_spacing)
            for name, size, before, after, colour in (
                ("Heading 1", 16, 18, 10, "2E74B5"),
                ("Heading 2", 13, 14, 7, "2E74B5"),
                ("Heading 3", 12, 10, 5, "1F4D78"),
            ):
                style = document.styles[name]
                self.assertAlmostEqual(size, style.font.size.pt, places=2)
                self.assertAlmostEqual(before, style.paragraph_format.space_before.pt, places=2)
                self.assertAlmostEqual(after, style.paragraph_format.space_after.pt, places=2)
                self.assertEqual(colour, str(style.font.color.rgb))

            for table in document.tables:
                properties = table._tbl.tblPr
                self.assertEqual("9360", properties.find(qn("w:tblW")).get(qn("w:w")))
                self.assertEqual("dxa", properties.find(qn("w:tblW")).get(qn("w:type")))
                self.assertEqual("120", properties.find(qn("w:tblInd")).get(qn("w:w")))
                self.assertEqual("fixed", properties.find(qn("w:tblLayout")).get(qn("w:type")))
                widths = [int(item.get(qn("w:w"))) for item in table._tbl.tblGrid]
                self.assertEqual(9360, sum(widths))
                for row in table.rows:
                    self.assertIsNone(row.height)
                    for cell, width in zip(row.cells, widths):
                        self.assertEqual(str(width), cell._tc.tcPr.tcW.get(qn("w:w")))

    def test_rejects_different_inventory_or_incomplete_classification(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            inventory, master = _inventory(root)
            analysis = root / "analysis.json"
            payload = _analysis_payload(inventory, master)
            payload["corpus"]["inventory_pair_fingerprint"] = "f" * 64
            _write_analysis(analysis, payload)
            with self.assertRaisesRegex(ContextTopicAnalysisError, "different compilation"):
                create_subtopic_plan_report(analysis, inventory)

            payload = _analysis_payload(inventory, master)
            payload["classifications"].pop()
            _write_analysis(analysis, payload)
            with self.assertRaisesRegex(ContextTopicAnalysisError, "coverage differs"):
                create_subtopic_plan_report(analysis, inventory)

    def test_accepts_final_orchestrator_shape_and_review_ambiguity_aliases(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            inventory, master = _inventory(root)
            analysis = root / "final-analysis.json"
            payload = _analysis_payload(inventory, master)
            payload["topic_analysis_profile"] = payload.pop("analysis_profile")
            payload["status"] = "complete"
            payload["corpus"]["query"] = payload.pop("query")
            payload["corpus"]["source_records_sha256"] = inventory.jsonl_sha256
            payload["corpus"]["master_docx_sha256"] = inventory.docx_sha256
            payload.pop("integrity")
            rows = payload.pop("classifications")
            review_values = ("accepted", "adjudicate", "human_review")
            ambiguity_values = (
                ["mixed_passage", "taxonomy_overlap"],
                ["insufficient_context"],
                ["boundary_uncertain"],
            )
            for row, review, ambiguity, region in zip(
                rows,
                review_values,
                ambiguity_values,
                inventory.regions,
            ):
                row.pop("status")
                row.pop("review_status")
                row.pop("ambiguity")
                row["model_review_status"] = review
                row["review_required"] = False
                row["ambiguity_codes"] = ambiguity
                row["taxonomy_gap"] = False
                row["selected_text_sha256"] = region.selected_text_sha256
            payload["regions"] = rows
            _write_analysis(analysis, payload)

            output = create_subtopic_plan_report(
                analysis,
                inventory,
                root / "orchestrator-plan.docx",
            )

            document = Document(output)
            combined = "\n".join(
                [paragraph.text for paragraph in document.paragraphs]
                + [
                    cell.text
                    for table in document.tables
                    for row in table.rows
                    for cell in row.cells
                ]
            )
            self.assertIn("accepted 1", combined)
            self.assertIn("review required 1", combined)
            self.assertIn("mixed passage", combined)
            self.assertIn("taxonomy overlap", combined)
            self.assertIn("boundary uncertain", combined)

    def test_never_allows_report_to_replace_master_or_analysis(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            inventory, master = _inventory(root)
            analysis = root / "analysis.json"
            _write_analysis(analysis, _analysis_payload(inventory, master))
            with self.assertRaises(ContextTopicReportOutputError):
                create_subtopic_plan_report(analysis, inventory, master)
            with self.assertRaises(ContextTopicReportOutputError):
                create_subtopic_plan_report(analysis, inventory, analysis)


if __name__ == "__main__":
    unittest.main()
