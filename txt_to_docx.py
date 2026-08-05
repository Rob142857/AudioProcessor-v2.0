import argparse
import os
import re
import tempfile
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.shared import RGBColor  # type: ignore

from console_compat import configure_safe_stdio
from publication_metadata import (
    PublicationMetadata,
    format_publication_date,
    infer_publication_metadata,
)


configure_safe_stdio()

try:
    from australian_spelling import normalize_text
    AUSTRALIAN_SPELLING_AVAILABLE = True
except ImportError:
    AUSTRALIAN_SPELLING_AVAILABLE = False
    def normalize_text(text, **kwargs): return text


EMU_PER_INCH = 914_400
EMU_PER_POINT = 12_700
TITLE_COLOUR = RGBColor(32, 55, 72)       # #203748
MUTED_COLOUR = RGBColor(96, 105, 112)     # restrained publication metadata
BODY_COLOUR = RGBColor(28, 31, 34)


def _set_font(run, *, size: float, color=None, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run.font.size = int(size * EMU_PER_POINT)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.font.italic = italic


def _set_paragraph_rhythm(
    paragraph,
    *,
    before: float = 0,
    after: float = 8,
    line_spacing: float = 1.333,
) -> None:
    paragraph.paragraph_format.space_before = int(before * EMU_PER_POINT)
    paragraph.paragraph_format.space_after = int(after * EMU_PER_POINT)
    paragraph.paragraph_format.line_spacing = line_spacing


def _configure_publication_document(doc: Document) -> None:
    """Apply the narrative_proposal preset and compact editorial override."""

    for section in getattr(doc, "sections", ()):
        section.page_width = int(8.5 * EMU_PER_INCH)
        section.page_height = int(11 * EMU_PER_INCH)
        section.top_margin = EMU_PER_INCH
        section.right_margin = EMU_PER_INCH
        section.bottom_margin = EMU_PER_INCH
        section.left_margin = EMU_PER_INCH
        section.header_distance = int(0.492 * EMU_PER_INCH)
        section.footer_distance = int(0.492 * EMU_PER_INCH)

    styles = getattr(doc, "styles", None)
    if styles is None:
        return
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = 11 * EMU_PER_POINT
    normal.font.color.rgb = BODY_COLOUR
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = 0
    normal.paragraph_format.space_after = 8 * EMU_PER_POINT
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, RGBColor(46, 116, 181), 18, 10),
        "Heading 2": (13, RGBColor(46, 116, 181), 12, 6),
        "Heading 3": (12, RGBColor(31, 77, 120), 8, 4),
    }
    for name, (size, colour, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = size * EMU_PER_POINT
        style.font.color.rgb = colour
        style.paragraph_format.space_before = before * EMU_PER_POINT
        style.paragraph_format.space_after = after * EMU_PER_POINT


_WORKFLOW_LINE_RE = re.compile(
    r"^(?:"
    r"(?:cleaned\s+up\s+at|processed\s+at|generated\s+at)(?:\s*:|\s+)|"
    r"processed\s+by\s+speech[- ]to[- ]text\s+from\s+a\s+digitised\s+tape\s+recording\s+"
    r"originally\s+(?:taken\s+from|recorded\s+in\s+person\s+by)\b|"
    r"needs\s+human\s+review\.?$|"
    r"(?:model|device|processing\s+time|audio\s+preprocessing|"
    r"pipeline(?:\s+version)?|cleanup\s+model|transcription\s+information)\s*:"
    r")",
    re.IGNORECASE,
)


def strip_workflow_metadata(body: str) -> str:
    """Remove old trailing processing/provenance notes without editing lecture prose."""

    lines = body.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    marker_index = None
    for index, line in enumerate(lines):
        stripped = line.strip()
        if stripped.casefold() == "transcription information":
            marker_index = index
    if marker_index is not None and marker_index >= max(0, len(lines) - 20):
        lines = lines[:marker_index]

    while lines:
        stripped = lines[-1].strip()
        if not stripped:
            lines.pop()
            continue
        if (
            _WORKFLOW_LINE_RE.match(stripped)
            or stripped.casefold() == "(this information can be deleted if not needed)"
            or re.fullmatch(r"[_=-]{8,}", stripped)
        ):
            lines.pop()
            continue
        break
    return "\n".join(lines).strip()


def _save_docx_atomically(doc: Document, out_path: Path) -> Path:
    """Save and validate a DOCX before atomically replacing its destination."""
    out_path = Path(out_path)
    if out_path.suffix.lower() != ".docx":
        raise ValueError(f"DOCX output path must end in .docx: {out_path}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path: Optional[Path] = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="wb",
            prefix=f".{out_path.stem}.",
            suffix=".tmp.docx",
            dir=out_path.parent,
            delete=False,
        ) as tmp:
            tmp_path = Path(tmp.name)

        doc.save(str(tmp_path))
        # python-docx reopening the package catches truncated/corrupt writes before
        # the destination is replaced.
        Document(str(tmp_path))
        os.replace(tmp_path, out_path)
        tmp_path = None
        return out_path
    finally:
        if tmp_path is not None:
            try:
                tmp_path.unlink()
            except FileNotFoundError:
                pass


def infer_year_from_parent(folder_name: str) -> int:
    """Extract a year from a folder name, e.g. '1988 MW' -> 1988 or '84-97' -> 1984.

    First tries to find a 4-digit year (19xx or 20xx).
    If not found, looks for 2-digit year patterns and assumes 19xx for years >= 50, 20xx for < 50.
    Raises ValueError if no year pattern is found.
    """
    # Try 4-digit year first
    m = re.search(r"(19\d{2}|20\d{2})", folder_name)
    if m:
        return int(m.group(1))
    
    # Try 2-digit year patterns like "84-97", "Recordings 92", etc.
    m = re.search(r"\b(\d{2})(?:-\d{2})?\b", folder_name)
    if m:
        two_digit = int(m.group(1))
        # Assume 19xx for years >= 50, 20xx for years < 50
        return 1900 + two_digit if two_digit >= 50 else 2000 + two_digit
    
    raise ValueError(f"Could not find year in folder name: {folder_name!r}")


def infer_year_from_ancestors(start: Path) -> Optional[int]:
    """Walk up from a path until we find a folder name containing a year.

    This allows layouts like '.../1988 MW/Temp/0202 Fishes.txt' and also handles
    2-digit year patterns like '84-97'. Returns None if no suitable folder is found.
    """
    for folder in [start] + list(start.parents):
        try:
            return infer_year_from_parent(folder.name)
        except ValueError:
            continue
    return None


def infer_date_from_filename(filename: str, year: Optional[int]) -> Optional[date]:
    """Infer a calendar date from a filename like '0202 Fishes.txt'.

    Assumes the first four digits are MMDD for the given year.
    Returns None if year is None or if no MMDD pattern is found.
    """
    if year is None:
        return None
    stem = Path(filename).stem
    m = re.match(r"(\d{4})(?:\D|$)", stem)
    if not m:
        return None
    mmdd = m.group(1)
    month = int(mmdd[:2])
    day = int(mmdd[2:])
    try:
        return date(year, month, day)
    except ValueError:
        return None


def make_title_from_filename(filename: str) -> str:
    """Use the words in the filename (after the leading MMDD) as the title.

    Example: '0202 Fishes.txt' -> 'Fishes'
    """
    stem = Path(filename).stem
    # Drop leading MMDD and optional separator
    m = re.match(r"\d{4}[_\- ]*(.*)", stem)
    title_part = m.group(1) if m and m.group(1) else stem
    return title_part.strip() or stem


def extract_lecture_number(filename: str) -> Optional[int]:
    """Extract lecture number from filename if present.
    
    Examples:
        '0114 1992 Mythology.mp3' -> 114 (from MMDD)
        'Lecture 05.mp3' -> 5
        'L23 Topic.mp3' -> 23
    
    Returns None if no clear lecture number pattern is found.
    """
    stem = Path(filename).stem
    
    # Try explicit lecture number patterns first
    m = re.search(r"(?:Lecture|L)\s*(\d+)", stem, re.IGNORECASE)
    if m:
        return int(m.group(1))
    
    # For MMDD format files, use the day portion as lecture number if it seems reasonable
    # (e.g., 0114 could be lecture 14)
    m = re.match(r"(\d{2})(\d{2})", stem)
    if m:
        month = int(m.group(1))
        day = int(m.group(2))
        # Only use if month is valid (1-12) and day looks reasonable (1-31)
        if 1 <= month <= 12 and 1 <= day <= 31:
            return day
    
    return None


def get_source_path_from_header(txt_path: Path) -> Optional[Path]:
    """Read the first 'Source:' line in the txt file and return its path, if any.

    The header is expected to look like:
        Source: C:\\path\\to\\audio.mp3
    """
    text = txt_path.read_text(encoding="utf-8")
    for line in text.splitlines():
        if line.startswith("Source:"):
            raw = line[len("Source:") :].strip()
            if raw:
                return Path(raw)
            break
    return None


def load_body_text(txt_path: Path) -> str:
    """Load transcript text, stripping the Source/Output header lines if present."""
    text = txt_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    body_lines = []
    for line in lines:
        if line.startswith("Source:") or line.startswith("Output:"):
            continue
        body_lines.append(line)
    body = "\n".join(body_lines).lstrip("\n")
    return body


def add_paragraphs_from_text(doc: Document, body: str) -> None:
    """Add justified paragraphs, preserving only editorial blank-line breaks.

    Paragraph boundaries are decided by the cleanup editor from changes in
    meaning and rhetorical function.  The renderer must not manufacture
    regular-length paragraphs merely to make the page look balanced.
    """
    # Split on blank lines (one or more empty/whitespace-only lines)
    blocks = re.split(r"\n\s*\n", body)
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        para = doc.add_paragraph(block)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_paragraph_rhythm(para)


def _add_publication_opening(doc: Document, publication: PublicationMetadata) -> None:
    """Compact editorial_cover override: polished opening, body still on page one."""

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_rhythm(title_paragraph, after=6, line_spacing=1.0)
    title_paragraph.paragraph_format.keep_with_next = True
    title_run = title_paragraph.add_run(publication.title)
    _set_font(title_run, size=30, color=TITLE_COLOUR, bold=False)

    metadata_values = [publication.artist]
    if publication.lecture_date:
        metadata_values.append(format_publication_date(publication.lecture_date))
    if publication.source_type:
        metadata_values.append(publication.source_type)
    metadata_paragraph = doc.add_paragraph()
    metadata_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_rhythm(metadata_paragraph, after=14, line_spacing=1.0)
    metadata_paragraph.paragraph_format.keep_with_next = True
    metadata_run = metadata_paragraph.add_run(" | ".join(metadata_values))
    _set_font(metadata_run, size=10, color=MUTED_COLOUR)


def _add_provenance_postscript(
    doc: Document,
    publication: PublicationMetadata,
    *,
    needs_human_review: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_rhythm(paragraph, before=12, after=0, line_spacing=1.0)
    run = paragraph.add_run(publication.postscript)
    _set_font(run, size=9, color=MUTED_COLOUR, italic=True)

    if needs_human_review:
        review_paragraph = doc.add_paragraph()
        review_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_rhythm(review_paragraph, before=3, after=0, line_spacing=1.0)
        review_run = review_paragraph.add_run("Needs human review.")
        _set_font(review_run, size=9, color=MUTED_COLOUR, italic=True, bold=True)


def _set_core_properties(doc: Document, publication: PublicationMetadata) -> None:
    properties = getattr(doc, "core_properties", None)
    if properties is None:
        return
    properties.title = publication.title
    properties.author = publication.artist
    properties.subject = (
        f"{publication.source_type} lecture transcript"
        if publication.source_type
        else "Lecture transcript"
    )


def convert_txt_to_docx(txt_path: Path, year: Optional[int] = None) -> Path:
    if not txt_path.is_file():
        raise FileNotFoundError(f"Input file not found: {txt_path}")

    # Infer year from the directory structure if not provided explicitly.
    # Preferred source is the path embedded in the header (first 'Source:'
    # line), which points to the original audio file location, e.g.
    #   Source: .../1988 MW/0202 Fishes.mp3
    # We walk up from that path to find a 4-digit year. If the header is
    # missing or unusable, we fall back to walking up from the txt location.
    if year is None:
        source_path = get_source_path_from_header(txt_path)
        if source_path is not None:
            year = infer_year_from_ancestors(source_path.parent)
        else:
            year = infer_year_from_ancestors(txt_path.parent)

    body = load_body_text(txt_path)
    source_path = get_source_path_from_header(txt_path) or txt_path
    return convert_txt_to_docx_from_text(
        body,
        source_path,
        year=year,
        use_australian_spelling=False,
        output_path=txt_path.with_suffix(".docx"),
    )


def convert_txt_to_docx_from_text(
    body_text: str,
    source_audio_path: Path,
    year: Optional[int] = None,
    metadata: Optional[dict] = None,
    use_australian_spelling: bool = True,
    *,
    output_path: Optional[Path] = None,
    relative_source_path: Optional[Path] = None,
    publication_metadata: Optional[PublicationMetadata] = None,
    needs_human_review: bool = False,
) -> Path:
    """Convert transcript text directly to an atomically written DOCX.
    
    Args:
        body_text: The formatted transcript text to include in the document
        source_audio_path: Path to the original audio/video file, used for title/date inference
        year: Optional year override for date inference
        metadata: Deprecated workflow metadata, retained for call compatibility and not published
        use_australian_spelling: Whether to convert to Australian spelling (default: True)
        output_path: Optional explicit DOCX destination (defaults to next to the source)
        relative_source_path: Optional archive-relative path used for metadata inference
        publication_metadata: Optional explicit publication fields, bypassing inference
        needs_human_review: Append a removable review notice to GLM-produced documents
    
    Returns:
        Path to the created DOCX file
    """
    # Diagnostic: Check input text length
    input_char_count = len(body_text)
    input_word_count = len(body_text.split())
    print(f"📊 Input text: {input_char_count} characters, {input_word_count} words")
    
    # Apply Australian spelling conversion and number formatting
    if use_australian_spelling and AUSTRALIAN_SPELLING_AVAILABLE:
        body_text = normalize_text(body_text, use_australian_spelling=True, fix_numbers=True)
        print("✅ Applied Australian spelling and number formatting")
        # Diagnostic: Check text length after normalization
        norm_char_count = len(body_text)
        norm_word_count = len(body_text.split())
        print(f"📊 After normalization: {norm_char_count} characters, {norm_word_count} words")
    
    # The renderer intentionally ignores model/device/timing metadata.  Those
    # details remain in the pipeline manifest, not in the publication artifact.
    _ = metadata
    body_text = strip_workflow_metadata(body_text)
    publication = publication_metadata or infer_publication_metadata(
        source_audio_path,
        relative_source_path,
        year=year,
    )

    doc = Document()
    _configure_publication_document(doc)
    _set_core_properties(doc, publication)
    _add_publication_opening(doc, publication)
    add_paragraphs_from_text(doc, body_text)
    _add_provenance_postscript(
        doc,
        publication,
        needs_human_review=needs_human_review,
    )

    # Explicit output destinations let callers preserve a source-relative tree in
    # a separate output root. The default remains next to the source audio.
    out_path = Path(output_path) if output_path is not None else source_audio_path.with_suffix(".docx")
    return _save_docx_atomically(doc, out_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert transcript .txt file(s) to formatted .docx files ready to print. Can process a single file or all .txt files in a folder.")
    parser.add_argument("input", help="Path to a transcript .txt file or a folder containing .txt files")
    parser.add_argument("--year", type=int, help="Override inferred year (e.g. 1988)")
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    
    # Determine if input is a file or folder
    if input_path.is_file():
        # Single file
        txt_files = [input_path]
    elif input_path.is_dir():
        # Folder: find all .txt files
        txt_files = sorted(input_path.glob("*.txt"))
        if not txt_files:
            print(f"No .txt files found in {input_path}")
            return
    else:
        raise FileNotFoundError(f"Input path does not exist: {input_path}")
    
    # Convert each file
    for txt_path in txt_files:
        try:
            out_path = convert_txt_to_docx(txt_path, year=args.year)
            print(f"Created DOCX: {out_path}")
        except Exception as e:
            print(f"Error processing {txt_path.name}: {e}")


if __name__ == "__main__":
    main()
