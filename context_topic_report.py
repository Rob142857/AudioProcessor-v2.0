"""Build a proposal-only Word report for a Context Finder topic analysis.

The report is a view over two immutable inputs: a locally bound Context Finder
inventory and a completed topic-analysis JSON snapshot.  It never edits the
master compilation, its JSONL companion, or the analysis record.  Only a
separate `` - Subtopic Plan.docx`` sibling is atomically published.
"""

from __future__ import annotations

from collections import Counter, defaultdict
from dataclasses import dataclass
import hashlib
from io import BytesIO
import json
import math
import os
from pathlib import Path
import re
import tempfile
from typing import Any, Mapping, Sequence

from docx import Document  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Inches, Pt, RGBColor  # type: ignore

from context_compilation_inventory import BoundContextCompilation


REPORT_SCHEMA_VERSION = "context-subtopic-plan-report-v1"
REPORT_MARKER = "AudioProcessor Context Finder Subtopic Plan"
REPORT_STATUS = "PROPOSED - NOT APPROVED"
DEFAULT_REPRESENTATIVE_SOURCE_LIMIT = 5
_WORD_RE = re.compile(r"\S+")
_SHA256_RE = re.compile(r"[0-9a-f]{64}")
_UNCLASSIFIED_TOPIC_IDS = frozenset(
    {"unclassified", "unclassified_needs_review", "taxonomy_gap"}
)
_AMBIGUITY_ALIASES = {
    "none": None,
    "topic_overlap": "taxonomy_overlap",
    "taxonomy_overlap": "taxonomy_overlap",
    "mixed_passage": "mixed_passage",
    "insufficient_context": "insufficient_context",
    "boundary_uncertain": "boundary_uncertain",
    "taxonomy_gap": "taxonomy_gap",
}
_REVIEW_ALIASES = {
    "accepted": "accepted",
    "review_recommended": "review_recommended",
    "recommended": "review_recommended",
    "review_required": "review_required",
    "required": "review_required",
    "adjudicate": "review_required",
    "human_review": "review_required",
}


class ContextTopicReportError(RuntimeError):
    """Base class for a report that cannot be safely produced."""


class ContextTopicAnalysisError(ContextTopicReportError):
    """The topic-analysis snapshot is malformed or does not bind locally."""


class ContextTopicReportOutputError(ContextTopicReportError):
    """The requested report destination is unsafe or could not be published."""


@dataclass(frozen=True, slots=True)
class _Classification:
    region_id: str
    primary_topic_id: str | None
    secondary_topic_ids: tuple[str, ...]
    certainty: str
    ambiguity_codes: tuple[str, ...]
    review_status: str
    taxonomy_gap: bool
    evidence_paragraph_numbers: tuple[int, ...]
    selected_text_sha256: str | None


@dataclass(frozen=True, slots=True)
class _Topic:
    topic_id: str
    family_id: str
    label: str
    definition: str
    include_cues: tuple[str, ...]
    exclude_cues: tuple[str, ...]


@dataclass(frozen=True, slots=True)
class _Family:
    family_id: str
    label: str
    definition: str


@dataclass(frozen=True, slots=True)
class _Analysis:
    path: Path
    sha256: str
    profile: str
    model: str
    query: str
    status: str
    recommendation: str
    master_docx_path: Path
    master_layout_pages: int
    reading_words_per_page: int
    inventory_pair_fingerprint: str
    taxonomy_sha256: str
    families: tuple[_Family, ...]
    topics: tuple[_Topic, ...]
    classifications: tuple[_Classification, ...]
    overlap_notes: tuple[Mapping[str, Any], ...]
    boundary_notes: tuple[Mapping[str, Any], ...]


@dataclass(frozen=True, slots=True)
class _TopicMetrics:
    topic: _Topic
    primary_region_ids: tuple[str, ...]
    secondary_region_ids: tuple[str, ...]
    unique_primary_passages: int
    primary_body_words: int
    unique_primary_body_words: int
    approximate_master_pages: float
    consolidated_reading_pages: int
    certainty_counts: Mapping[str, int]
    review_counts: Mapping[str, int]
    ambiguity_counts: Mapping[str, int]


def create_subtopic_plan_report(
    analysis_json_path: Path | str,
    inventory: BoundContextCompilation,
    output_path: Path | str | None = None,
    *,
    representative_source_limit: int = DEFAULT_REPRESENTATIVE_SOURCE_LIMIT,
) -> Path:
    """Create an atomically written proposal report without changing inputs.

    ``inventory`` must be the local immutable binding used to construct the
    analysis.  The JSON's recorded real master path controls the default output
    location, which remains correct when ``inventory`` was bound from a
    temporary read snapshot.
    """

    if representative_source_limit < 1:
        raise ValueError("representative_source_limit must be at least 1")
    analysis = _load_analysis(analysis_json_path, inventory)
    output = _resolve_output(analysis, inventory, output_path)
    protected = _protected_signatures(analysis, inventory)

    region_by_id = {region.region_id: region for region in inventory.regions}
    metrics = _calculate_topic_metrics(analysis, inventory, region_by_id)
    document = _build_document(
        analysis,
        inventory,
        metrics,
        region_by_id,
        representative_source_limit=representative_source_limit,
    )
    _assert_protected_unchanged(protected)
    result = _save_docx_atomically(document, output)
    _assert_protected_unchanged(protected)
    return result


def _load_analysis(
    path: Path | str,
    inventory: BoundContextCompilation,
) -> _Analysis:
    candidate = Path(path).expanduser().resolve()
    if candidate.suffix.casefold() != ".json" or not candidate.is_file():
        raise ContextTopicAnalysisError(
            f"Topic analysis must be an existing .json file: {candidate}"
        )
    before = candidate.stat()
    data = candidate.read_bytes()
    after = candidate.stat()
    if before.st_size != after.st_size or before.st_mtime_ns != after.st_mtime_ns:
        raise ContextTopicAnalysisError("Topic analysis changed while it was read")
    try:
        payload = json.loads(data.decode("utf-8-sig"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ContextTopicAnalysisError("Topic analysis is not valid UTF-8 JSON") from exc
    if not isinstance(payload, Mapping):
        raise ContextTopicAnalysisError("Topic analysis must contain a JSON object")

    corpus = _mapping(payload.get("corpus"), "corpus")
    integrity_value = payload.get("integrity")
    integrity = integrity_value if isinstance(integrity_value, Mapping) else {}
    taxonomy = _mapping(payload.get("taxonomy"), "taxonomy")
    fingerprint = _required_string(
        corpus.get("inventory_pair_fingerprint")
        or integrity.get("inventory_pair_fingerprint"),
        "corpus.inventory_pair_fingerprint",
    )
    if fingerprint != inventory.pair_fingerprint:
        raise ContextTopicAnalysisError(
            "Topic analysis was produced from a different compilation inventory"
        )
    source_records_sha256 = corpus.get("source_records_sha256")
    if (
        source_records_sha256 is not None
        and _required_sha256(source_records_sha256, "corpus.source_records_sha256")
        != inventory.jsonl_sha256
    ):
        raise ContextTopicAnalysisError(
            "Topic analysis source-record hash does not match the bound inventory"
        )
    master_snapshot_sha256 = corpus.get("master_docx_sha256")
    if (
        master_snapshot_sha256 is not None
        and _required_sha256(master_snapshot_sha256, "corpus.master_docx_sha256")
        != inventory.docx_sha256
    ):
        raise ContextTopicAnalysisError(
            "Topic analysis master hash does not match the bound inventory"
        )
    query = _required_string(payload.get("query") or corpus.get("query"), "query")
    if query.casefold() != inventory.query.casefold():
        raise ContextTopicAnalysisError("Topic analysis query does not match inventory")

    master_value = corpus.get("master_docx_path") or payload.get("master_docx_path")
    master_docx_path = Path(
        _required_string(master_value, "corpus.master_docx_path")
    ).expanduser().resolve()
    if master_docx_path.suffix.casefold() != ".docx":
        raise ContextTopicAnalysisError("Recorded master document is not a .docx path")
    master_layout_pages = _positive_int(
        corpus.get("master_layout_pages"), "corpus.master_layout_pages"
    )
    reading_words_per_page = _positive_int(
        corpus.get("reading_words_per_page"), "corpus.reading_words_per_page"
    )

    raw_families = taxonomy.get("families")
    raw_topics = taxonomy.get("topics")
    raw_classifications = payload.get("classifications") or payload.get("regions")
    if not isinstance(raw_families, list) or not raw_families:
        raise ContextTopicAnalysisError("taxonomy.families must be a non-empty list")
    if not isinstance(raw_topics, list) or not raw_topics:
        raise ContextTopicAnalysisError("taxonomy.topics must be a non-empty list")
    if not isinstance(raw_classifications, list):
        raise ContextTopicAnalysisError("classifications must be a list")

    families = tuple(_parse_family(item, index) for index, item in enumerate(raw_families, 1))
    topics = tuple(_parse_topic(item, index) for index, item in enumerate(raw_topics, 1))
    classifications = tuple(
        _parse_classification(item, index)
        for index, item in enumerate(raw_classifications, 1)
    )
    _validate_analysis_membership(families, topics, classifications, inventory)

    notes = payload.get("review")
    notes_map = notes if isinstance(notes, Mapping) else {}
    overlap_notes = _mapping_sequence(
        payload.get("overlaps") or notes_map.get("overlaps"), "overlaps"
    )
    boundary_notes = _mapping_sequence(
        payload.get("boundary_reviews") or notes_map.get("boundary_reviews"),
        "boundary_reviews",
    )
    taxonomy_sha256 = _required_sha256(payload.get("taxonomy_sha256"), "taxonomy_sha256")
    if _sha256_json(taxonomy) != taxonomy_sha256:
        raise ContextTopicAnalysisError("taxonomy_sha256 does not match the taxonomy")
    status = str(payload.get("status") or "proposed").strip().casefold()
    if status not in {"complete", "proposed", "proposal", "draft"}:
        raise ContextTopicAnalysisError(
            "Analysis must be complete or explicitly marked as a proposal"
        )
    recommendation = str(payload.get("recommendation") or "").strip()
    if not recommendation:
        recommendation = (
            "Retain the exact-source master compilation as the complete authority, "
            "and use the proposed subtopics as smaller reading and review volumes."
        )

    return _Analysis(
        path=candidate,
        sha256=hashlib.sha256(data).hexdigest(),
        profile=_required_string(
            payload.get("analysis_profile") or payload.get("topic_analysis_profile"),
            "topic_analysis_profile",
        ),
        model=_required_string(payload.get("model"), "model"),
        query=query,
        status=status,
        recommendation=recommendation,
        master_docx_path=master_docx_path,
        master_layout_pages=master_layout_pages,
        reading_words_per_page=reading_words_per_page,
        inventory_pair_fingerprint=fingerprint,
        taxonomy_sha256=taxonomy_sha256,
        families=families,
        topics=topics,
        classifications=classifications,
        overlap_notes=overlap_notes,
        boundary_notes=boundary_notes,
    )


def _parse_family(value: Any, index: int) -> _Family:
    item = _mapping(value, f"taxonomy.families[{index}]")
    return _Family(
        family_id=_required_string(item.get("family_id"), f"family {index} id"),
        label=_required_string(item.get("label"), f"family {index} label"),
        definition=_required_string(item.get("definition"), f"family {index} definition"),
    )


def _parse_topic(value: Any, index: int) -> _Topic:
    item = _mapping(value, f"taxonomy.topics[{index}]")
    return _Topic(
        topic_id=_required_string(item.get("topic_id"), f"topic {index} id"),
        family_id=_required_string(item.get("family_id"), f"topic {index} family_id"),
        label=_required_string(item.get("label"), f"topic {index} label"),
        definition=_required_string(item.get("definition"), f"topic {index} definition"),
        include_cues=_string_sequence(item.get("include_cues"), f"topic {index} include_cues"),
        exclude_cues=_string_sequence(item.get("exclude_cues"), f"topic {index} exclude_cues"),
    )


def _parse_classification(value: Any, index: int) -> _Classification:
    item = _mapping(value, f"classifications[{index}]")
    status_value = item.get("status")
    status = (
        _choice(
            status_value,
            {"classified", "taxonomy_gap"},
            f"classification {index} status",
        )
        if status_value is not None
        else None
    )
    primary = item.get("primary_topic_id")
    if primary is not None and not isinstance(primary, str):
        raise ContextTopicAnalysisError(f"classification {index} primary topic is invalid")
    secondary = _string_sequence(
        item.get("secondary_topic_ids"), f"classification {index} secondary topics"
    )
    evidence = item.get("evidence_paragraph_numbers")
    if not isinstance(evidence, list) or any(
        isinstance(number, bool) or not isinstance(number, int) or number < 1
        for number in evidence
    ):
        raise ContextTopicAnalysisError(
            f"classification {index} evidence paragraph numbers are invalid"
        )
    primary_id = primary.strip() if isinstance(primary, str) and primary.strip() else None
    taxonomy_gap_value = item.get("taxonomy_gap")
    if taxonomy_gap_value is not None and not isinstance(taxonomy_gap_value, bool):
        raise ContextTopicAnalysisError(
            f"classification {index} taxonomy_gap must be boolean"
        )
    taxonomy_gap = bool(taxonomy_gap_value) or status == "taxonomy_gap"
    if primary_id is not None and primary_id.casefold() in _UNCLASSIFIED_TOPIC_IDS:
        primary_id = None
        taxonomy_gap = True
    if status == "classified" and primary_id is None:
        raise ContextTopicAnalysisError(
            f"classification {index} is classified but has no primary topic"
        )
    if taxonomy_gap and primary_id is not None:
        raise ContextTopicAnalysisError(
            f"classification {index} is a taxonomy gap but has a primary topic"
        )
    if primary_id is None and not taxonomy_gap:
        raise ContextTopicAnalysisError(
            f"classification {index} has no primary topic and is not a taxonomy gap"
        )
    ambiguity_value = item.get("ambiguity_codes")
    if ambiguity_value is None:
        ambiguity_value = item.get("ambiguity")
    ambiguity_codes = _normalise_ambiguity_codes(
        ambiguity_value,
        index=index,
        taxonomy_gap=taxonomy_gap,
    )
    review_status = _normalise_review_status(item, index=index)
    return _Classification(
        region_id=_required_string(item.get("region_id"), f"classification {index} region_id"),
        primary_topic_id=primary_id,
        secondary_topic_ids=secondary,
        certainty=_choice(item.get("certainty"), {"high", "medium", "low"}, "certainty"),
        ambiguity_codes=ambiguity_codes,
        review_status=review_status,
        taxonomy_gap=taxonomy_gap,
        evidence_paragraph_numbers=tuple(evidence),
        selected_text_sha256=(
            _required_sha256(
                item.get("selected_text_sha256"),
                f"classification {index} selected_text_sha256",
            )
            if item.get("selected_text_sha256") is not None
            else None
        ),
    )


def _normalise_ambiguity_codes(
    value: Any,
    *,
    index: int,
    taxonomy_gap: bool,
) -> tuple[str, ...]:
    if value is None:
        raw: tuple[str, ...] = ()
    elif isinstance(value, str):
        raw = (value,)
    elif isinstance(value, list) and all(
        isinstance(item, str) and item.strip() for item in value
    ):
        raw = tuple(value)
    else:
        raise ContextTopicAnalysisError(
            f"classification {index} ambiguity must be a string or string list"
        )
    result: list[str] = []
    for item in raw:
        key = item.strip().casefold().replace("-", "_").replace(" ", "_")
        if key not in _AMBIGUITY_ALIASES:
            raise ContextTopicAnalysisError(
                f"classification {index} uses unsupported ambiguity code: {item}"
            )
        canonical = _AMBIGUITY_ALIASES[key]
        if canonical is not None and canonical not in result:
            result.append(canonical)
    if taxonomy_gap and "taxonomy_gap" not in result:
        result.append("taxonomy_gap")
    return tuple(result)


def _normalise_review_status(item: Mapping[str, Any], *, index: int) -> str:
    required = item.get("review_required")
    if required is not None and not isinstance(required, bool):
        raise ContextTopicAnalysisError(
            f"classification {index} review_required must be boolean"
        )
    if required is True:
        return "review_required"
    value = item.get("review_status") or item.get("model_review_status")
    if value is None:
        return "accepted" if required is False else "review_recommended"
    if not isinstance(value, str) or not value.strip():
        raise ContextTopicAnalysisError(
            f"classification {index} review status is invalid"
        )
    key = value.strip().casefold().replace("-", "_").replace(" ", "_")
    canonical = _REVIEW_ALIASES.get(key)
    if canonical is None:
        raise ContextTopicAnalysisError(
            f"classification {index} uses unsupported review status: {value}"
        )
    return canonical


def _validate_analysis_membership(
    families: Sequence[_Family],
    topics: Sequence[_Topic],
    classifications: Sequence[_Classification],
    inventory: BoundContextCompilation,
) -> None:
    family_ids = [family.family_id for family in families]
    topic_ids = [topic.topic_id for topic in topics]
    if len(set(family_ids)) != len(family_ids):
        raise ContextTopicAnalysisError("Taxonomy contains duplicate family IDs")
    if len(set(topic_ids)) != len(topic_ids):
        raise ContextTopicAnalysisError("Taxonomy contains duplicate topic IDs")
    if any(topic.family_id not in set(family_ids) for topic in topics):
        raise ContextTopicAnalysisError("Taxonomy topic refers to an unknown family")
    known_topics = set(topic_ids)
    region_by_id = {region.region_id: region for region in inventory.regions}
    known_regions = set(region_by_id)
    classified_regions = [item.region_id for item in classifications]
    if len(set(classified_regions)) != len(classified_regions):
        raise ContextTopicAnalysisError("A region is classified more than once")
    if set(classified_regions) != known_regions:
        missing = len(known_regions - set(classified_regions))
        extra = len(set(classified_regions) - known_regions)
        raise ContextTopicAnalysisError(
            f"Classification coverage differs from inventory (missing={missing}, extra={extra})"
        )
    if classified_regions != [region.region_id for region in inventory.regions]:
        raise ContextTopicAnalysisError(
            "Classification order does not match the canonical inventory order"
        )
    for item in classifications:
        if item.primary_topic_id is not None and item.primary_topic_id not in known_topics:
            raise ContextTopicAnalysisError("Classification uses an unknown primary topic")
        if any(topic_id not in known_topics for topic_id in item.secondary_topic_ids):
            raise ContextTopicAnalysisError("Classification uses an unknown secondary topic")
        if item.primary_topic_id in item.secondary_topic_ids:
            raise ContextTopicAnalysisError("Primary topic is repeated as a secondary topic")
        region = region_by_id[item.region_id]
        if (
            item.selected_text_sha256 is not None
            and item.selected_text_sha256 != region.selected_text_sha256
        ):
            raise ContextTopicAnalysisError(
                "Classification selected-text hash differs from its bound passage"
            )
        paragraph_numbers = {paragraph.number for paragraph in region.selected_paragraphs}
        if any(number not in paragraph_numbers for number in item.evidence_paragraph_numbers):
            raise ContextTopicAnalysisError(
                "Classification evidence refers outside its selected local passage"
            )


def _calculate_topic_metrics(
    analysis: _Analysis,
    inventory: BoundContextCompilation,
    region_by_id: Mapping[str, Any],
) -> tuple[_TopicMetrics, ...]:
    primary: dict[str, list[str]] = defaultdict(list)
    secondary: dict[str, list[str]] = defaultdict(list)
    by_region = {item.region_id: item for item in analysis.classifications}
    for item in analysis.classifications:
        if item.primary_topic_id is not None:
            primary[item.primary_topic_id].append(item.region_id)
        for topic_id in item.secondary_topic_ids:
            secondary[topic_id].append(item.region_id)

    corpus_words = sum(_word_count(region.selected_text) for region in inventory.regions)
    results: list[_TopicMetrics] = []
    for topic in analysis.topics:
        primary_ids = tuple(primary.get(topic.topic_id, ()))
        secondary_ids = tuple(secondary.get(topic.topic_id, ()))
        primary_regions = [region_by_id[region_id] for region_id in primary_ids]
        unique: dict[str, Any] = {}
        for region in primary_regions:
            unique.setdefault(region.selected_text_sha256, region)
        primary_words = sum(_word_count(region.selected_text) for region in primary_regions)
        unique_words = sum(_word_count(region.selected_text) for region in unique.values())
        master_pages = (
            analysis.master_layout_pages * primary_words / corpus_words
            if corpus_words
            else 0.0
        )
        classifications = [by_region[region_id] for region_id in primary_ids]
        results.append(
            _TopicMetrics(
                topic=topic,
                primary_region_ids=primary_ids,
                secondary_region_ids=secondary_ids,
                unique_primary_passages=len(unique),
                primary_body_words=primary_words,
                unique_primary_body_words=unique_words,
                approximate_master_pages=master_pages,
                consolidated_reading_pages=(
                    math.ceil(unique_words / analysis.reading_words_per_page)
                    if unique_words
                    else 0
                ),
                certainty_counts=Counter(item.certainty for item in classifications),
                review_counts=Counter(item.review_status for item in classifications),
                ambiguity_counts=Counter(
                    code
                    for classification in classifications
                    for code in (classification.ambiguity_codes or ("none",))
                ),
            )
        )
    return tuple(results)


def _resolve_output(
    analysis: _Analysis,
    inventory: BoundContextCompilation,
    output_path: Path | str | None,
) -> Path:
    if output_path is None:
        output = analysis.master_docx_path.with_name(
            f"{analysis.master_docx_path.stem} - Subtopic Plan.docx"
        )
    else:
        output = Path(output_path).expanduser().resolve()
    if output.suffix.casefold() != ".docx":
        raise ContextTopicReportOutputError("Report output must use the .docx extension")
    protected = {
        analysis.master_docx_path,
        Path(inventory.docx_path).resolve(),
        Path(inventory.jsonl_path).resolve(),
        analysis.path,
    }
    if output in protected:
        raise ContextTopicReportOutputError("Report output cannot replace an input artefact")
    return output


def _build_document(
    analysis: _Analysis,
    inventory: BoundContextCompilation,
    metrics: Sequence[_TopicMetrics],
    region_by_id: Mapping[str, Any],
    *,
    representative_source_limit: int,
) -> Document:
    document = Document()
    _configure_document(document, analysis)
    _add_opening(document, analysis, inventory, metrics)
    document.add_page_break()
    _add_integrity_ledger(document, analysis, inventory, metrics)
    document.add_page_break()
    _add_overview(document, analysis, inventory, metrics, region_by_id)
    document.add_page_break()
    _add_topic_details(
        document,
        analysis,
        metrics,
        region_by_id,
        representative_source_limit=representative_source_limit,
    )
    document.add_page_break()
    _add_overlap_review(document, analysis, region_by_id)
    _add_boundary_review(document, analysis, region_by_id)
    _add_methodology(document, analysis)
    return document


def _configure_document(document: Document, analysis: _Analysis) -> None:
    # compact_reference_guide exact preset tokens.
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    # Explicitly populate all three Word footer variants.  A restrained footer
    # carries status and pagination; the header is intentionally left clear so
    # Word installations with inherited odd/even settings render consistently.
    section.different_first_page_header_footer = True
    document.settings.odd_and_even_pages_header_footer = True

    normal = document.styles["Normal"]
    _set_style_font(normal, "Calibri", 11, "1C1F22")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    heading_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for name, (size, colour, before, after) in heading_tokens.items():
        style = document.styles[name]
        _set_style_font(style, "Calibri", size, colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    dense = document.styles.add_style("Subtopic Table Text", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(dense, "Calibri", 9, "1C1F22")
    dense.paragraph_format.space_before = Pt(0)
    dense.paragraph_format.space_after = Pt(0)
    dense.paragraph_format.line_spacing = 1.0
    citation = document.styles.add_style("Subtopic Table Note", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(citation, "Calibri", 9, "606970")
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    citation.paragraph_format.line_spacing = 1.0

    properties = document.core_properties
    properties.title = f'{analysis.query.title()} - Proposed Subtopic Plan'
    properties.author = "AudioProcessor Context Finder"
    properties.subject = REPORT_MARKER
    properties.keywords = REPORT_MARKER

    for footer in (
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        _configure_footer(footer)


def _configure_footer(footer: Any) -> None:
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    status = paragraph.add_run(f"{REPORT_STATUS} | Page ")
    _set_run_font(status, 8.5, "787E84")
    page_run = paragraph.add_run()
    _set_run_font(page_run, 8.5, "787E84")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    page_run._r.extend((begin, instruction, end))


def _add_opening(
    document: Document,
    analysis: _Analysis,
    inventory: BoundContextCompilation,
    metrics: Sequence[_TopicMetrics],
) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("RESEARCH ORGANISATION PLAN")
    _set_run_font(run, 9.5, "7A5A00", bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(f"{analysis.query.title()}: Proposed Subtopic Structure")
    _set_run_font(run, 25, "203748", bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.keep_with_next = True
    run = subtitle.add_run(
        f"A subdivision plan for the {analysis.master_layout_pages}-page exact-source master compilation"
    )
    _set_run_font(run, 11.5, "606970")

    _add_callout(
        document,
        REPORT_STATUS,
        "This is an analytical proposal for human review. The master compilation remains "
        "the complete authority; no source passage has been edited, removed, or approved "
        "for subdivision by this report.",
    )
    document.add_heading("Executive recommendation", level=1)
    paragraph = document.add_paragraph(analysis.recommendation)
    paragraph.paragraph_format.keep_with_next = True
    assigned = sum(bool(item.primary_region_ids) for item in metrics)
    document.add_paragraph(
        f"The proposal contains {_counted(len(analysis.families), 'topic family')} and "
        f"{_counted(len(analysis.topics), 'topic')} ({assigned} with at least one primary passage). "
        f"It organises {inventory.region_count:,} exact context sections while retaining "
        f"the original {analysis.master_layout_pages}-page master as a single, "
        "searchable reference volume."
    )


def _add_integrity_ledger(
    document: Document,
    analysis: _Analysis,
    inventory: BoundContextCompilation,
    metrics: Sequence[_TopicMetrics],
) -> None:
    document.add_heading("Corpus integrity ledger", level=1)
    unique_hashes = {region.selected_text_sha256 for region in inventory.regions}
    body_words = sum(_word_count(region.selected_text) for region in inventory.regions)
    unique_words = sum(
        _word_count(region.selected_text)
        for key, region in _first_by_hash(inventory.regions).items()
    )
    unclassified = sum(
        item.primary_topic_id is None for item in analysis.classifications
    )
    rows = [
        ("Status", REPORT_STATUS),
        ("Master document", str(analysis.master_docx_path)),
        ("Query", analysis.query),
        ("Master layout", f"{analysis.master_layout_pages:,} pages (recorded)"),
        ("Corpus", f"{_counted(inventory.region_count, 'region')} | {_counted(inventory.source_count, 'source')} | {_counted(inventory.occurrence_count, 'exact occurrence')}"),
        ("Passages", f"{_counted(len(unique_hashes), 'unique text passage')} | {body_words:,} body words | {unique_words:,} duplicate-consolidated words"),
        ("Classification", f"{inventory.region_count - unclassified:,} primary assignments | {unclassified:,} taxonomy gaps/unassigned"),
        ("Inventory fingerprint", inventory.pair_fingerprint),
        ("Bound master snapshot", str(inventory.docx_path)),
        ("Bound snapshot SHA-256", inventory.docx_sha256),
        ("Context JSONL SHA-256", inventory.jsonl_sha256),
        ("Analysis JSON SHA-256", analysis.sha256),
        ("Taxonomy SHA-256", analysis.taxonomy_sha256),
        ("Analysis model", analysis.model),
    ]
    _add_label_detail_table(document, rows)
    note = document.add_paragraph(style="Subtopic Table Note")
    note.add_run(
        "Integrity rule: topic membership is joined by canonical region_id against the "
        "bound local inventory; quotations are not copied into this planning report."
    )


def _add_overview(
    document: Document,
    analysis: _Analysis,
    inventory: BoundContextCompilation,
    metrics: Sequence[_TopicMetrics],
    region_by_id: Mapping[str, Any],
) -> None:
    document.add_heading("Family and topic overview", level=1)
    document.add_paragraph(
        "Counts and sizes below use primary assignments so the proposed volumes remain "
        "additive. Secondary memberships are shown separately as cross-topic reading cues. "
        "Master pages are proportional estimates; consolidated pages count each identical "
        "passage once at the recorded reading-page density."
    )
    family_by_id = {family.family_id: family for family in analysis.families}
    corpus_words = sum(_word_count(region.selected_text) for region in inventory.regions)
    document.add_heading("Family totals", level=2)
    family_rows: list[tuple[str, ...]] = []
    for family in analysis.families:
        members = [item for item in metrics if item.topic.family_id == family.family_id]
        primary_ids = tuple(
            dict.fromkeys(
                region_id for item in members for region_id in item.primary_region_ids
            )
        )
        secondary_memberships = sum(len(item.secondary_region_ids) for item in members)
        primary_regions = [region_by_id[region_id] for region_id in primary_ids]
        unique = _first_by_hash(primary_regions)
        words = sum(_word_count(region.selected_text) for region in primary_regions)
        unique_words = sum(_word_count(region.selected_text) for region in unique.values())
        master_pages = analysis.master_layout_pages * words / corpus_words if corpus_words else 0.0
        reading_pages = (
            math.ceil(unique_words / analysis.reading_words_per_page)
            if unique_words
            else 0
        )
        family_rows.append(
            (
                family.label,
                f"{len(members):,}",
                f"{len(primary_ids):,} / {secondary_memberships:,}",
                f"{len(unique):,}",
                f"{words:,}",
                f"{master_pages:.1f} / {reading_pages:,}",
            )
        )
    _add_matrix_table(
        document,
        ("Family", "Topics", "Regions\nP / S", "Unique", "Body words", "Pages\nmaster / read"),
        family_rows,
        widths=(2850, 850, 1100, 900, 1250, 2410),
        numeric_columns={1, 2, 3, 4, 5},
    )

    document.add_heading("Topic totals", level=2)
    headers = (
        "Family",
        "Topic",
        "Regions\nP / S",
        "Unique",
        "Body words",
        "Pages\nmaster / read",
    )
    rows: list[tuple[str, ...]] = []
    for item in metrics:
        family = family_by_id[item.topic.family_id]
        rows.append(
            (
                family.label,
                item.topic.label,
                f"{len(item.primary_region_ids):,} / {len(item.secondary_region_ids):,}",
                f"{item.unique_primary_passages:,}",
                f"{item.primary_body_words:,}",
                f"{item.approximate_master_pages:.1f} / {item.consolidated_reading_pages:,}",
            )
        )
    _add_matrix_table(
        document,
        headers,
        rows,
        widths=(1500, 3100, 1000, 800, 1150, 1810),
        numeric_columns={2, 3, 4, 5},
    )


def _add_topic_details(
    document: Document,
    analysis: _Analysis,
    metrics: Sequence[_TopicMetrics],
    region_by_id: Mapping[str, Any],
    *,
    representative_source_limit: int,
) -> None:
    document.add_heading("Detailed topic proposals", level=1)
    family_by_id = {family.family_id: family for family in analysis.families}
    for topic_index, item in enumerate(metrics):
        if topic_index:
            document.add_page_break()
        family = family_by_id[item.topic.family_id]
        document.add_heading(f"{family.label} | {item.topic.label}", level=2)
        document.add_paragraph(item.topic.definition)
        rows = [
            ("Family", f"{family.label} - {family.definition}"),
            ("Membership", f"{_counted(len(item.primary_region_ids), 'primary region')}; {_counted(len(item.secondary_region_ids), 'secondary membership')}"),
            ("Reading size", f"{item.primary_body_words:,} body words; {_counted(item.unique_primary_passages, 'unique passage')}; about {item.approximate_master_pages:.1f} master-layout pages; {_counted(item.consolidated_reading_pages, 'consolidated reading page')}"),
            ("Certainty", _format_counts(item.certainty_counts, ("high", "medium", "low"))),
            ("Review", _format_counts(item.review_counts, ("review_required", "review_recommended", "accepted"))),
            ("Ambiguity", _format_counts(item.ambiguity_counts, ("taxonomy_overlap", "mixed_passage", "boundary_uncertain", "insufficient_context", "taxonomy_gap", "none"))),
            ("Include cues", "; ".join(item.topic.include_cues) or "None recorded"),
            ("Exclude cues", "; ".join(item.topic.exclude_cues) or "None recorded"),
        ]
        _add_label_detail_table(document, rows)
        representatives = _representative_regions(
            item,
            region_by_id,
            representative_source_limit,
        )
        document.add_heading("Representative local sources", level=3)
        if representatives:
            source_rows = tuple(
                (region.heading_text, region.source_relative_path)
                for region in representatives
            )
            _add_matrix_table(
                document,
                ("Source heading", "Bound source path"),
                source_rows,
                widths=(3600, 5760),
            )
        else:
            document.add_paragraph("No primary passage is assigned to this proposed topic.")


def _add_overlap_review(
    document: Document,
    analysis: _Analysis,
    region_by_id: Mapping[str, Any],
) -> None:
    document.add_heading("Overlap and ambiguity review", level=1)
    topic_labels = {topic.topic_id: topic.label for topic in analysis.topics}
    overlap = [item for item in analysis.classifications if item.secondary_topic_ids]
    ambiguous = [item for item in analysis.classifications if item.ambiguity_codes]
    document.add_paragraph(
        f"{_counted(len(overlap), 'region')} "
        f"{'has' if len(overlap) == 1 else 'have'} one or more secondary topic memberships; "
        f"{_counted(len(ambiguous), 'region')} "
        f"{'carries' if len(ambiguous) == 1 else 'carry'} an explicit ambiguity flag. These are "
        "cross-reading or review signals, not permission to duplicate or remove passages."
    )
    rows: list[tuple[str, ...]] = []
    for item in ambiguous[:40]:
        region = region_by_id[item.region_id]
        primary = topic_labels.get(item.primary_topic_id or "", "Unassigned")
        secondary = ", ".join(topic_labels[value] for value in item.secondary_topic_ids) or "-"
        ambiguity = ", ".join(
            code.replace("_", " ") for code in item.ambiguity_codes
        )
        rows.append((region.heading_text, primary, secondary, ambiguity, item.review_status.replace("_", " ")))
    if rows:
        _add_matrix_table(
            document,
            ("Representative section", "Primary", "Secondary", "Ambiguity", "Review"),
            rows,
            widths=(2700, 1750, 1750, 1550, 1610),
        )
        if len(ambiguous) > len(rows):
            document.add_paragraph(
                f"The table shows the first {len(rows):,} of {len(ambiguous):,} ambiguous regions in canonical inventory order.",
                style="Subtopic Table Note",
            )
    else:
        document.add_paragraph("No classification carries an explicit ambiguity flag.")
    for note in analysis.overlap_notes:
        summary = str(note.get("summary") or note.get("reason") or "").strip()
        if summary:
            document.add_paragraph(summary)


def _add_boundary_review(
    document: Document,
    analysis: _Analysis,
    region_by_id: Mapping[str, Any],
) -> None:
    document.add_heading("Boundary-review queue", level=1)
    review = [
        item for item in analysis.classifications
        if item.review_status == "review_required" or item.certainty == "low"
    ]
    document.add_paragraph(
        f"{_counted(len(review), 'region')} "
        f"{'requires' if len(review) == 1 else 'require'} boundary or taxonomy review before any subtopic "
        "documents are generated. Human review should confirm the complete local source "
        "section, not only its topic label."
    )
    rows = []
    for item in review[:60]:
        region = region_by_id[item.region_id]
        rows.append(
            (
                region.heading_text,
                region.source_relative_path,
                item.certainty,
                ", ".join(code.replace("_", " ") for code in item.ambiguity_codes)
                or "review required",
                region.locator,
            )
        )
    if rows:
        _add_matrix_table(
            document,
            ("Section", "Source path", "Cert.", "Issue", "Locator"),
            rows,
            widths=(2450, 2700, 1000, 1450, 1760),
        )
        if len(review) > len(rows):
            document.add_paragraph(
                f"The table shows the first {len(rows):,} of {len(review):,} required-review regions in canonical inventory order.",
                style="Subtopic Table Note",
            )
    else:
        document.add_paragraph("No region is currently marked low-certainty or review-required.")
    for note in analysis.boundary_notes:
        summary = str(note.get("summary") or note.get("reason") or "").strip()
        if summary:
            document.add_paragraph(summary)


def _add_methodology(document: Document, analysis: _Analysis) -> None:
    document.add_heading("Methodology and decision status", level=1)
    paragraphs = (
        ("Immutable bind.", " The exact Context Finder DOCX and its canonical JSONL snapshot were bound locally by region ID and content fingerprints before analysis."),
        ("Model-assisted taxonomy.", f" {analysis.model} proposed topic families, inclusion and exclusion cues, and per-region classifications. No remote call is made by this report builder."),
        ("Local metric join.", " Region counts, source headings, source paths and word totals in this report are calculated from the bound inventory, not invented by the model response."),
        ("Duplicate handling.", " Duplicate-consolidated estimates count identical selected-text hashes once within each primary topic; the exact master remains unchanged and may intentionally repeat material."),
        ("Human decision gate.", " Topic labels, overlaps, exclusions and boundaries remain proposals. Approval should follow review of low-certainty and ambiguous regions against the linked primary sources."),
    )
    for label, text in paragraphs:
        paragraph = document.add_paragraph()
        lead = paragraph.add_run(label)
        lead.bold = True
        paragraph.add_run(text)
    _add_callout(
        document,
        "DECISION GATE",
        "Do not generate, replace, publish, or delete subtopic volumes until a human has approved the taxonomy and resolved the boundary-review queue.",
    )


def _add_callout(document: Document, label: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_table_geometry(table, (9360,))
    cell = table.cell(0, 0)
    _set_cell_fill(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.style = document.styles["Normal"]
    paragraph.paragraph_format.space_after = Pt(0)
    lead = paragraph.add_run(f"{label}. ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string("7A5A00")
    paragraph.add_run(text)


def _add_label_detail_table(document: Document, rows: Sequence[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    _set_table_geometry(table, (1700, 7660))
    table.rows[0]._element.getparent().remove(table.rows[0]._element)
    for label, detail in rows:
        cells = table.add_row().cells
        _replace_cell_text(cells[0], label, bold=True, fill="E8EEF5")
        _replace_cell_text(cells[1], detail)


def _add_matrix_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    widths: Sequence[int],
    numeric_columns: set[int] | None = None,
) -> None:
    if len(headers) != len(widths) or sum(widths) != 9360:
        raise ContextTopicReportError("Internal report table geometry is invalid")
    table = document.add_table(rows=1, cols=len(headers))
    _set_table_geometry(table, widths)
    _mark_repeat_header(table.rows[0])
    for index, value in enumerate(headers):
        _replace_cell_text(table.rows[0].cells[index], value, bold=True, fill="E8EEF5")
    for row in rows:
        if len(row) != len(headers):
            raise ContextTopicReportError("Internal report table row is malformed")
        cells = table.add_row().cells
        for index, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.RIGHT if numeric_columns and index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
            _replace_cell_text(cells[index], str(value), alignment=align)


def _set_table_geometry(table: Any, widths: Sequence[int]) -> None:
    table.autofit = False
    table.style = "Table Grid"
    table.alignment = 0
    table_element = table._tbl
    properties = table_element.tblPr
    _set_child_attr(properties, "w:tblW", "w:w", str(sum(widths)), {"w:type": "dxa"})
    _set_child_attr(properties, "w:tblInd", "w:w", "120", {"w:type": "dxa"})
    _set_child_attr(properties, "w:tblLayout", "w:type", "fixed")
    margins = properties.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        properties.append(margins)
    for name, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        _set_child_attr(margins, f"w:{name}", "w:w", str(value), {"w:type": "dxa"})
    grid = table_element.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            _set_cell_width(cell, width)


def _set_child_attr(parent: Any, tag: str, key: str, value: str, extra: Mapping[str, str] | None = None) -> None:
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    child.set(qn(key), value)
    for name, item in (extra or {}).items():
        child.set(qn(name), item)


def _set_cell_width(cell: Any, width: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:tcW"))
    if element is None:
        element = OxmlElement("w:tcW")
        properties.append(element)
    element.set(qn("w:w"), str(width))
    element.set(qn("w:type"), "dxa")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _replace_cell_text(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    fill: str | None = None,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "Subtopic Table Text"
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    _set_run_font(run, 9, "1C1F22", bold=bold)
    if fill:
        _set_cell_fill(cell, fill)


def _set_cell_fill(cell: Any, colour: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), colour)


def _mark_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def _representative_regions(
    item: _TopicMetrics,
    region_by_id: Mapping[str, Any],
    limit: int,
) -> tuple[Any, ...]:
    selected = []
    seen_sources: set[str] = set()
    for region_id in item.primary_region_ids + item.secondary_region_ids:
        region = region_by_id[region_id]
        key = region.source_relative_path.casefold()
        if key in seen_sources:
            continue
        seen_sources.add(key)
        selected.append(region)
        if len(selected) == limit:
            break
    return tuple(selected)


def _protected_signatures(
    analysis: _Analysis,
    inventory: BoundContextCompilation,
) -> Mapping[Path, tuple[int, int]]:
    paths = {
        analysis.path,
        analysis.master_docx_path,
        Path(inventory.docx_path).resolve(),
        Path(inventory.jsonl_path).resolve(),
    }
    result: dict[Path, tuple[int, int]] = {}
    for path in paths:
        if path.is_file():
            stat = path.stat()
            result[path] = (stat.st_size, stat.st_mtime_ns)
    return result


def _assert_protected_unchanged(signatures: Mapping[Path, tuple[int, int]]) -> None:
    for path, expected in signatures.items():
        if not path.is_file():
            raise ContextTopicReportOutputError(f"Protected input disappeared: {path}")
        stat = path.stat()
        if (stat.st_size, stat.st_mtime_ns) != expected:
            raise ContextTopicReportOutputError(f"Protected input changed during report build: {path}")


def _save_docx_atomically(document: Document, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="wb",
            prefix=f".{path.stem}.",
            suffix=".tmp.docx",
            dir=path.parent,
            delete=False,
        ) as handle:
            temporary = Path(handle.name)
        document.save(str(temporary))
        Document(BytesIO(temporary.read_bytes()))
        os.replace(temporary, path)
        temporary = None
        return path
    finally:
        if temporary is not None:
            temporary.unlink(missing_ok=True)


def _set_style_font(style: Any, name: str, size: float, colour: str) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(colour)


def _set_run_font(run: Any, size: float, colour: str, *, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(colour)
    if bold is not None:
        run.bold = bold


def _first_by_hash(regions: Sequence[Any]) -> Mapping[str, Any]:
    result: dict[str, Any] = {}
    for region in regions:
        result.setdefault(region.selected_text_sha256, region)
    return result


def _format_counts(counts: Mapping[str, int], order: Sequence[str]) -> str:
    return "; ".join(f"{key.replace('_', ' ')} {counts.get(key, 0):,}" for key in order)


def _counted(count: int, singular: str) -> str:
    suffix = "" if count == 1 else "s"
    return f"{count:,} {singular}{suffix}"


def _word_count(text: str) -> int:
    return len(_WORD_RE.findall(text))


def _sha256_json(value: Any) -> str:
    encoded = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _mapping(value: Any, label: str) -> Mapping[str, Any]:
    if not isinstance(value, Mapping):
        raise ContextTopicAnalysisError(f"{label} must be an object")
    return value


def _mapping_sequence(value: Any, label: str) -> tuple[Mapping[str, Any], ...]:
    if value is None:
        return ()
    if not isinstance(value, list) or any(not isinstance(item, Mapping) for item in value):
        raise ContextTopicAnalysisError(f"{label} must be a list of objects")
    return tuple(value)


def _string_sequence(value: Any, label: str) -> tuple[str, ...]:
    if not isinstance(value, list) or any(not isinstance(item, str) or not item.strip() for item in value):
        raise ContextTopicAnalysisError(f"{label} must be a list of non-empty strings")
    cleaned = tuple(item.strip() for item in value)
    if len(set(cleaned)) != len(cleaned):
        raise ContextTopicAnalysisError(f"{label} contains duplicates")
    return cleaned


def _required_string(value: Any, label: str) -> str:
    if not isinstance(value, str) or not value.strip():
        raise ContextTopicAnalysisError(f"{label} must be a non-empty string")
    return value.strip()


def _required_sha256(value: Any, label: str) -> str:
    text = _required_string(value, label).casefold()
    if _SHA256_RE.fullmatch(text) is None:
        raise ContextTopicAnalysisError(f"{label} must be a SHA-256 digest")
    return text


def _positive_int(value: Any, label: str) -> int:
    if isinstance(value, bool) or not isinstance(value, int) or value < 1:
        raise ContextTopicAnalysisError(f"{label} must be a positive integer")
    return value


def _choice(value: Any, choices: set[str], label: str) -> str:
    text = _required_string(value, label)
    if text not in choices:
        raise ContextTopicAnalysisError(f"{label} has unsupported value: {text}")
    return text


__all__ = [
    "REPORT_MARKER",
    "REPORT_SCHEMA_VERSION",
    "REPORT_STATUS",
    "ContextTopicAnalysisError",
    "ContextTopicReportError",
    "ContextTopicReportOutputError",
    "create_subtopic_plan_report",
]
