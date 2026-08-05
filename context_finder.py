"""Exact, resumable context search for a local transcript library.

The core deliberately separates retrieval from optional language-model work:

* source text is extracted and matched deterministically;
* broad paragraph windows are stored verbatim in JSONL records;
* a model may select narrower paragraph boundaries, but cannot rewrite text;
* the publication DOCX is always rebuilt from the stored source paragraphs.

This makes the JSONL file a safe resume boundary for a future GLM boundary
refiner while keeping every published quotation traceable to its source.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import tempfile
from dataclasses import asdict, dataclass, replace
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Inches, Pt, RGBColor  # type: ignore
from docx.table import Table  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore


SCHEMA_VERSION = "context-finder-v1"
COMPILATION_MARKER = "AudioProcessor exact-context compilation"
DEFAULT_CONTEXT_WORDS_EACH_SIDE = 500
DEFAULT_MAX_REGION_PARAGRAPHS = 48
DEFAULT_MAX_REGION_CHARACTERS = 24_000
MAX_REGION_PARAGRAPH_SPAN = 64
MAX_REGION_PARAGRAPH_MAP_BYTES = 40_000
SUPPORTED_SUFFIXES = (".docx", ".md", ".txt")
GLM_REVIEW_SUFFIX = " - GLM Review"


@dataclass(frozen=True, slots=True)
class QuerySpec:
    """A validated one-to-three-word exact-search query."""

    text: str
    canonical: str
    word_count: int


@dataclass(frozen=True, slots=True)
class ParagraphSnapshot:
    """One source paragraph, numbered in deterministic extraction order."""

    number: int
    text: str
    page_number: int | None = None


@dataclass(frozen=True, slots=True)
class OccurrenceRecord:
    """An exact match located within a source paragraph."""

    occurrence_id: str
    paragraph_number: int
    start: int
    end: int
    matched_text: str
    page_number: int | None = None


@dataclass(frozen=True, slots=True)
class BoundarySelection:
    """Exact paragraph boundaries selected for publication.

    A GLM integration should return only ``start_paragraph`` and
    ``end_paragraph`` (plus optional metadata) and then call
    :func:`apply_boundary_selection`. It must never return replacement prose.
    """

    start_paragraph: int
    end_paragraph: int
    method: str = "deterministic_context_window"
    model: str | None = None
    confidence: float | None = None
    note: str | None = None


@dataclass(frozen=True, slots=True)
class ContextRegion:
    """A merged, de-duplicated context region from one source document."""

    region_id: str
    query: str
    source_relative_path: str
    source_absolute_path: str
    source_sha256: str
    source_suffix: str
    broad_start_paragraph: int
    broad_end_paragraph: int
    paragraphs: tuple[ParagraphSnapshot, ...]
    occurrences: tuple[OccurrenceRecord, ...]
    selection: BoundarySelection

    @property
    def selected_paragraphs(self) -> tuple[ParagraphSnapshot, ...]:
        """Return the exact stored paragraphs inside the selected bounds."""

        return tuple(
            paragraph
            for paragraph in self.paragraphs
            if self.selection.start_paragraph
            <= paragraph.number
            <= self.selection.end_paragraph
        )

    @property
    def occurrence_count(self) -> int:
        return len(self.occurrences)

    def boundary_payload(self) -> dict[str, Any]:
        """Return a model-ready payload containing only exact source text.

        The consumer is expected to choose inclusive paragraph numbers. The
        validation function prevents it from dropping a matched occurrence.
        """

        return {
            "schema_version": SCHEMA_VERSION,
            "region_id": self.region_id,
            "query": self.query,
            "source_relative_path": self.source_relative_path,
            "allowed_start_paragraph": self.broad_start_paragraph,
            "allowed_end_paragraph": self.broad_end_paragraph,
            "must_include_paragraphs": sorted(
                {occurrence.paragraph_number for occurrence in self.occurrences}
            ),
            "paragraphs": [asdict(paragraph) for paragraph in self.paragraphs],
        }

    def to_record(self) -> dict[str, Any]:
        return {
            "record_type": "context_region",
            "schema_version": SCHEMA_VERSION,
            "region_id": self.region_id,
            "query": self.query,
            "source_relative_path": self.source_relative_path,
            "source_absolute_path": self.source_absolute_path,
            "source_sha256": self.source_sha256,
            "source_suffix": self.source_suffix,
            "broad_start_paragraph": self.broad_start_paragraph,
            "broad_end_paragraph": self.broad_end_paragraph,
            "paragraphs": [asdict(paragraph) for paragraph in self.paragraphs],
            "occurrences": [asdict(occurrence) for occurrence in self.occurrences],
            "selection": asdict(self.selection),
        }

    @classmethod
    def from_record(cls, record: dict[str, Any]) -> "ContextRegion":
        if record.get("schema_version") != SCHEMA_VERSION:
            raise ValueError(
                f"Unsupported context record schema: {record.get('schema_version')!r}"
            )
        return cls(
            region_id=str(record["region_id"]),
            query=str(record["query"]),
            source_relative_path=str(record["source_relative_path"]),
            source_absolute_path=str(record["source_absolute_path"]),
            source_sha256=str(record["source_sha256"]),
            source_suffix=str(record["source_suffix"]),
            broad_start_paragraph=int(record["broad_start_paragraph"]),
            broad_end_paragraph=int(record["broad_end_paragraph"]),
            paragraphs=tuple(
                ParagraphSnapshot(**paragraph) for paragraph in record["paragraphs"]
            ),
            occurrences=tuple(
                OccurrenceRecord(**occurrence) for occurrence in record["occurrences"]
            ),
            selection=BoundarySelection(**record["selection"]),
        )


@dataclass(frozen=True, slots=True)
class ScanIssue:
    source_relative_path: str
    error_type: str
    message: str


@dataclass(frozen=True, slots=True)
class SearchOptions:
    context_words_each_side: int = DEFAULT_CONTEXT_WORDS_EACH_SIDE
    max_region_paragraphs: int = DEFAULT_MAX_REGION_PARAGRAPHS
    max_region_characters: int = DEFAULT_MAX_REGION_CHARACTERS
    prefer_glm_review: bool = True
    suffixes: tuple[str, ...] = SUPPORTED_SUFFIXES
    strict: bool = False

    def __post_init__(self) -> None:
        if self.context_words_each_side < 0:
            raise ValueError("context_words_each_side must be zero or greater")
        if self.max_region_paragraphs < 1:
            raise ValueError("max_region_paragraphs must be at least one")
        if self.max_region_paragraphs > DEFAULT_MAX_REGION_PARAGRAPHS:
            raise ValueError(
                f"max_region_paragraphs cannot exceed the endpoint-safe limit "
                f"of {DEFAULT_MAX_REGION_PARAGRAPHS}"
            )
        if self.max_region_characters < 1:
            raise ValueError("max_region_characters must be at least one")
        normalised = tuple(sorted({suffix.casefold() for suffix in self.suffixes}))
        if not normalised or any(not suffix.startswith(".") for suffix in normalised):
            raise ValueError("suffixes must contain extensions such as '.docx'")
        object.__setattr__(self, "suffixes", normalised)


@dataclass(frozen=True, slots=True)
class SearchResult:
    schema_version: str
    root: str
    query: QuerySpec
    options: SearchOptions
    scanned_files: int
    ignored_generated_files: int
    regions: tuple[ContextRegion, ...]
    issues: tuple[ScanIssue, ...] = ()

    @property
    def occurrence_count(self) -> int:
        return sum(region.occurrence_count for region in self.regions)

    @property
    def source_count(self) -> int:
        return len({region.source_relative_path for region in self.regions})

    def with_regions(self, regions: Sequence[ContextRegion]) -> "SearchResult":
        """Return an updated immutable result, useful after GLM selections."""

        expected = {region.region_id for region in self.regions}
        actual = {region.region_id for region in regions}
        if len(regions) != len(self.regions) or expected != actual:
            raise ValueError("Replacement regions must preserve every region_id exactly once")
        ordered = tuple(sorted(regions, key=_region_sort_key))
        return replace(self, regions=ordered)


class _GeneratedCompilation(Exception):
    pass


class SourceIntegrityError(RuntimeError):
    """Raised when a source no longer matches the immutable search snapshot."""


def validate_query(query: str) -> QuerySpec:
    """Normalise and validate an exact word/phrase query (maximum three words)."""

    text = " ".join(str(query).split())
    if not text:
        raise ValueError("Search query cannot be empty")
    words = text.split(" ")
    if len(words) > 3:
        raise ValueError("Search query must contain no more than three words")
    word_re = re.compile(r"[^\W_](?:[\w'’\-]*[^\W_])?", re.UNICODE)
    if any(word_re.fullmatch(word) is None for word in words):
        raise ValueError(
            "Each query word must begin and end with a letter or number; "
            "apostrophes and hyphens may occur inside a word"
        )
    return QuerySpec(text=text, canonical=text.casefold(), word_count=len(words))


def compile_query_pattern(query: QuerySpec | str) -> re.Pattern[str]:
    """Compile a case-insensitive, whole-word exact phrase pattern."""

    spec = validate_query(query) if isinstance(query, str) else query
    phrase = r"\s+".join(re.escape(word) for word in spec.text.split(" "))
    # Apostrophes and hyphens form compounds in the endpoint contract. Thus a
    # query such as ``wake`` must not match ``wake-up`` or ``wake's``.
    compound_word_character = r"[\w'’\-]"
    return re.compile(
        rf"(?<!{compound_word_character}){phrase}(?!{compound_word_character})",
        re.IGNORECASE | re.UNICODE,
    )


def apply_boundary_selection(
    region: ContextRegion,
    start_paragraph: int,
    end_paragraph: int,
    *,
    method: str = "glm_boundary_refinement",
    model: str | None = None,
    confidence: float | None = None,
    note: str | None = None,
) -> ContextRegion:
    """Apply exact paragraph boundaries without accepting rewritten prose."""

    if start_paragraph > end_paragraph:
        raise ValueError("start_paragraph must not exceed end_paragraph")
    if (
        start_paragraph < region.broad_start_paragraph
        or end_paragraph > region.broad_end_paragraph
    ):
        raise ValueError("Selected boundaries must remain inside the broad context")
    available = {paragraph.number for paragraph in region.paragraphs}
    if start_paragraph not in available or end_paragraph not in available:
        raise ValueError("Selected boundaries must identify stored source paragraphs")
    missing_matches = [
        occurrence.paragraph_number
        for occurrence in region.occurrences
        if not start_paragraph <= occurrence.paragraph_number <= end_paragraph
    ]
    if missing_matches:
        raise ValueError("Selected boundaries cannot exclude a matched occurrence")
    if confidence is not None and not 0 <= confidence <= 1:
        raise ValueError("confidence must be between 0 and 1")
    selection = BoundarySelection(
        start_paragraph=start_paragraph,
        end_paragraph=end_paragraph,
        method=method,
        model=model,
        confidence=confidence,
        note=note,
    )
    return replace(region, selection=selection)


def find_contexts(
    root: Path | str,
    query: str,
    *,
    options: SearchOptions | None = None,
    exclude_paths: Iterable[Path | str] = (),
) -> SearchResult:
    """Recursively search supported documents and return exact context regions."""

    root_path = Path(root).expanduser().resolve()
    if not root_path.is_dir():
        raise NotADirectoryError(f"Search root is not a folder: {root_path}")
    spec = validate_query(query)
    pattern = compile_query_pattern(spec)
    options = options or SearchOptions()
    excluded = {Path(path).expanduser().resolve() for path in exclude_paths}

    source_files = sorted(
        (
            path
            for path in root_path.rglob("*")
            if path.is_file()
            and path.suffix.casefold() in options.suffixes
            and not path.name.startswith("~$")
            and path.resolve() not in excluded
        ),
        key=lambda path: path.relative_to(root_path).as_posix().casefold(),
    )
    if options.prefer_glm_review:
        source_files = _prefer_glm_review_sources(source_files)

    regions: list[ContextRegion] = []
    issues: list[ScanIssue] = []
    scanned_files = 0
    ignored_generated_files = 0
    for source_path in source_files:
        relative_path = source_path.relative_to(root_path).as_posix()
        try:
            paragraphs = _extract_paragraphs(source_path)
        except _GeneratedCompilation:
            ignored_generated_files += 1
            continue
        except Exception as exc:
            if options.strict:
                raise
            issues.append(
                ScanIssue(
                    source_relative_path=relative_path,
                    error_type=type(exc).__name__,
                    message=str(exc),
                )
            )
            continue
        scanned_files += 1
        source_sha256 = _sha256_file(source_path)
        regions.extend(
            _regions_for_source(
                source_path=source_path,
                relative_path=relative_path,
                source_sha256=source_sha256,
                paragraphs=paragraphs,
                query=spec,
                pattern=pattern,
                context_words_each_side=options.context_words_each_side,
                max_region_paragraphs=options.max_region_paragraphs,
                max_region_characters=options.max_region_characters,
            )
        )

    return SearchResult(
        schema_version=SCHEMA_VERSION,
        root=str(root_path),
        query=spec,
        options=options,
        scanned_files=scanned_files,
        ignored_generated_files=ignored_generated_files,
        regions=tuple(sorted(regions, key=_region_sort_key)),
        issues=tuple(issues),
    )


def write_result_records(result: SearchResult, path: Path | str) -> Path:
    """Atomically write a deterministic JSONL resume file."""

    destination = Path(path).expanduser().resolve()
    manifest = {
        "record_type": "search_manifest",
        "schema_version": SCHEMA_VERSION,
        "root": result.root,
        "query": asdict(result.query),
        "options": {
            "context_words_each_side": result.options.context_words_each_side,
            "max_region_paragraphs": result.options.max_region_paragraphs,
            "max_region_characters": result.options.max_region_characters,
            "prefer_glm_review": result.options.prefer_glm_review,
            "suffixes": list(result.options.suffixes),
            "strict": result.options.strict,
        },
        "scanned_files": result.scanned_files,
        "ignored_generated_files": result.ignored_generated_files,
        "region_count": len(result.regions),
        "occurrence_count": result.occurrence_count,
        "source_count": result.source_count,
        "issues": [asdict(issue) for issue in result.issues],
    }
    lines = [json.dumps(manifest, ensure_ascii=False, sort_keys=True)]
    lines.extend(
        json.dumps(region.to_record(), ensure_ascii=False, sort_keys=True)
        for region in result.regions
    )
    _write_text_atomically(destination, "\n".join(lines) + "\n")
    return destination


def read_result_records(path: Path | str) -> SearchResult:
    """Load a JSONL resume file and validate its record set."""

    source = Path(path).expanduser().resolve()
    records = [
        json.loads(line)
        for line in source.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    if not records or records[0].get("record_type") != "search_manifest":
        raise ValueError("Context records must begin with a search_manifest")
    manifest = records[0]
    if manifest.get("schema_version") != SCHEMA_VERSION:
        raise ValueError(
            f"Unsupported context record schema: {manifest.get('schema_version')!r}"
        )
    regions = tuple(
        ContextRegion.from_record(record)
        for record in records[1:]
        if record.get("record_type") == "context_region"
    )
    if len(regions) != int(manifest["region_count"]):
        raise ValueError("Context record count does not match its manifest")
    query = QuerySpec(**manifest["query"])
    options_record = manifest["options"]
    options = SearchOptions(
        context_words_each_side=int(options_record["context_words_each_side"]),
        max_region_paragraphs=int(
            options_record.get(
                "max_region_paragraphs", DEFAULT_MAX_REGION_PARAGRAPHS
            )
        ),
        max_region_characters=int(
            options_record.get(
                "max_region_characters", DEFAULT_MAX_REGION_CHARACTERS
            )
        ),
        prefer_glm_review=bool(options_record.get("prefer_glm_review", True)),
        suffixes=tuple(options_record["suffixes"]),
        strict=bool(options_record["strict"]),
    )
    result = SearchResult(
        schema_version=SCHEMA_VERSION,
        root=str(manifest["root"]),
        query=query,
        options=options,
        scanned_files=int(manifest["scanned_files"]),
        ignored_generated_files=int(manifest["ignored_generated_files"]),
        regions=tuple(sorted(regions, key=_region_sort_key)),
        issues=tuple(ScanIssue(**issue) for issue in manifest.get("issues", [])),
    )
    if result.occurrence_count != int(manifest["occurrence_count"]):
        raise ValueError("Occurrence count does not match the records manifest")
    return result


def create_compilation_docx(
    result: SearchResult,
    output_path: Path | str,
    *,
    highlight_occurrences: bool = True,
) -> Path:
    """Create an exact-quotation, publication-quality Word compilation."""

    output = Path(output_path).expanduser().resolve()
    if output.suffix.casefold() != ".docx":
        raise ValueError("Compilation output must use the .docx extension")

    document = Document()
    _configure_compilation_document(document, result)
    _add_opening(document, result)
    query_pattern = compile_query_pattern(result.query)

    source_region_counts: dict[str, int] = {}
    for region_index, region in enumerate(result.regions, start=1):
        source_region_counts[region.source_relative_path] = (
            source_region_counts.get(region.source_relative_path, 0) + 1
        )
        _add_region(
            document,
            region,
            region_index=region_index,
            source_region_index=source_region_counts[region.source_relative_path],
            query_pattern=query_pattern,
            highlight_occurrences=highlight_occurrences,
        )

    if not result.regions:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run("No exact whole-word or whole-phrase occurrences were found.")

    # Validate after the potentially lengthy document build, immediately before
    # the atomic save. This prevents publishing stored quotations with links to
    # source files that changed after the search snapshot was taken.
    validate_source_integrity(result)
    return _save_docx_atomically(document, output)


def validate_source_integrity(result: SearchResult) -> None:
    """Fail closed unless every unique quoted source still matches its scan hash."""

    sources: dict[str, tuple[Path, str, str]] = {}
    for region in result.regions:
        source_path = Path(region.source_absolute_path).expanduser().resolve()
        key = os.path.normcase(str(source_path))
        previous = sources.get(key)
        if previous is not None and previous[1] != region.source_sha256:
            raise SourceIntegrityError(
                f"Conflicting stored hashes for source: {region.source_relative_path}"
            )
        sources[key] = (
            source_path,
            region.source_sha256,
            region.source_relative_path,
        )

    for source_path, expected_hash, relative_path in sources.values():
        try:
            actual_hash = _sha256_file(source_path)
        except OSError as exc:
            raise SourceIntegrityError(
                f"Source is unavailable before publication: {relative_path}"
            ) from exc
        if actual_hash != expected_hash:
            raise SourceIntegrityError(
                f"Source changed after the search snapshot: {relative_path}"
            )


def _extract_paragraphs(source_path: Path) -> tuple[ParagraphSnapshot, ...]:
    suffix = source_path.suffix.casefold()
    if suffix == ".docx":
        return _extract_docx_paragraphs(source_path)
    if suffix in {".txt", ".md"}:
        return _extract_plain_paragraphs(source_path)
    raise ValueError(f"Unsupported source type: {suffix}")


def _prefer_glm_review_sources(source_files: Sequence[Path]) -> list[Path]:
    """Prefer a generated GLM-review DOCX over its raw sibling by default."""

    lookup = {
        (str(path.parent.resolve()).casefold(), path.name.casefold()): path
        for path in source_files
    }
    superseded: set[Path] = set()
    suffix = GLM_REVIEW_SUFFIX.casefold()
    for path in source_files:
        if path.suffix.casefold() != ".docx" or not path.stem.casefold().endswith(suffix):
            continue
        raw_stem = path.stem[: -len(GLM_REVIEW_SUFFIX)]
        raw = lookup.get(
            (str(path.parent.resolve()).casefold(), f"{raw_stem}.docx".casefold())
        )
        if raw is not None:
            superseded.add(raw)
    return [path for path in source_files if path not in superseded]


def _extract_docx_paragraphs(source_path: Path) -> tuple[ParagraphSnapshot, ...]:
    document = Document(str(source_path))
    properties = document.core_properties
    if (
        properties.subject == COMPILATION_MARKER
        or properties.keywords == COMPILATION_MARKER
    ):
        raise _GeneratedCompilation

    texts: list[str] = []
    seen_cells: set[int] = set()
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            texts.append(item.text)
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    cell_identity = id(cell._tc)
                    if cell_identity in seen_cells:
                        continue
                    seen_cells.add(cell_identity)
                    texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return tuple(
        ParagraphSnapshot(number=index, text=text)
        for index, text in enumerate(texts, start=1)
    )


def _extract_plain_paragraphs(source_path: Path) -> tuple[ParagraphSnapshot, ...]:
    data = source_path.read_bytes()
    if data.startswith((b"\xff\xfe", b"\xfe\xff")):
        text = data.decode("utf-16")
    else:
        try:
            text = data.decode("utf-8-sig")
        except UnicodeDecodeError:
            text = data.decode("cp1252")
    blocks = re.split(r"(?:\r?\n)[ \t]*(?:\r?\n)+", text)
    texts = [block.strip("\r\n") for block in blocks if block.strip("\r\n")]
    return tuple(
        ParagraphSnapshot(number=index, text=block)
        for index, block in enumerate(texts, start=1)
    )


def _regions_for_source(
    *,
    source_path: Path,
    relative_path: str,
    source_sha256: str,
    paragraphs: tuple[ParagraphSnapshot, ...],
    query: QuerySpec,
    pattern: re.Pattern[str],
    context_words_each_side: int,
    max_region_paragraphs: int,
    max_region_characters: int,
) -> list[ContextRegion]:
    occurrences: list[OccurrenceRecord] = []
    occurrence_indices: dict[str, int] = {}
    for paragraph_index, paragraph in enumerate(paragraphs):
        for match in pattern.finditer(paragraph.text):
            occurrence_key = (
                f"{relative_path}\0{source_sha256}\0{query.canonical}\0"
                f"{paragraph.number}\0{match.start()}\0{match.end()}"
            )
            occurrence_id = _stable_id("occ", occurrence_key)
            occurrence = OccurrenceRecord(
                occurrence_id=occurrence_id,
                paragraph_number=paragraph.number,
                start=match.start(),
                end=match.end(),
                matched_text=match.group(0),
                page_number=paragraph.page_number,
            )
            occurrences.append(occurrence)
            occurrence_indices[occurrence_id] = paragraph_index
    if not occurrences:
        return []

    spans: list[tuple[int, int, list[OccurrenceRecord]]] = []
    for occurrence in occurrences:
        paragraph_index = occurrence_indices[occurrence.occurrence_id]
        start, end = _context_window(
            paragraphs,
            paragraph_index,
            context_words_each_side=context_words_each_side,
            max_region_paragraphs=max_region_paragraphs,
            max_region_characters=max_region_characters,
        )
        spans.append((start, end, [occurrence]))
    spans.sort(key=lambda value: (value[0], value[1], value[2][0].start))

    merged: list[tuple[int, int, list[OccurrenceRecord]]] = []
    for start, end, span_occurrences in spans:
        can_merge = False
        if merged and start <= merged[-1][1] + 1:
            candidate_start = merged[-1][0]
            candidate_end = max(merged[-1][1], end)
            can_merge = _span_within_limits(
                paragraphs,
                candidate_start,
                candidate_end,
                max_region_paragraphs=max_region_paragraphs,
                max_region_characters=max_region_characters,
            )
        if can_merge:
            old_start, old_end, old_occurrences = merged[-1]
            merged[-1] = (
                old_start,
                max(old_end, end),
                old_occurrences + span_occurrences,
            )
        else:
            merged.append((start, end, list(span_occurrences)))

    regions: list[ContextRegion] = []
    for start, end, region_occurrences in merged:
        snapshots = paragraphs[start : end + 1]
        broad_start = snapshots[0].number
        broad_end = snapshots[-1].number
        unique_occurrences = tuple(
            sorted(
                {occurrence.occurrence_id: occurrence for occurrence in region_occurrences}.values(),
                key=lambda occurrence: (
                    occurrence.paragraph_number,
                    occurrence.start,
                    occurrence.end,
                ),
            )
        )
        region_key = (
            f"{relative_path}\0{source_sha256}\0{query.canonical}\0"
            + "\0".join(occurrence.occurrence_id for occurrence in unique_occurrences)
        )
        region_id = _stable_id("region", region_key)
        regions.append(
            ContextRegion(
                region_id=region_id,
                query=query.text,
                source_relative_path=relative_path,
                source_absolute_path=str(source_path.resolve()),
                source_sha256=source_sha256,
                source_suffix=source_path.suffix.casefold(),
                broad_start_paragraph=broad_start,
                broad_end_paragraph=broad_end,
                paragraphs=snapshots,
                occurrences=unique_occurrences,
                selection=BoundarySelection(
                    start_paragraph=broad_start,
                    end_paragraph=broad_end,
                ),
            )
        )
    return regions


def _context_window(
    paragraphs: Sequence[ParagraphSnapshot],
    occurrence_index: int,
    *,
    context_words_each_side: int,
    max_region_paragraphs: int,
    max_region_characters: int,
) -> tuple[int, int]:
    start = occurrence_index
    end = occurrence_index
    left_words = 0
    right_words = 0
    left_open = True
    right_open = True
    while (
        (left_words < context_words_each_side and left_open)
        or (right_words < context_words_each_side and right_open)
    ):
        sides = (
            ("left", left_words),
            ("right", right_words),
        )
        progressed = False
        for side, _words in sorted(sides, key=lambda item: item[1]):
            if side == "left":
                if left_words >= context_words_each_side or not left_open:
                    continue
                candidate_start, candidate_end = start - 1, end
                if candidate_start < 0:
                    left_open = False
                    continue
                if not _span_within_limits(
                    paragraphs,
                    candidate_start,
                    candidate_end,
                    max_region_paragraphs=max_region_paragraphs,
                    max_region_characters=max_region_characters,
                ):
                    left_open = False
                    continue
                start = candidate_start
                left_words += _word_count(paragraphs[start].text)
                progressed = True
                break
            if right_words >= context_words_each_side or not right_open:
                continue
            candidate_start, candidate_end = start, end + 1
            if candidate_end >= len(paragraphs):
                right_open = False
                continue
            if not _span_within_limits(
                paragraphs,
                candidate_start,
                candidate_end,
                max_region_paragraphs=max_region_paragraphs,
                max_region_characters=max_region_characters,
            ):
                right_open = False
                continue
            end = candidate_end
            right_words += _word_count(paragraphs[end].text)
            progressed = True
            break
        if not progressed and not left_open and not right_open:
            break
    return start, end


def _span_within_limits(
    paragraphs: Sequence[ParagraphSnapshot],
    start: int,
    end: int,
    *,
    max_region_paragraphs: int,
    max_region_characters: int,
) -> bool:
    snapshots = paragraphs[start : end + 1]
    if len(snapshots) > MAX_REGION_PARAGRAPH_SPAN:
        return False
    nonempty = [paragraph for paragraph in snapshots if paragraph.text]
    if len(nonempty) > max_region_paragraphs:
        return False
    if sum(len(paragraph.text) for paragraph in nonempty) > max_region_characters:
        return False
    return _paragraph_map_bytes(nonempty) <= MAX_REGION_PARAGRAPH_MAP_BYTES


def _paragraph_map_bytes(paragraphs: Sequence[ParagraphSnapshot]) -> int:
    """Return endpoint-compatible UTF-8 bytes for the numbered source map."""

    payload = [
        {"number": paragraph.number, "text": paragraph.text}
        for paragraph in paragraphs
    ]
    return len(
        json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode(
            "utf-8"
        )
    )


def _word_count(text: str) -> int:
    return len(re.findall(r"\S+", text))


def _region_sort_key(region: ContextRegion) -> tuple[str, int, str]:
    return (
        region.source_relative_path.casefold(),
        region.broad_start_paragraph,
        region.region_id,
    )


def _stable_id(prefix: str, value: str) -> str:
    digest = hashlib.sha256(value.encode("utf-8")).hexdigest()[:20]
    return f"{prefix}_{digest}"


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _write_text_atomically(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="w",
            encoding="utf-8",
            newline="\n",
            prefix=f".{path.name}.",
            suffix=".tmp",
            dir=path.parent,
            delete=False,
        ) as handle:
            handle.write(text)
            temporary = Path(handle.name)
        os.replace(temporary, path)
        temporary = None
    finally:
        if temporary is not None:
            temporary.unlink(missing_ok=True)


def _save_docx_atomically(document: Document, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="wb",
            prefix=f".{path.stem}.",
            suffix=".tmp.docx",
            dir=path.parent,
            delete=False,
        ) as handle:
            temporary = Path(handle.name)
        document.save(str(temporary))
        Document(str(temporary))
        os.replace(temporary, path)
        temporary = None
        return path
    finally:
        if temporary is not None:
            temporary.unlink(missing_ok=True)


def _set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _configure_compilation_document(document: Document, result: SearchResult) -> None:
    """Apply the narrative_proposal preset plus a compact editorial opening."""

    navy = RGBColor(32, 55, 72)
    body = RGBColor(28, 31, 34)
    blue = RGBColor(46, 116, 181)
    dark_blue = RGBColor(31, 77, 120)

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = body
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, blue, 18, 10),
        "Heading 2": (13, blue, 12, 6),
        "Heading 3": (12, dark_blue, 8, 4),
    }
    for style_name, (size, colour, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = colour
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    metadata_style = document.styles.add_style("Context Source Metadata", WD_STYLE_TYPE.PARAGRAPH)
    metadata_style.font.name = "Calibri"
    metadata_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    metadata_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    metadata_style.font.size = Pt(9.5)
    metadata_style.font.color.rgb = RGBColor(96, 105, 112)
    metadata_style.paragraph_format.space_before = Pt(0)
    metadata_style.paragraph_format.space_after = Pt(4)
    metadata_style.paragraph_format.line_spacing = 1.0

    properties = document.core_properties
    properties.title = f'Context Finder: "{result.query.text}"'
    properties.author = "AudioProcessor Context Finder"
    properties.subject = COMPILATION_MARKER
    properties.keywords = COMPILATION_MARKER

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(f'Context Finder | "{result.query.text}"')
    _set_run_font(run, size=8.5, color=RGBColor(120, 126, 132))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    prefix = footer.add_run("Page ")
    _set_run_font(prefix, size=8.5, color=RGBColor(120, 126, 132))
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    page_run = footer.add_run()
    _set_run_font(page_run, size=8.5, color=RGBColor(120, 126, 132))
    page_run._r.extend((field_begin, instruction, field_end))


def _add_opening(document: Document, result: SearchResult) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.0
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(f'Context Finder: "{result.query.text}"')
    _set_run_font(title_run, size=28, color=RGBColor(32, 55, 72))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle.paragraph_format.line_spacing = 1.0
    subtitle.paragraph_format.keep_with_next = True
    count_text = (
        f"{result.occurrence_count} exact occurrence"
        f"{'s' if result.occurrence_count != 1 else ''} in "
        f"{result.source_count} source document"
        f"{'s' if result.source_count != 1 else ''}, presented as "
        f"{len(result.regions)} distinct context section"
        f"{'s' if len(result.regions) != 1 else ''}."
    )
    count_run = subtitle.add_run(count_text)
    _set_run_font(count_run, size=11, color=RGBColor(96, 105, 112))

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(16)
    note.paragraph_format.line_spacing = 1.1
    note.paragraph_format.keep_with_next = bool(result.regions)
    note_run = note.add_run(
        "Local-source edition. Quoted text is reproduced exactly from the source. "
        "Highlighting marks the requested word or phrase; source links open the "
        "local document on the machine where this compilation was created."
    )
    _set_run_font(note_run, size=9.5, color=RGBColor(96, 105, 112), italic=True)


def _add_region(
    document: Document,
    region: ContextRegion,
    *,
    region_index: int,
    source_region_index: int,
    query_pattern: re.Pattern[str],
    highlight_occurrences: bool,
) -> None:
    source_name = Path(region.source_relative_path).stem
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run(f"{source_name} | Context {source_region_index}")

    locator = _format_locator(region)
    metadata = document.add_paragraph(style="Context Source Metadata")
    metadata.add_run(
        f"Section {region_index} | {locator} | "
        f"{region.occurrence_count} occurrence"
        f"{'s' if region.occurrence_count != 1 else ''}"
    )

    link_paragraph = document.add_paragraph(style="Context Source Metadata")
    label_run = link_paragraph.add_run("Source: ")
    label_run.bold = True
    _add_external_hyperlink(
        link_paragraph,
        Path(region.source_absolute_path).as_uri(),
        region.source_relative_path,
    )
    link_paragraph.paragraph_format.keep_with_next = True

    for snapshot in region.selected_paragraphs:
        if not snapshot.text:
            continue
        paragraph = document.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_highlighted_text(
            paragraph,
            snapshot.text,
            query_pattern if highlight_occurrences else None,
        )


def _format_locator(region: ContextRegion) -> str:
    start = region.selection.start_paragraph
    end = region.selection.end_paragraph
    paragraph_text = f"Paragraph {start}" if start == end else f"Paragraphs {start}-{end}"
    pages = {
        paragraph.page_number
        for paragraph in region.selected_paragraphs
        if paragraph.page_number is not None
    }
    if len(pages) == 1:
        return f"Page {next(iter(pages))} | {paragraph_text}"
    if pages:
        return f"Pages {min(pages)}-{max(pages)} | {paragraph_text}"
    return paragraph_text


def _add_highlighted_text(
    paragraph,
    text: str,
    pattern: re.Pattern[str] | None,
) -> None:
    if pattern is None:
        paragraph.add_run(text)
        return
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        run = paragraph.add_run(match.group(0))
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _add_external_hyperlink(paragraph, target: str, display_text: str) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((colour, underline))
    text_element = OxmlElement("w:t")
    text_element.text = display_text
    run.extend((run_properties, text_element))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _safe_filename_component(value: str) -> str:
    value = re.sub(r"[^\w .'-]+", "_", value, flags=re.UNICODE).strip(" .")
    return value[:80] or "search"


def main(argv: Sequence[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Recursively find an exact word/phrase in DOCX, TXT and Markdown "
            "sources and publish verbatim context excerpts."
        )
    )
    parser.add_argument("folder", type=Path, help="Library folder to search recursively")
    parser.add_argument("query", help="Exact word or phrase (one to three words)")
    parser.add_argument("--output", type=Path, help="Compilation .docx path")
    parser.add_argument("--records", type=Path, help="Resume/audit .jsonl path")
    parser.add_argument(
        "--context-words",
        type=int,
        default=DEFAULT_CONTEXT_WORDS_EACH_SIDE,
        help="Approximate words to capture on each side (default: 500)",
    )
    parser.add_argument(
        "--no-highlight",
        action="store_true",
        help="Do not highlight occurrences in the compilation",
    )
    args = parser.parse_args(argv)

    root = args.folder.expanduser().resolve()
    query = validate_query(args.query)
    safe_query = _safe_filename_component(query.text)
    output = (
        args.output.expanduser().resolve()
        if args.output
        else root.parent / f"{root.name} - Context - {safe_query}.docx"
    )
    records = (
        args.records.expanduser().resolve()
        if args.records
        else output.with_suffix(".jsonl")
    )
    result = find_contexts(
        root,
        query.text,
        options=SearchOptions(context_words_each_side=args.context_words),
        exclude_paths=(output,),
    )
    write_result_records(result, records)
    create_compilation_docx(
        result,
        output,
        highlight_occurrences=not args.no_highlight,
    )
    print(
        f"Found {result.occurrence_count} occurrence(s) in "
        f"{result.source_count} source document(s); wrote {len(result.regions)} "
        f"context section(s)."
    )
    print(f"Compilation: {output}")
    print(f"Resume records: {records}")
    if result.issues:
        print(f"Warning: {len(result.issues)} source document(s) could not be read.")
        for issue in result.issues:
            print(f"  {issue.source_relative_path}: {issue.error_type}: {issue.message}")
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
